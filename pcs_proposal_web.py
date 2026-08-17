from flask import Flask, render_template, request, send_file, redirect, url_for, flash, has_request_context, jsonify, session
from docx2pdf import convert
from docx import Document
import os
import math
import shutil
import datetime
import glob
import hashlib
import html
from openpyxl import load_workbook
from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.utils.datetime import from_excel
from openpyxl.worksheet.datavalidation import DataValidation
from copy import copy as _copy_style
from decimal import Decimal, ROUND_HALF_UP
from email import policy
from email.parser import BytesParser
from email.utils import parseaddr
from functools import lru_cache
import base64
import csv
import json
import re
import tempfile
import uuid

import subprocess
import threading
import shlex
import sys
import time
from urllib.parse import quote, urlsplit, urlunsplit

from roof_intelligence_jobs import (
    DEFAULT_DATA_DIR,
    SUPPORTED_ROOF_TYPES,
    get_job_store,
)
from roof_intelligence_cutover_flags import load_cutover_flags
from proposal_tracking_cutover_flags import load_proposal_tracking_cutover_flags
from pcs_runtime_config import load_runtime_configuration
from roof_report_naming import roof_report_pdf_filename
from pcs_local_settings import (
    google_maps_api_key,
    flask_secret_key,
    report_export_directory,
    remove_google_maps_api_key,
    remove_supabase_configuration,
    save_google_maps_api_key,
    save_report_export_directory,
    save_supabase_configuration,
    supabase_configuration,
)
from contact_store import ContactConfigurationError, ContactStore, ContactStoreError, get_contact_store
from proposal_tracking_store import (
    ProposalContactOrganizationRequired,
    ProposalTrackingStoreError,
    get_proposal_tracking_store,
)
from tenant_context import (
    TenantAuthenticationError,
    current_tenant_context,
    sign_in as tenant_sign_in,
    sign_out as tenant_sign_out,
)
from tenant_settings_store import TenantSettingsStore

APP_FOLDER = os.path.dirname(os.path.abspath(__file__))
RUNTIME_CONFIGURATION = load_runtime_configuration()
APP_VARIANT = RUNTIME_CONFIGURATION.app_variant
APP_IS_BETA = APP_VARIANT == "beta"
MULTI_TENANT_ENABLED = RUNTIME_CONFIGURATION.multi_tenant_enabled
PROPOSAL_STORAGE_MODE = RUNTIME_CONFIGURATION.proposal_storage_mode
PROPOSAL_DATABASE_SOURCE_ENABLED = (
    RUNTIME_CONFIGURATION.proposal_database_source_enabled
)
APP_DISPLAY_NAME = os.environ.get("PCS_APP_DISPLAY_NAME", "PCS Proposal").strip() or "PCS Proposal"
APP_ERROR_LOG = str(DEFAULT_DATA_DIR / "pcs_app_error.log")
os.makedirs(os.path.dirname(APP_ERROR_LOG), exist_ok=True)
ROOF_INTELLIGENCE_PROJECT_DIR = os.environ.get(
    "ROOF_INTELLIGENCE_PROJECT_DIR",
    "/Users/vernabbott/Library/CloudStorage/OneDrive-Personal/Visual Studio/PilotPoint IQ Roof Intelligence Report",
)
ROOF_INTELLIGENCE_SCRIPT = os.path.join(APP_FOLDER, "roof_intelligence_single_address.py")
ROOF_INTELLIGENCE_AREA_SCRIPT = os.path.join(APP_FOLDER, "roof_intelligence_area_batch.py")
ROOF_INTELLIGENCE_USER_KEY = os.environ.get("ROOF_INTELLIGENCE_USER_KEY", "local-user")
HAS_XLWINGS = None
xw = None


def _roof_intelligence_user_key() -> str:
    """Namespace compatibility SQLite records by trusted tenant and user."""
    if not MULTI_TENANT_ENABLED:
        return ROOF_INTELLIGENCE_USER_KEY
    context = current_tenant_context()
    return f"{context.tenant_id}:{context.user_id}"


def _tenant_report_output_paths(user_key: str) -> tuple[str, str]:
    tenant_id = str(user_key or "local").split(":", 1)[0]
    if not re.fullmatch(r"[0-9a-fA-F-]{36}", tenant_id):
        tenant_id = "local"
    configured_root = report_export_directory()
    root = configured_root or str(DEFAULT_DATA_DIR)
    tenant_root = os.path.join(root, "tenants", tenant_id)
    return (
        os.path.join(tenant_root, "roof-intelligence-reports"),
        os.path.join(tenant_root, "roof-intelligence-images"),
    )

def _safe_debug(message: str):
    try:
        stream = getattr(sys, "stdout", None)
        if stream and hasattr(stream, "write"):
            stream.write(f"{message}\n")
            stream.flush()
            return
    except Exception:
        pass

    try:
        with open(APP_ERROR_LOG, "a", encoding="utf-8") as handle:
            handle.write(f"{message}\n")
    except Exception:
        pass

def _log_timing(label: str, start_time: float):
    try:
        elapsed = time.perf_counter() - start_time
        _safe_debug(f"[TIMING] {label}: {elapsed:.3f}s")
    except Exception:
        pass

def _notify_user(message: str, category: str = "warning"):
    try:
        if has_request_context():
            flash(message, category)
    except Exception:
        pass
    _safe_debug(f"[{category.upper()}] {message}")

def _get_xlwings():
    global HAS_XLWINGS, xw
    if HAS_XLWINGS is False:
        return None
    if xw is not None:
        return xw
    try:
        import xlwings as _xw
        xw = _xw
        HAS_XLWINGS = True
        return xw
    except Exception as exc:
        HAS_XLWINGS = False
        _safe_debug(f"[WARNING] xlwings unavailable: {exc}")
        return None

def _run_background_task(task_name: str, func):
    def _worker():
        started = time.perf_counter()
        try:
            func()
            _log_timing(f"{task_name} (background)", started)
        except Exception as exc:
            _safe_debug(f"[ERROR] Background task '{task_name}' failed: {exc}")

    threading.Thread(target=_worker, daemon=True).start()

def _read_defaults_flag(domain: str, key: str) -> bool | None:
    try:
        result = subprocess.run(
            ["defaults", "read", domain, key],
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.DEVNULL,
            text=True,
        )
        raw_value = (result.stdout or "").strip().lower()
        if raw_value in {"1", "true", "yes"}:
            return True
        if raw_value in {"0", "false", "no"}:
            return False
    except Exception:
        pass
    return None

def _is_running_new_outlook() -> bool:
    for key in ("RunningNewOutlook", "IsRunningNewOutlook"):
        value = _read_defaults_flag("com.microsoft.Outlook", key)
        if value is not None:
            return value
    return False

def _format_currency(value):
    try:
        return f"${float(value or 0):,.2f}"
    except Exception:
        return "$0.00"

def _format_square_count(value):
    try:
        return f"{int(round(float(value or 0))):,}"
    except Exception:
        return "0"

# Centralized Profit Summary formulas (expanded import list)
from profit_summary_formulas import (
    PS_F_M3,   # 10-year price per sq (M3)
    PS_F_M5,   # 15-year price per sq (M5)
    PS_F_M7,   # 20-year price per sq (M7)
    PS_F_C11,  # silicone_units_10
    PS_F_C12,  # gaco_patch_units
    PS_F_C13,  # bleed_trap_units
    PS_F_C14,  # gaco_e5320_units
    PS_F_C15,  # sw_1flash_units
    PS_F_C16,  # sw_bleed_block_units
    PS_F_C17,  # drainage_mat_units
    PS_F_C18,  # foam_units
    PS_F_D11, PS_F_D12, PS_F_D13, PS_F_D14, PS_F_D15, PS_F_D16, PS_F_D17, PS_F_D18, PS_F_D19, PS_F_D21,
    PS_F_I14, PS_F_O14, PS_F_T14,
    PS_F_H11, PS_F_K11, PS_F_N11, PS_F_P11,  # silicone price/units/totals by term
    PS_F_E11, PS_F_E12, PS_F_E13, PS_F_E15, PS_F_E16, PS_F_E17, PS_F_E18,  # line totals
    PS_F_E19, PS_F_E21,  # labor totals
    PS_F_E24,            # warranty total
    PS_F_P3, PS_F_P5, PS_F_P7,  # total prices by term
    PS_F_E25, PS_F_K25, PS_F_P25, PS_F_U25, PS_F_E27, PS_F_K27, PS_F_P27, PS_F_U27,
    PS_F_E29, PS_F_K29, PS_F_P29, PS_F_E30, PS_F_K30, PS_F_P30,
    PS_F_E31, PS_F_K31, PS_F_P31, PS_F_E32, PS_F_K32, PS_F_P32,
    PS_F_E33, PS_F_K33, PS_F_P33, PS_F_U33,  # fees/profit/commission
)

# Optional centralized formula for labor_days (E7). Falls back to inline if not present.
try:
    from profit_summary_formulas import PS_F_E7 as _PS_F_E7_LABOR
except Exception:
    _PS_F_E7_LABOR = None
import pathlib, traceback

# Directory constants (editable in one place)
PROPOSAL_TEMP_DIR = os.environ.get("PCS_PROPOSAL_TEMP_DIR", "/Users/vernabbott/Library/CloudStorage/OneDrive-ProfessionalCoatingSystems/Test Site/1. Open Proposals")
CONTRACTS_DIR = os.environ.get("PCS_CONTRACTS_DIR", "/Users/vernabbott/Library/CloudStorage/OneDrive-ProfessionalCoatingSystems/Test Site/2. Signed Contracts")
COMPLETED_DIR = os.environ.get("PCS_COMPLETED_DIR", "/Users/vernabbott/Library/CloudStorage/OneDrive-ProfessionalCoatingSystems/Test Site/3. Finished Jobs")
DEADFILE_DIR = os.environ.get("PCS_DEADFILE_DIR", "/Users/vernabbott/Library/CloudStorage/OneDrive-ProfessionalCoatingSystems/Test Site/4. Dead Proposals")
TEMPLATE_DIR = os.environ.get("PCS_TEMPLATE_DIR", "/Users/vernabbott/Library/CloudStorage/OneDrive-ProfessionalCoatingSystems/Test Site/Job Jacket Template")
LIBREOFFICE_PATH = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
PCS_PROPOSALS_DIR = os.environ.get("PCS_PROPOSALS_DIR", "/Users/vernabbott/Library/CloudStorage/OneDrive-ProfessionalCoatingSystems/PCS/1 - Open Proposals")
DAVIDS_PROPOSALS_DIR = os.environ.get("PCS_DAVIDS_PROPOSALS_DIR", "/Users/vernabbott/Library/CloudStorage/OneDrive-ProfessionalCoatingSystems/PCS/David's Accounts/1 - Open Proposals")
LYDIAS_PROPOSALS_DIR = os.environ.get("PCS_LYDIAS_PROPOSALS_DIR", "/Users/vernabbott/Library/CloudStorage/OneDrive-ProfessionalCoatingSystems/PCS/Lydia's Accounts/1 - Open Proposals")
RANDYS_PROPOSALS_DIR = os.environ.get("PCS_RANDYS_PROPOSALS_DIR", "/Users/vernabbott/Library/CloudStorage/OneDrive-ProfessionalCoatingSystems/PCS/Randy's Accounts/1 - Open Proposals")
OPEN_PROPOSAL_DIRS = (
    PCS_PROPOSALS_DIR,
    DAVIDS_PROPOSALS_DIR,
    LYDIAS_PROPOSALS_DIR,
    RANDYS_PROPOSALS_DIR,
)

# Optional OneDrive/SharePoint folder roots for email hyperlinks.
# These default to the web locations for the synced work OneDrive folders so
# sent email links are usable by other PCS users instead of local file:// links.
PCS_PROPOSALS_WEB_URL = os.environ.get(
    "PCS_PROPOSALS_WEB_URL",
    "https://procoatingsystems-my.sharepoint.com/personal/admin_procoatingsystems_onmicrosoft_com/Documents/PCS/1%20-%20Open%20Proposals",
).strip()
DAVIDS_PROPOSALS_WEB_URL = os.environ.get(
    "DAVIDS_PROPOSALS_WEB_URL",
    "https://procoatingsystems-my.sharepoint.com/personal/admin_procoatingsystems_onmicrosoft_com/Documents/PCS/David%27s%20Accounts/1%20-%20Open%20Proposals",
).strip()
LYDIAS_PROPOSALS_WEB_URL = os.environ.get(
    "LYDIAS_PROPOSALS_WEB_URL",
    "https://procoatingsystems-my.sharepoint.com/personal/admin_procoatingsystems_onmicrosoft_com/Documents/PCS/Lydia%27s%20Accounts/1%20-%20Open%20Proposals",
).strip()
RANDYS_PROPOSALS_WEB_URL = os.environ.get(
    "RANDYS_PROPOSALS_WEB_URL",
    "https://procoatingsystems-my.sharepoint.com/personal/admin_procoatingsystems_onmicrosoft_com/Documents/PCS/Randy%27s%20Accounts/1%20-%20Open%20Proposals",
).strip()
PROPOSAL_SUMMARY_TEMPLATE_PATH = os.environ.get("PCS_PROPOSAL_SUMMARY_TEMPLATE_PATH", "/Users/vernabbott/Library/CloudStorage/OneDrive-Personal/1. Proposal Summary Template.emltpl")
OUTLOOK_SENDER_EMAIL = "vern@procoatingsystems.com"
EMAIL_TEMPLATE_DIR = os.environ.get("PCS_EMAIL_TEMPLATE_DIR", "/Users/vernabbott/Library/CloudStorage/OneDrive-ProfessionalCoatingSystems/PCS/Marketing/Email Templates")
EMAIL_LIST_DIR = os.environ.get("PCS_EMAIL_LIST_DIR", "/Users/vernabbott/Library/CloudStorage/OneDrive-ProfessionalCoatingSystems/PCS/Marketing/Email Lists")
REPAIR_COSTS_PROPOSAL_LANGUAGE = "*PCS will perform all necessary repairs to bring the roof to coating ready*"

PROPOSAL_TRACKER = os.environ.get("PCS_PROPOSAL_TRACKER", "/Users/vernabbott/Library/CloudStorage/OneDrive-ProfessionalCoatingSystems/PCS/1 - Open Proposals/Proposal Tracking.xlsx")
TRACKER_IO_LOCK = threading.RLock()

_PROPOSAL_TRACKER_CANONICAL_HEADERS = (
    "Customer",
    "Contact",
    "Email Address",
    "Lead Generated",
    "Submitted By",
    "Estimate Dt",
    "Proposal Dt",
    "Follow-Up",
    "Status",
    "Estimated By",
    "Response",
)
_PROPOSAL_TRACKER_DEFAULT_COLUMNS = {
    # Legacy workbooks used this order. Header names override every fallback.
    "customer": 1,
    "contact": 2,
    "email_address": 3,
    "lead_source": 4,
    "submitted_by": 5,
    "proposal_date": 6,
    "follow_up_date": 7,
    "estimated_by": 8,
    "response": 9,
    "estimate_date": 10,
    "status": 11,
}
_PROPOSAL_TRACKER_HEADER_ALIASES = {
    "customer": ("Customer", "Proposal Folder Name", "Proposal Name"),
    "contact": ("Contact", "Contact Name"),
    "email_address": ("Email Address", "Email"),
    "lead_source": ("Lead Generated", "Lead", "Lead Source"),
    "submitted_by": ("Submitted By", "Submitted"),
    "estimate_date": ("Estimate Dt", "Estimate Date", "Estimate Completed Date"),
    "proposal_date": ("Proposal Dt", "Proposal Date", "Proposal Sent Date"),
    "follow_up_date": ("Follow-Up", "Follow Up", "Follow-Up Date", "Follow Up Date"),
    "status": ("Status", "Proposal Status"),
    "estimated_by": ("Estimated By", "Estimator", "Vern"),
    "response": ("Response", "Response Notes"),
}


def _normalize_proposal_tracker_header(value):
    return re.sub(r"[^a-z0-9]+", "", str(value or "").strip().casefold())


def _proposal_tracker_column_map_from_headers(headers):
    indexed_headers = {}
    for column_number, header in enumerate(headers or (), start=1):
        normalized = _normalize_proposal_tracker_header(header)
        if normalized and normalized not in indexed_headers:
            indexed_headers[normalized] = column_number

    columns = {}
    for field_name, aliases in _PROPOSAL_TRACKER_HEADER_ALIASES.items():
        for alias in aliases:
            column_number = indexed_headers.get(
                _normalize_proposal_tracker_header(alias)
            )
            if column_number is not None:
                columns[field_name] = column_number
                break
        columns.setdefault(
            field_name, _PROPOSAL_TRACKER_DEFAULT_COLUMNS[field_name]
        )
    return columns


def _proposal_tracker_column_map(ws):
    max_column = max(getattr(ws, "max_column", 0) or 0, 11)
    headers = [
        ws.cell(row=1, column=column_number).value
        for column_number in range(1, max_column + 1)
    ]
    return _proposal_tracker_column_map_from_headers(headers)


def _initialize_proposal_tracker_headers(ws):
    for column_number, header in enumerate(
        _PROPOSAL_TRACKER_CANONICAL_HEADERS, start=1
    ):
        ws.cell(row=1, column=column_number).value = header


def _ensure_proposal_tracker_status_column(ws):
    headers = [
        ws.cell(row=1, column=column_number).value
        for column_number in range(1, (ws.max_column or 0) + 1)
    ]
    normalized = {
        _normalize_proposal_tracker_header(value): column_number
        for column_number, value in enumerate(headers, start=1)
        if _normalize_proposal_tracker_header(value)
    }
    if _normalize_proposal_tracker_header("Status") in normalized:
        return

    follow_up_column = None
    for alias in _PROPOSAL_TRACKER_HEADER_ALIASES["follow_up_date"]:
        follow_up_column = normalized.get(_normalize_proposal_tracker_header(alias))
        if follow_up_column is not None:
            break
    insert_column = (follow_up_column + 1) if follow_up_column else (ws.max_column + 1)
    ws.insert_cols(insert_column)
    style_column = insert_column + 1 if insert_column + 1 <= ws.max_column else insert_column - 1
    for row_number in range(1, ws.max_row + 1):
        source = ws.cell(row=row_number, column=style_column)
        target = ws.cell(row=row_number, column=insert_column)
        if source.has_style:
            target._style = _copy_style(source._style)
        target.number_format = source.number_format
    ws.cell(row=1, column=insert_column).value = "Status"


def _proposal_tracker_row_value(row, column_number):
    index = int(column_number) - 1
    return row[index] if 0 <= index < len(row) else None


def _proposal_tracker_previous_path(tracker_path):
    base, ext = os.path.splitext(tracker_path)
    return f"{base} (previous){ext}"


def _proposal_tracker_source_path(tracker_path):
    if tracker_path and os.path.exists(tracker_path):
        return tracker_path
    prev_path = _proposal_tracker_previous_path(tracker_path)
    if os.path.exists(prev_path):
        return prev_path
    return None


def _proposal_tracker_temp_path(tracker_path):
    tracker_dir = os.path.dirname(tracker_path)
    os.makedirs(tracker_dir, exist_ok=True)
    _, ext = os.path.splitext(tracker_path)
    fd, temp_path = tempfile.mkstemp(
        prefix=".proposal-tracking-",
        suffix=ext or ".xlsx",
        dir=tracker_dir,
    )
    os.close(fd)
    return temp_path


def _replace_proposal_tracker_file(temp_path, tracker_path):
    prev_path = _proposal_tracker_previous_path(tracker_path)
    if tracker_path and os.path.exists(tracker_path):
        shutil.copy2(tracker_path, prev_path)
    os.replace(temp_path, tracker_path)

# ---- Roof types and base pricing/coverage matrices ----
roof_types = ["TPO/EPDM", "Metal", "Mod Bit", "Ballasted 60 mil", "Ballasted 45 mil", "Rock/Foam/Coat"]

# Pricing arrays moved here for global access
 
# Roofer-facing pricing (default)
roofer_pricing10 = [320, 325, 330, 470, 575, 670]
roofer_pricing15 = [360, 365, 370, 510, 615, 710]
roofer_pricing20 = [400, 405, 410, 550, 655, 750]

# PCS Direct pricing (initially identical unless adjusted later)
pcs_pricing10 = [350, 355, 360, 500, 605, 700]
pcs_pricing15 = [390, 395, 400, 540, 645, 740]
pcs_pricing20 = [430, 435, 440, 580, 685, 780]

# Active pricing arrays used throughout the system — these will be switched at runtime

coverage_amounts = {
    "Gaco": {
        "TPO/EPDM":    {10: 1.25, 15: 1.75, 20: 2.25},
        "Metal":       {10: 1.35, 15: 1.85, 20: 2.35},
        "Mod Bit":     {10: 1.25, 15: 1.75, 20: 2.25},
        "Ballasted 60 mil": {10: 2.5, 15: 3.25, 20: 3.75},
        "Ballasted 45 mil": {10: 3.0, 15: 4.5,  20: 5.5},
        "Rock/Foam/Coat":   {10: 1.25, 15: 1.75, 20: 2.25},
    },
    "Uniflex": {
        "TPO/EPDM":    {10: 1.25, 15: 1.75, 20: 2.25},
        "Metal":       {10: 1.35, 15: 1.85, 20: 2.35},
        "Mod Bit":     {10: 1.25, 15: 1.75, 20: 2.25},
        "Ballasted 60 mil": {10: 2.5, 15: 3.25, 20: 3.75},
        "Ballasted 45 mil": {10: 3.0, 15: 4.5,  20: 5.5},
        "Rock/Foam/Coat":   {10: 1.25, 15: 1.75, 20: 2.25},
    }
}

# Base prices
PCS_BASE_LABOR_RATE = 3250
TRAVEL_GAS_PER_JOB = 250
TRAVEL_HOTEL_PER_NIGHT = 175
TRAVEL_ROOMS_PER_NIGHT = 6
TRAVEL_FOOD_PER_DAY = 700
TRAVEL_MISC_500 = 500
TRAVEL_MISC_250 = 250
GACO_S42_BASE_PRICE = 190
GACO_PATCH_BASE_PRICE = 125
GACO_E5320_PRICE = 185
BLEED_TRAP_BASE_PRICE = 168
DRAINAGE_MAT_BASE_PRICE = 164
UNIFLEX_BASE_PRICE = 185
SW_1FLASH_BASE_PRICE = 110
SW_BLEED_BLOCK_BASE_PRICE = 100
GACO_FOAM_BASE_PRICE = 2600
UNIFLEX_FOAM_BASE_PRICE = 2600
RFC_LABOR_RATE = 250
BASE_OFFICE_FEE_PCT = 0.05
SALES_STAFF_OFFICE_FEE_PCT = 0.05
PROFIT_SHARE_PCT = 0.10
COMMISSION_PCT = 0.10

def office_fee_pct_for_submitter(submitted_by):
    return BASE_OFFICE_FEE_PCT if str(submitted_by or "").strip() == "Mark" else SALES_STAFF_OFFICE_FEE_PCT

def commission_pct_for_submitter(submitted_by):
    return 0.0 if str(submitted_by or "").strip() in ("Mark", "Richard") else COMMISSION_PCT

def _append_to_proposal_tracking_xlwings(tracker_path,
                                         folder_name,
                                         total_squares,
                                         lead_value,
                                         submitted_by,
                                         estimate_completed_date=""):
    app = None
    wb = None
    temp_path = None
    xw_module = _get_xlwings()
    if xw_module is None:
        raise RuntimeError("xlwings is not available")
    try:
        app = xw_module.App(visible=False, add_book=False)
        # Open existing workbook; if missing, create a new one with a header row
        source_path = _proposal_tracker_source_path(tracker_path)
        if source_path:
            wb = xw_module.Book(source_path)
            ws = wb.sheets[0]
        else:
            wb = xw_module.Book()
            ws = wb.sheets[0]
            ws.name = "Tracking"
            ws.range("A1:J1").value = [list(_PROPOSAL_TRACKER_CANONICAL_HEADERS)]
        header_values = ws.range("A1:J1").value or []
        if header_values and isinstance(header_values[0], list):
            header_values = header_values[0]
        columns = _proposal_tracker_column_map_from_headers(header_values)

        # Calculate insertion point by scanning A2..last
        used = ws.used_range
        last_row = used.last_cell.row if used is not None else 1
        first_data_row = 2
        if last_row < first_data_row:
            last_row = first_data_row - 1

        new_key = (str(folder_name).strip().lower() if folder_name is not None else "")
        insert_at = last_row + 1  # default append
        existing_row = None
        for r in range(first_data_row, last_row + 1):
            a_val = ws.range((r, columns["customer"])).value
            a_key = (str(a_val).strip().lower() if a_val is not None else "")
            if a_key == new_key:
                existing_row = r
                break
            if insert_at == (last_row + 1) and a_key > new_key:
                insert_at = r

        if existing_row is not None:
            target_row = existing_row
        else:
            # Insert a new row at insert_at. Excel shifts formats down automatically.
            ws.api.Rows(insert_at).Insert()
            # Copy style from the row that was previously at insert_at (now at insert_at+1) when present
            try:
                if insert_at + 1 <= ws.used_range.last_cell.row:
                    ws.range(f"{insert_at+1}:{insert_at+1}").api.Copy()
                    ws.range(f"{insert_at}:{insert_at}").api.PasteSpecial(Paste=-4122)  # xlPasteFormats
            except Exception:
                pass
            target_row = insert_at

        values = {
            "customer": folder_name or "",
            "lead_source": lead_value or "",
            "submitted_by": submitted_by or "",
            "estimated_by": "Vern",
            "estimate_date": estimate_completed_date or "",
        }
        for field_name, value in values.items():
            ws.range((target_row, columns[field_name])).value = value

        temp_path = _proposal_tracker_temp_path(tracker_path)
        wb.save(temp_path)
        wb.close()
        wb = None
        _replace_proposal_tracker_file(temp_path, tracker_path)
        temp_path = None
    finally:
        try:
            if wb is not None:
                wb.close()
        except Exception:
            pass
        try:
            if app is not None:
                app.quit()
        except Exception:
            pass
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass


def _append_to_proposal_tracking_openpyxl_simple(
    tracker_path,
    folder_name,
    lead_value,
    submitted_by,
    estimate_completed_date="",
):
    source_path = _proposal_tracker_source_path(tracker_path)
    wb = None
    temp_path = None
    try:
        if source_path:
            wb = load_workbook(source_path)
            ws = wb.active
        else:
            wb = Workbook()
            ws = wb.active
        if ws is None:
            raise RuntimeError("Proposal tracker workbook does not have an active worksheet")
        if not source_path:
            ws.title = "Tracking"
            _initialize_proposal_tracker_headers(ws)
        columns = _proposal_tracker_column_map(ws)

        first_data_row = 2
        last_row = max(ws.max_row or 1, 1)
        new_key = str(folder_name or "").strip().lower()
        existing_row = None
        insert_at = last_row + 1

        for row_idx in range(first_data_row, last_row + 1):
            row_key = str(
                ws.cell(row=row_idx, column=columns["customer"]).value or ""
            ).strip().lower()
            if row_key == new_key:
                existing_row = row_idx
                break
            if row_key and row_key > new_key and insert_at == last_row + 1:
                insert_at = row_idx

        target_row = existing_row or insert_at
        if existing_row is None and insert_at <= last_row:
            ws.insert_rows(insert_at)

        values = {
            "customer": folder_name or "",
            "lead_source": lead_value or "",
            "submitted_by": submitted_by or "",
            "estimated_by": "Vern",
            "estimate_date": estimate_completed_date or "",
        }
        for field_name, value in values.items():
            ws.cell(row=target_row, column=columns[field_name]).value = value

        temp_path = _proposal_tracker_temp_path(tracker_path)
        wb.save(temp_path)
        wb.close()
        wb = None
        _replace_proposal_tracker_file(temp_path, tracker_path)
        temp_path = None
    finally:
        try:
            if wb is not None:
                wb.close()
        except Exception:
            pass
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass


# Helper to append a row to the Proposal Tracking workbook
def _append_to_proposal_tracking_unlocked(created_date,
                                          customer_name,
                                          street_address,
                                          city,
                                          state,
                                          zip_code,
                                          product,
                                          roof_type,
                                          total_squares,
                                          warranty_incl,
                                          submitted_by,
                                          folder_name,
                                          proposal_folder,
                                          tp10,
                                          tp15,
                                          tp20,
                                          lead_value=""):
    """Append or update one tracking row while preserving a recoverable previous copy.

    Strategy:
      1) Read the current tracker (or previous backup if the live file is missing)
      2) Create a new workbook in memory
      3) Copy every row from the source workbook to the new one **unchanged** (values + styles)
      4) Update the matching row or insert a new row in alphabetical order by Column A
      5) Save to a temp file and atomically replace the live tracker
    """
    try:
        _append_to_proposal_tracking_openpyxl_simple(
            PROPOSAL_TRACKER,
            folder_name,
            lead_value,
            submitted_by,
            created_date,
        )
        return
    except Exception as _simple_err:
        try:
            with open(APP_ERROR_LOG, "a", encoding="utf-8") as _f:
                _f.write(f"[ERROR] simple tracking append failed for {folder_name}: {_simple_err}\n")
        except Exception:
            pass

    # Prefer Excel-native insert (xlwings) to fully preserve all formatting. Fallback to openpyxl rebuild if unavailable.
    if _get_xlwings() is not None:
        try:
            _append_to_proposal_tracking_xlwings(
                PROPOSAL_TRACKER,
                folder_name,
                total_squares,
                lead_value,
                submitted_by,
                created_date,
            )
            return
        except Exception as _xw_err:
            try:
                with open(APP_ERROR_LOG, "a", encoding="utf-8") as _f:
                    _f.write(f"[WARN] xlwings path failed, falling back to openpyxl: {_xw_err}\n")
            except Exception:
                pass
    try:
        # Ensure directory exists
        os.makedirs(os.path.dirname(PROPOSAL_TRACKER), exist_ok=True)

        tracker_path = PROPOSAL_TRACKER
        source_path = _proposal_tracker_source_path(tracker_path)

        # Load or create a previous file snapshot
        prev_wb = None
        prev_ws = None
        new_wb = None
        temp_path = None
        if source_path:
            prev_wb = load_workbook(source_path)
            prev_ws = prev_wb.active

        # Create the new workbook (fresh)
        new_wb = Workbook()
        new_ws = new_wb.active
        if new_ws is None:
            raise RuntimeError("Proposal tracker workbook does not have an active worksheet")
        new_ws.title = "Tracking"
        if prev_ws is None:
            _initialize_proposal_tracker_headers(new_ws)
        # Preserve column widths, row heights, and freeze panes from previous sheet (if any)
        prev_freeze = None
        if prev_ws is not None:
            try:
                # Column widths
                for key, dim in prev_ws.column_dimensions.items():
                    if dim.width is not None:
                        new_ws.column_dimensions[key].width = dim.width
                # Row heights
                for idx, rdim in prev_ws.row_dimensions.items():
                    if rdim.height is not None:
                        new_ws.row_dimensions[idx].height = rdim.height
                # Freeze panes
                prev_freeze = getattr(prev_ws, 'freeze_panes', None)
                if prev_freeze:
                    new_ws.freeze_panes = prev_freeze
            except Exception:
                pass

        # Also preserve merged cells, conditional formatting, and named styles when available
        if prev_ws is not None:
            try:
                # Merged cells
                merged_ranges = getattr(getattr(prev_ws, 'merged_cells', None), 'ranges', [])
                for mc in merged_ranges:
                    try:
                        new_ws.merge_cells(str(mc))
                    except Exception:
                        pass

                # Conditional formatting rules
                try:
                    cf = getattr(prev_ws, 'conditional_formatting', None)
                    if cf is not None and getattr(cf, 'cf_rules', None):
                        for rng, rules in cf.cf_rules.items():
                            for rule in rules:
                                try:
                                    new_ws.conditional_formatting.add(rng, rule)
                                except Exception:
                                    pass
                except Exception:
                    pass

                # Named styles (fonts, fills, borders referenced by name)
                try:
                    prev_parent = prev_ws.parent
                    new_parent = new_ws.parent
                    if (
                        prev_parent is not None
                        and new_parent is not None
                        and hasattr(prev_parent, 'named_styles')
                        and hasattr(new_parent, 'named_styles')
                    ):
                        existing = {
                            ns if isinstance(ns, str) else getattr(ns, "name", None)
                            for ns in new_parent.named_styles
                        }
                        for ns in prev_parent.named_styles:
                            try:
                                if isinstance(ns, str):
                                    continue
                                style_name = ns if isinstance(ns, str) else getattr(ns, "name", None)
                                if style_name not in existing:
                                    new_parent.add_named_style(ns)
                            except Exception:
                                pass
                except Exception:
                    pass
            except Exception:
                pass

        # Helper to copy a single cell's value and style
        def _copy_cell(src, dst):
            try:
                dst.value = src.value
                # Deep-copy style objects so bold/color/etc. are preserved reliably
                try:
                    dst.font = _copy_style(src.font)
                    dst.fill = _copy_style(src.fill)
                    dst.border = _copy_style(src.border)
                    dst.alignment = _copy_style(src.alignment)
                    dst.protection = _copy_style(src.protection)
                    # number_format is a plain string in most cases
                    dst.number_format = src.number_format
                except Exception:
                    # Fallback to direct assignment
                    dst.font = src.font
                    dst.fill = src.fill
                    dst.border = src.border
                    dst.alignment = src.alignment
                    dst.protection = src.protection
                    dst.number_format = src.number_format
            except Exception:
                dst.value = src.value

        # If we have a previous sheet, copy everything over unchanged
        max_col = 11  # at least A..K
        max_row_prev = 0
        if prev_ws is not None:
            max_col = max(max_col, prev_ws.max_column or 11)
            max_row_prev = prev_ws.max_row or 0
            for r in range(1, max_row_prev + 1):
                for c in range(1, max_col + 1):
                    _copy_cell(prev_ws.cell(row=r, column=c), new_ws.cell(row=r, column=c))
        _ensure_proposal_tracker_status_column(new_ws)
        columns = _proposal_tracker_column_map(new_ws)
        row_values = {
            "customer": folder_name or "",
            "lead_source": lead_value or "",
            "submitted_by": submitted_by or "",
            "estimated_by": "Vern",
            "estimate_date": created_date or "",
            "status": "Draft",
        }

        # Determine insertion point by Column A (case-insensitive), preserving header at row 1
        first_data_row = 2
        insert_at = max(first_data_row, max_row_prev + 1)  # default: append at bottom
        existing_row = None
        new_key = (str(folder_name).strip().lower() if folder_name is not None else "")

        # Scan existing data rows only from the *new* worksheet copy
        if max_row_prev >= first_data_row:
            for r in range(first_data_row, max_row_prev + 1):
                a_val = new_ws.cell(row=r, column=columns["customer"]).value
                a_key = (str(a_val).strip().lower() if a_val is not None else "")
                if a_key == new_key:
                    existing_row = r
                    break
                if a_key > new_key:
                    insert_at = r
                    break

        # Shift rows down by 1 from bottom to insert_at (A..K only) to make space
        if existing_row is None and insert_at <= max(max_row_prev, first_data_row - 1):
            for rr in range(max_row_prev, insert_at - 1, -1):
                for c in range(1, 12):  # A..K only
                    src = new_ws.cell(row=rr, column=c)
                    dst = new_ws.cell(row=rr + 1, column=c)
                    # Move value + full style (deep copy) for A..K to retain bold/color/etc.
                    dst.value = src.value
                    try:
                        dst.font = _copy_style(src.font)
                        dst.fill = _copy_style(src.fill)
                        dst.border = _copy_style(src.border)
                        dst.alignment = _copy_style(src.alignment)
                        dst.protection = _copy_style(src.protection)
                        dst.number_format = src.number_format
                    except Exception:
                        dst.font = src.font
                        dst.fill = src.fill
                        dst.border = src.border
                        dst.alignment = src.alignment
                        dst.protection = src.protection
                        dst.number_format = src.number_format

        # Choose a template row to clone styles from: prefer the row now at insert_at+1 (the one we shifted down)
        if existing_row is not None:
            template_row = existing_row
            insert_target = existing_row
        elif insert_at + 1 <= (max_row_prev + 1):
            template_row = insert_at + 1
            insert_target = insert_at
        elif max_row_prev >= first_data_row:
            template_row = first_data_row
            insert_target = insert_at
        else:
            template_row = insert_at
            insert_target = insert_at

        for c in range(1, 12):
            cell = new_ws.cell(row=insert_target, column=c)
            if existing_row is None:
                cell.value = None
            # Apply styles from template row to preserve fonts/sizes/bold/color
            try:
                tmpl = new_ws.cell(row=template_row, column=c)
                cell.font = _copy_style(tmpl.font)
                cell.fill = _copy_style(tmpl.fill)
                cell.border = _copy_style(tmpl.border)
                cell.alignment = _copy_style(tmpl.alignment)
                cell.protection = _copy_style(tmpl.protection)
                cell.number_format = tmpl.number_format
            except Exception:
                try:
                    cell.font = tmpl.font
                    cell.fill = tmpl.fill
                    cell.border = tmpl.border
                    cell.alignment = tmpl.alignment
                    cell.protection = tmpl.protection
                    cell.number_format = tmpl.number_format
                except Exception:
                    pass

        for field_name, value in row_values.items():
            new_ws.cell(
                row=insert_target, column=columns[field_name]
            ).value = value

        # Save the new tracker atomically so the live file is never half-written.
        temp_path = _proposal_tracker_temp_path(tracker_path)
        new_wb.save(temp_path)
        new_wb.close()
        new_wb = None
        if prev_wb is not None:
            prev_wb.close()
            prev_wb = None
        _replace_proposal_tracker_file(temp_path, tracker_path)
        temp_path = None
    except Exception as e:
        try:
            with open(APP_ERROR_LOG, "a", encoding="utf-8") as _f:
                _f.write(f"[ERROR] append_to_proposal_tracking (rebuild) failed: {e}\n")
        except Exception:
            pass
    finally:
        try:
            if prev_wb is not None:
                prev_wb.close()
        except Exception:
            pass
        try:
            if new_wb is not None:
                new_wb.close()
        except Exception:
            pass
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass


def _append_to_proposal_tracking_spreadsheet(created_date,
                                customer_name,
                                street_address,
                                city,
                                state,
                                zip_code,
                                product,
                                roof_type,
                                total_squares,
                                warranty_incl,
                                submitted_by,
                                folder_name,
                                proposal_folder,
                                tp10,
                                tp15,
                                tp20,
                                lead_value=""):
    with TRACKER_IO_LOCK:
        _append_to_proposal_tracking_unlocked(
            created_date=created_date,
            customer_name=customer_name,
            street_address=street_address,
            city=city,
            state=state,
            zip_code=zip_code,
            product=product,
            roof_type=roof_type,
            total_squares=total_squares,
            warranty_incl=warranty_incl,
            submitted_by=submitted_by,
            folder_name=folder_name,
            proposal_folder=proposal_folder,
            tp10=tp10,
            tp15=tp15,
            tp20=tp20,
            lead_value=lead_value,
        )


def append_to_proposal_tracking(created_date,
                                customer_name,
                                street_address,
                                city,
                                state,
                                zip_code,
                                product,
                                roof_type,
                                total_squares,
                                warranty_incl,
                                submitted_by,
                                folder_name,
                                proposal_folder,
                                tp10,
                                tp15,
                                tp20,
                                lead_value=""):
    flags = load_proposal_tracking_cutover_flags()
    supabase_error = None
    if flags.writes_enabled:
        try:
            get_proposal_tracking_store().upsert_from_proposal_save(
                created_date=created_date,
                customer_name=customer_name,
                street_address=street_address,
                city=city,
                state=state,
                zip_code=zip_code,
                submitted_by=submitted_by,
                folder_name=folder_name,
                lead_value=lead_value,
            )
        except Exception as exc:
            supabase_error = exc
            _safe_debug(f"[ERROR] Supabase proposal tracking write failed: {exc}")
    if flags.spreadsheet_writes_active:
        _append_to_proposal_tracking_spreadsheet(
            created_date=created_date,
            customer_name=customer_name,
            street_address=street_address,
            city=city,
            state=state,
            zip_code=zip_code,
            product=product,
            roof_type=roof_type,
            total_squares=total_squares,
            warranty_incl=warranty_incl,
            submitted_by=submitted_by,
            folder_name=folder_name,
            proposal_folder=proposal_folder,
            tp10=tp10,
            tp15=tp15,
            tp20=tp20,
            lead_value=lead_value,
        )
    elif supabase_error is not None:
        raise supabase_error

def create_proposal_from_fields(customer_name,
                                street_address,
                                city,
                                state,
                                zip_code,
                                roof_type,
                                total_squares,
                                warranty_incl,
                                product,
                                proposal_language,
                                submitted_by,
                                target_folder: str | None = None,
                                mapped_data: dict | None = None,
                                pdf_async: bool = True,
                                use_libreoffice: bool = True,
                                update_tracking: bool = True,
                                copy_destination: bool = True,
                                create_email_draft: bool = True):
    
    today = datetime.date.today()
    formatted_date = today.strftime("%m/%d/%Y")

    # Determine placeholder values based on warranty and product
    if product == "Gaco":
        if warranty_incl == "Yes":
            price_includes_text = "* Price Includes material, labor, trash pickup, haul away and Gaco Warranty Fee"
            warranty_text = "IS INCLUDED"
        else:
            price_includes_text = "* Price Includes material, labor, trash pickup and haul away"
            warranty_text = "IS NOT INCLUDED"
    else:
        price_includes_text = proposal_language if proposal_language else ' '
        warranty_text = warranty_incl

    # Use computed totals passed from the UI/calculation_routine (authoritative)
    # Fall back to 0 if not provided.
    tp10 = 0
    tp15 = 0
    tp20 = 0
    if mapped_data:
        try:
            tp10 = float(mapped_data.get("total_price_10", 0) or 0)
            tp15 = float(mapped_data.get("total_price_15", 0) or 0)
            tp20 = float(mapped_data.get("total_price_20", 0) or 0)
        except Exception:
            tp10, tp15, tp20 = 0, 0, 0

    # Use full submitter names in the Word document while keeping short names elsewhere.
    submitter_name_map = {
        "david": "David Estes",
        "lydia": "Lydia Williams",
        "mark": "Mark Burcham",
        "randy": "Randy",
        "richard": "Richard Winger",
        "vern": "Vern Abbott",
    }
    submitted_by_key = str(submitted_by or "").strip()
    submitted_by_for_doc = submitter_name_map.get(submitted_by_key.lower(), submitted_by_key)

    # Prepare replacements for placeholders (using double-bracket format as in template)
    replacements = {
        '[[CustomerName]]': customer_name,
        '[[ProjectStreetAddr]]': street_address,
        '[[ProjectCity]]': city,
        '[[ProjectState]]': state,
        '[[ProjectZip]]': zip_code,
        '[[Date]]': formatted_date,
        '[[Squares]]': total_squares,
        '[[PriceIncludesLanguage]]': price_includes_text,
        '[[WarrantyIncluded]]': warranty_text,
        '[[SubmittedBy]]': submitted_by_for_doc,
        '[[10YrTotalPrice]]': f"{tp10:,.0f}",
        '[[15YrTotalPrice]]': f"{tp15:,.0f}",
        '[[20YrTotalPrice]]': f"{tp20:,.0f}",
        '[[AdditionalLanguage]]': proposal_language if proposal_language else ' '
    }

    # Folder and file names
    # Folder and file names
    if target_folder:
        proposal_folder = target_folder
        os.makedirs(proposal_folder, exist_ok=True)
        folder_name = os.path.basename(proposal_folder)
    else:
        folder_name = f"{customer_name} - {street_address}"
        proposal_folder = os.path.join(PROPOSAL_TEMP_DIR, folder_name)
        os.makedirs(proposal_folder, exist_ok=True)

    # Map roof type to suffix
    roof_suffix_map = {
        "TPO/EPDM": "TPO EPDM Metal.docx",
        "Metal": "TPO EPDM Metal.docx",
        "Mod Bit": "Mod Bit.docx",
        "Rock/Foam/Coat": "RFC.docx",
        "Ballasted 45 mil": "Ballasted 45mil.docx",
        "Ballasted 60 mil": "Ballasted 60mil.docx"
    }
    roof_suffix = roof_suffix_map.get(roof_type, "Unknown.docx")

    # Select template based on product and roof type
    prefix = "Gaco S42 Proposal - " if product == "Gaco" else "Uniflex Proposal - "
    doc_template_name = f"{prefix}{roof_suffix}"
    doc_output_name = f"{prefix}{street_address}.docx"
    doc_template_path = os.path.join(TEMPLATE_DIR, doc_template_name)
    doc_output_path = os.path.join(proposal_folder, doc_output_name)

    create_started = time.perf_counter()

    # Copy template and replace placeholders directly in the output file
    doc_started = time.perf_counter()
    shutil.copy(doc_template_path, doc_output_path)
    doc = Document(doc_output_path)
    replace_placeholder_blocks(doc, replacements)
    doc.save(doc_output_path)
    _log_timing(f"word generation for {folder_name}", doc_started)

    # Convert Word doc to PDF and save in same folder (headless if possible)
    pdf_started = time.perf_counter()
    _convert_to_pdf(
        doc_output_path,
        proposal_folder,
        use_libreoffice=use_libreoffice,
        async_mode=pdf_async,
    )
    if not pdf_async:
        _log_timing(f"pdf conversion for {folder_name}", pdf_started)

    # Copy Excel files
    profit_template = os.path.join(TEMPLATE_DIR, "Profit Summary.xlsx")
    profit_output = os.path.join(proposal_folder, f"Profit Summary - {street_address}.xlsx")
    shutil.copy(profit_template, profit_output)

    # Update Profit Summary.xlsx using openpyxl (no Excel automation required)
    try:
        workbook_started = time.perf_counter()
        wb_profit = load_workbook(profit_output)

        # Prepare default header-only map, then merge any provided mapped_data
        default_header_map = {
            "customer_name": customer_name,
            "street_address": street_address,
            "city": city,
            "state": state,
            "zip_code": zip_code,
            "pcs_or_roofer_ind": (mapped_data.get("pcs_or_roofer_ind") if mapped_data else ""),
            "flat_roof_squares": (mapped_data.get("flat_roof_squares") if mapped_data else 0),
            "wall_squares": (mapped_data.get("wall_squares") if mapped_data else 0),
            "squares": total_squares,
            "current_roof": roof_type,
            "product": product,
            "warranty_incl": warranty_incl,
            "submitted_by": submitted_by,
            "proposal_note": "",
            "proposal_language": proposal_language if proposal_language else "",
            "lead": (mapped_data.get("lead") if mapped_data else ""),
        }
        merged_map = dict(default_header_map)
        if mapped_data:
            merged_map.update({
                k: v for k, v in mapped_data.items()
                if (k in EXCEL_CELL_MAP and EXCEL_CELL_MAP[k]) or (k in HIDDEN_SHEET_CELL_MAP) or (k in CALC_ONLY_FIELDS)
            })

        # Compute all derived totals so the saved .xlsx has concrete numbers (no Excel recalc needed)
        def _get_num(key, default=0.0):
            try:
                v = merged_map.get(key)
                if v is None:
                    return default
                if isinstance(v, str):
                    v = v.replace('$', '').replace(',', '').strip()
                    if v == '':
                        return default
                return float(v)
            except Exception:
                return default

        def _get_int(key, default=0):
            try:
                return int(_get_num(key, default))
            except Exception:
                return default

        # Derive office fee/commission from submitted_by when not provided
        _submitted_by = submitted_by
        _office_fee_pct = merged_map.get("office_fee_pct")
        if _office_fee_pct in (None, "", 0):
            _office_fee_pct = office_fee_pct_for_submitter(_submitted_by)
        try:
            _office_fee_pct = float(_office_fee_pct)
        except Exception:
            _office_fee_pct = office_fee_pct_for_submitter(_submitted_by)

        _commission_pct = commission_pct_for_submitter(_submitted_by)

        # Inputs for calculation
        calc_result = calculation_routine(
            squares=_get_num("squares", float(total_squares) if total_squares else 0.0),
            product=str(merged_map.get("product") or product or ""),
            roof_type=str(merged_map.get("current_roof") or roof_type or ""),
            labor_days=_get_int("labor_days", 0),
            warranty_incl=str(merged_map.get("warranty_incl") or warranty_incl or "No"),
            include_travel=str(merged_map.get("include_travel") or "No"),
            previous_include_travel=str(merged_map.get("previous_include_travel") or "No"),
            previous_calc_travel_total=_get_num("previous_calc_travel_total", 0.0),
            price_per_sq_10=_get_num("price_per_sq_10", 0.0),
            commission_pct=_commission_pct,
            submitted_by=_submitted_by,
            previous_submitted_by=_submitted_by,
            office_fee_pct=_office_fee_pct,
            adjusted_coverage=_get_num("adjusted_coverage", 0.0),
            silicone_units_10=_get_num("silicone_units_10", 0.0),
            silicone_price=_get_num("silicone_price", 0.0),
            gaco_patch_units=_get_num("gaco_patch_units", 0.0),
            gaco_patch_price=_get_num("gaco_patch_price", 0.0),
            sw_1flash_units=_get_num("sw_1flash_units", 0.0),
            sw_1flash_price=_get_num("sw_1flash_price", 0.0),
            bleed_trap_units=_get_num("bleed_trap_units", 0.0),
            bleed_trap_price=_get_num("bleed_trap_price", 0.0),
            gaco_e5320_units=_get_num("gaco_e5320_units", 0.0),
            gaco_e5320_price=_get_num("gaco_e5320_price", GACO_E5320_PRICE),
            sw_bleed_block_units=_get_num("sw_bleed_block_units", 0.0),
            sw_bleed_block_price=_get_num("sw_bleed_block_price", 0.0),
            drainage_mat_units=_get_num("drainage_mat_units", 0.0),
            drainage_mat_price=_get_num("drainage_mat_price", 0.0),
            foam_units=_get_num("foam_units", 0.0),
            foam_price=_get_num("foam_price", 0.0),
            rfc_labor_price=_get_num("rfc_labor_price", 0.0),
            pcs_labor_price=_get_num("pcs_labor_price", 0.0),
            scarifying_total=_get_num("scarifying_total", 0.0),
            travel_total=_get_num("travel_total", 0.0),
            repair_costs_total=_get_num("repair_costs_total", 0.0),
            previous_squares=_get_num("previous_squares", _get_num("squares", 0.0)),
            previous_roof_type=str(merged_map.get("previous_roof_type") or merged_map.get("current_roof") or roof_type or ""),
            previous_product=str(merged_map.get("previous_product") or merged_map.get("product") or product or ""),
            previous_adjusted_coverage=_get_num("previous_adjusted_coverage", _get_num("adjusted_coverage", 0.0)),
            previous_silicone_units_10=_get_num("previous_silicone_units_10", _get_num("silicone_units_10", 0.0)),
            proposal_note=str(merged_map.get("proposal_note") or ""),
            pcs_or_roofer_ind=str(merged_map.get("pcs_or_roofer_ind") or "").strip(),
            previous_pcs_or_roofer_ind=str(merged_map.get("previous_pcs_or_roofer_ind") or merged_map.get("pcs_or_roofer_ind") or "").strip(),
        )

        # Only overlay keys we know how to write (present in EXCEL_CELL_MAP and not None)
        for k, v in calc_result.items():
            if k in EXCEL_CELL_MAP and EXCEL_CELL_MAP[k] and v is not None:
                merged_map[k] = v
        merged_map["include_travel"] = include_travel_from_travel_total(
            merged_map.get("travel_total", calc_result.get("travel_total", 0))
        )

        # --- Write labor_days as formula when it matches the calculated baseline; otherwise, keep user value ---
        def _to_int_safe(v, d=None):
            try:
                if v is None:
                    return d
                if isinstance(v, str):
                    s = v.replace(',', '').strip()
                    if s == '':
                        return d
                    return int(float(s))
                return int(float(v))
            except Exception:
                return d
        # Calculate baseline from current header fields
        try:
            _squares_int = int(float(merged_map.get("squares") or total_squares or 0))
        except Exception:
            _squares_int = 0
        _roof_type = str(merged_map.get("current_roof") or roof_type or "")
        if _roof_type in ("Ballasted 60 mil", "Ballasted 45 mil"):
            _calc_ld = int(math.ceil(_squares_int / 30.0))
            _ld_formula = "=ROUNDUP(E3/30,0)"
        elif _roof_type == "Rock/Foam/Coat":
            _calc_ld = int(math.ceil(_squares_int / 75.0))
            _ld_formula = "=ROUNDUP(E3/75,0)"
        else:
            _calc_ld = int(math.ceil(_squares_int / 45.0))
            _ld_formula = "=ROUNDUP(E3/45,0)"
        _user_ld = _to_int_safe(merged_map.get("labor_days"), None)
        # Rule:
        # - If user value is blank/None, treat as calculated -> write formula
        # - If user value equals calculated baseline -> write formula
        # - Otherwise (differs from baseline) -> write the user number
        if (_user_ld is None) or (_user_ld == _calc_ld):
            # Prefer centralized formula from profit_summary_formulas if available
            if _PS_F_E7_LABOR:
                merged_map["labor_days"] = _PS_F_E7_LABOR.replace("\n", "")
            else:
                merged_map["labor_days"] = _ld_formula
        else:
            merged_map["labor_days"] = _user_ld

        # --- Write silicone_units_10 as formula when it matches the calculated baseline; otherwise, keep user value ---
        # Determine effective 10-year coverage from product + roof + adjusted_coverage
        _prod = str(merged_map.get("product") or product or "")
        _roof = str(merged_map.get("current_roof") or roof_type or "")
        try:
            _adj_cov = float(merged_map.get("adjusted_coverage") or 0.0)
        except Exception:
            _adj_cov = 0.0
        eff_cov_10 = adjusted_coverage_rates(_prod, _roof, _adj_cov).get(10, 0.0)
        # Baseline units (rounded up)
        try:
            _squares_float = float(merged_map.get("squares") or total_squares or 0.0)
        except Exception:
            _squares_float = 0.0
        _calc_units_10 = int(math.ceil((_squares_float / 5.0) * eff_cov_10))
        # Parse the user value (rounded up if provided)
        def _to_int_ceil(v, d=None):
            try:
                if v is None:
                    return d
                if isinstance(v, str):
                    s = v.replace(',', '').strip()
                    if s == '':
                        return d
                    return int(math.ceil(float(s)))
                return int(math.ceil(float(v)))
            except Exception:
                return d
        _user_units_10 = _to_int_ceil(merged_map.get("silicone_units_10"), None)
        # Rule:
        # - If user value is blank/None, treat as calculated -> write formula
        # - If user value equals calculated baseline -> write formula
        # - Otherwise (differs from baseline) -> write the user number
        if abs(float(_adj_cov or 0.0)) > 1e-9:
            merged_map["silicone_units_10"] = PS_F_C11.replace("\n", "")
        elif (_user_units_10 is None) or (_user_units_10 == _calc_units_10):
            # Use centralized formula for silicone_units_10
            merged_map["silicone_units_10"] = PS_F_C11.replace("\n", "")
        else:
            merged_map["silicone_units_10"] = _user_units_10

        merged_map["silicone_units_15"] = PS_F_H11.replace("\n", "")
        merged_map["silicone_units_20"] = PS_F_N11.replace("\n", "")

        # --- Override-or-formula behavior for additional inputs (incremental) ---
        # Helpers
        def _to_float_clean(v, d=None):
            try:
                if v is None:
                    return d
                if isinstance(v, str):
                    s = v.replace("$","").replace(",","").strip()
                    if s == "":
                        return d
                    return float(s)
                return float(v)
            except Exception:
                return d
        def _to_int_ceil_clean(v, d=None):
            try:
                if v is None:
                    return d
                if isinstance(v, str):
                    s = v.replace(",","").strip()
                    if s == "":
                        return d
                    return int(math.ceil(float(s)))
                return int(math.ceil(float(v)))
            except Exception:
                return d
        def _store_formula_or_value(user_val, base_val, formula_str, tol=1e-6):
            """Return `formula_str` when the field is baseline (blank or equals base), else return the user value.
            - user_val: possibly None/str/number
            - base_val: numeric baseline used for comparison
            - formula_str: Excel formula string to store when baseline
            - tol: numeric tolerance for equality
            """
            # If user value is missing, use formula
            if user_val is None:
                return formula_str
            # Try numeric compare; if it fails, fall back to formula-or-user depending on emptiness
            try:
                uv = float(user_val)
                bv = float(base_val or 0.0)
                if abs(uv - bv) <= tol:
                    return formula_str
                return user_val
            except Exception:
                # If user provided a non-numeric but non-empty, keep it; else formula
                s = str(user_val).strip()
                return user_val if s != "" else formula_str

        _prodF = str(merged_map.get("product") or product or "")
        _roofF = str(merged_map.get("current_roof") or roof_type or "")
        _sqF   = _to_float_clean(merged_map.get("squares"), float(total_squares) if total_squares else 0.0)

        # 0) 10-year price per square (M3)
        # Use canonical formula from profit_summary_formulas.py when baseline; else store user's number
        try:
            _roof_idxF = roof_types.index(_roofF)
            _indF = str(merged_map.get("pcs_or_roofer_ind") or "").strip()
            if _indF == "PCS Direct":
                _pricing10 = pcs_pricing10
            else:
                _pricing10 = roofer_pricing10
            _base_pps10 = float(_pricing10[_roof_idxF])
        except Exception:
            _base_pps10 = 0.0
        _user_pps10 = _to_float_clean(merged_map.get("price_per_sq_10"), None)
        merged_map["price_per_sq_10"] = _store_formula_or_value(
            _user_pps10,
            _base_pps10,
            PS_F_M3.replace("\n", "")
        )

        # 1) Units fields
        # gaco_patch_units: use centralized formula constant unless user-overridden
        if _prodF == "Gaco":
            _base_gaco_patch_units = (
                math.ceil((_sqF or 0.0) * 0.03)
                if _roofF == "Rock/Foam/Coat"
                else math.ceil((_sqF or 0.0) / 10.0)
            )
        else:
            _base_gaco_patch_units = 0
        _user_gaco_patch_units = _to_int_ceil_clean(merged_map.get("gaco_patch_units"), None)
        if (_user_gaco_patch_units is None) or (_user_gaco_patch_units == _base_gaco_patch_units):
            merged_map["gaco_patch_units"] = PS_F_C12.replace("\n", "")
        else:
            merged_map["gaco_patch_units"] = _user_gaco_patch_units

        # bleed_trap_units: use centralized formula constant unless user-overridden
        _base_bleed_units = math.ceil((_sqF or 0.0) / 5.0) if (_prodF == "Gaco" and _roofF == "Mod Bit") else 0
        _user_bleed_units = _to_int_ceil_clean(merged_map.get("bleed_trap_units"), None)
        if (_user_bleed_units is None) or (_user_bleed_units == _base_bleed_units):
            merged_map["bleed_trap_units"] = PS_F_C13.replace("\n", "")
        else:
            merged_map["bleed_trap_units"] = _user_bleed_units

        _user_gaco_e5320_units = _to_int_ceil_clean(merged_map.get("gaco_e5320_units"), None)
        merged_map["gaco_e5320_units"] = PS_F_C14.replace("\n", "") if not _user_gaco_e5320_units else _user_gaco_e5320_units

        # sw_1flash_units: use centralized formula constant unless user-overridden
        if _prodF == "Uniflex":
            _div = 20 if _roofF in ["TPO/EPDM", "Mod Bit", "Rock/Foam/Coat"] else 10
            _base_sw1_units = math.ceil((_sqF or 0.0) / _div)
        else:
            _base_sw1_units = 0
        _user_sw1_units = _to_int_ceil_clean(merged_map.get("sw_1flash_units"), None)
        if (_user_sw1_units is None) or (_user_sw1_units == _base_sw1_units):
            merged_map["sw_1flash_units"] = PS_F_C15.replace("\n", "")
        else:
            merged_map["sw_1flash_units"] = _user_sw1_units

        # sw_bleed_block_units: use centralized formula constant unless user-overridden
        _base_swbb_units = math.ceil((_sqF or 0.0) / 5.0) if (_prodF == "Uniflex" and _roofF == "Mod Bit") else 0
        _user_swbb_units = _to_int_ceil_clean(merged_map.get("sw_bleed_block_units"), None)
        if (_user_swbb_units is None) or (_user_swbb_units == _base_swbb_units):
            merged_map["sw_bleed_block_units"] = PS_F_C16.replace("\n", "")
        else:
            merged_map["sw_bleed_block_units"] = _user_swbb_units

        # drainage_mat_units: use centralized formula constant unless user-overridden
        _base_drain_units = math.ceil((_sqF or 0.0) / 18.0) if _roofF in ["Ballasted 60 mil", "Ballasted 45 mil"] else 0
        _user_drain_units = _to_int_ceil_clean(merged_map.get("drainage_mat_units"), None)
        if (_user_drain_units is None) or (_user_drain_units == _base_drain_units):
            merged_map["drainage_mat_units"] = PS_F_C17.replace("\n", "")
        else:
            merged_map["drainage_mat_units"] = _user_drain_units

        # foam_units: use centralized formula constant unless user-overridden
        _base_foam_units = math.ceil((_sqF or 0.0) / 25.0) if _roofF == "Rock/Foam/Coat" else 0
        _user_foam_units = _to_int_ceil_clean(merged_map.get("foam_units"), None)
        if (_user_foam_units is None) or (_user_foam_units == _base_foam_units):
            merged_map["foam_units"] = PS_F_C18.replace("\n", "")
        else:
            merged_map["foam_units"] = _user_foam_units

        # 2) Unit price fields (store formula when baseline, number when overridden)

        _base_sil_price = GACO_S42_BASE_PRICE if _prodF == "Gaco" else (UNIFLEX_BASE_PRICE if _prodF == "Uniflex" else 0)
        _user_sil_price = _to_float_clean(merged_map.get("silicone_price"), None)
        merged_map["silicone_price"] = _store_formula_or_value(
            _user_sil_price,
            _base_sil_price,
            PS_F_D11.replace("\n", "")
        )

        _base_gp_price = GACO_PATCH_BASE_PRICE if _prodF == "Gaco" else 0
        _user_gp_price = _to_float_clean(merged_map.get("gaco_patch_price"), None)
        merged_map["gaco_patch_price"] = _store_formula_or_value(
            _user_gp_price,
            _base_gp_price,
            PS_F_D12.replace("\n", "")
        )

        _base_bt_price = BLEED_TRAP_BASE_PRICE if (_prodF == "Gaco" and _roofF == "Mod Bit") else 0
        _user_bt_price = _to_float_clean(merged_map.get("bleed_trap_price"), None)
        merged_map["bleed_trap_price"] = _store_formula_or_value(
            _user_bt_price,
            _base_bt_price,
            PS_F_D13.replace("\n", "")
        )

        _user_gaco_e5320_price = _to_float_clean(merged_map.get("gaco_e5320_price"), None)
        merged_map["gaco_e5320_price"] = (
            PS_F_D14.replace("\n", "")
            if _user_gaco_e5320_price in (None, 0, GACO_E5320_PRICE)
            else _user_gaco_e5320_price
        )

        _base_sw1_price = SW_1FLASH_BASE_PRICE if _prodF == "Uniflex" else 0
        _user_sw1_price = _to_float_clean(merged_map.get("sw_1flash_price"), None)
        merged_map["sw_1flash_price"] = _store_formula_or_value(
            _user_sw1_price,
            _base_sw1_price,
            PS_F_D15.replace("\n", "")
        )

        _base_swbb_price = SW_BLEED_BLOCK_BASE_PRICE if (_prodF == "Uniflex" and _roofF == "Mod Bit") else 0
        _user_swbb_price = _to_float_clean(merged_map.get("sw_bleed_block_price"), None)
        merged_map["sw_bleed_block_price"] = _store_formula_or_value(
            _user_swbb_price,
            _base_swbb_price,
            PS_F_D16.replace("\n", "")
        )

        _base_drain_price = DRAINAGE_MAT_BASE_PRICE if _roofF in ["Ballasted 60 mil", "Ballasted 45 mil"] else 0
        _user_drain_price = _to_float_clean(merged_map.get("drainage_mat_price"), None)
        merged_map["drainage_mat_price"] = _store_formula_or_value(
            _user_drain_price,
            _base_drain_price,
            PS_F_D17.replace("\n", "")
        )

        if _roofF == "Rock/Foam/Coat":
            _base_foam_price = GACO_FOAM_BASE_PRICE if _prodF == "Gaco" else (UNIFLEX_FOAM_BASE_PRICE if _prodF == "Uniflex" else 0)
        else:
            _base_foam_price = 0
        _user_foam_price = _to_float_clean(merged_map.get("foam_price"), None)
        merged_map["foam_price"] = _store_formula_or_value(
            _user_foam_price,
            _base_foam_price,
            PS_F_D18.replace("\n", "")
        )

        _base_rfc_price = RFC_LABOR_RATE if _roofF == "Rock/Foam/Coat" else 0
        _user_rfc_price = _to_float_clean(merged_map.get("rfc_labor_price"), None)
        merged_map["rfc_labor_price"] = _store_formula_or_value(
            _user_rfc_price,
            _base_rfc_price,
            PS_F_D19.replace("\n", "")
        )

        _user_pcs_labor_price = _to_float_clean(merged_map.get("pcs_labor_price"), None)
        merged_map["pcs_labor_price"] = _store_formula_or_value(
            _user_pcs_labor_price,
            PCS_BASE_LABOR_RATE,
            PS_F_D21.replace("\n", "")
        )

        # === Force key fields to be written as Excel formulas ===
        # Price per sq (15/20): centralized formulas
        merged_map["price_per_sq_15"] = PS_F_M5.replace("\n", "")
        merged_map["price_per_sq_20"] = PS_F_M7.replace("\n", "")

        # Warranty total (E24)
        merged_map['warranty_10_total'] = PS_F_E24.replace("\n", "")

        # Line-item totals (E11–E18)
        merged_map['silicone_total'] = PS_F_E11.replace("\n", "")
        merged_map['silicone_15_total'] = PS_F_K11.replace("\n", "")
        merged_map['silicone_20_total'] = PS_F_P11.replace("\n", "")
        merged_map['gaco_patch_total']      = PS_F_E12.replace("\n", "")
        merged_map['bleed_trap_total']      = PS_F_E13.replace("\n", "")
        merged_map['gaco_e5320_total']      = '=IF(C14<>"",C14*D14,0)'
        merged_map['sw_1flash_total']       = PS_F_E15.replace("\n", "")
        merged_map['sw_bleed_block_total']  = PS_F_E16.replace("\n", "")
        merged_map['drainage_mat_total']    = PS_F_E17.replace("\n", "")
        merged_map['foam_total']            = PS_F_E18.replace("\n", "")

        # Labor totals (E19, E21)
        merged_map['rfc_labor_total'] = PS_F_E19.replace("\n", "")
        merged_map['pcs_labor_total'] = PS_F_E21.replace("\n", "")

        # Totals by term (P3, P5, P7)
        merged_map['total_price_10'] = PS_F_P3.replace("\n", "")
        merged_map['total_price_15'] = PS_F_P5.replace("\n", "")
        merged_map['total_price_20'] = PS_F_P7.replace("\n", "")

        # Fees / commission / totals / profit calcs
        merged_map['commission_amt']   = PS_F_E25.replace("\n", "")
        merged_map['office_fee_total'] = PS_F_E33.replace("\n", "")
        merged_map['total_cost']       = PS_F_E27.replace("\n", "")
        merged_map['profit_share']     = PS_F_E32.replace("\n", "")
        merged_map['profit_share_15']  = PS_F_K32.replace("\n", "")
        merged_map['profit_share_20']  = PS_F_P32.replace("\n", "")
        merged_map['pcs_profit']       = PS_F_E29.replace("\n", "")
        merged_map['pcs_profit_15']    = PS_F_K29.replace("\n", "")
        merged_map['pcs_profit_20']    = PS_F_P29.replace("\n", "")
        merged_map['profit_pct']       = PS_F_E30.replace("\n", "")
        merged_map['profit_pct_15']    = PS_F_K30.replace("\n", "")
        merged_map['profit_pct_20']    = PS_F_P30.replace("\n", "")
        merged_map['daily_profit']     = PS_F_E31.replace("\n", "")
        merged_map['daily_profit_15']  = PS_F_K31.replace("\n", "")
        merged_map['daily_profit_20']  = PS_F_P31.replace("\n", "")

        # Build ghost values (evaluated numbers) for UI display
        ghost_fields = [
            "labor_days",
            "silicone_units_10","silicone_units_15","silicone_units_20","gaco_patch_units","bleed_trap_units","gaco_e5320_units",
            "sw_1flash_units","sw_bleed_block_units","drainage_mat_units","foam_units",
            "silicone_price","gaco_patch_price","bleed_trap_price","gaco_e5320_price",
            "sw_1flash_price","sw_bleed_block_price","drainage_mat_price",
            "foam_price","rfc_labor_price","pcs_labor_price",
            "price_per_sq_10","price_per_sq_15","price_per_sq_20",
            "total_price_10","total_price_15","total_price_20",
            "silicone_total","silicone_15_total","silicone_20_total","gaco_patch_total","bleed_trap_total","gaco_e5320_total","sw_1flash_total",
            "sw_bleed_block_total","drainage_mat_total","foam_total",
            "rfc_labor_total","pcs_labor_total",
            "warranty_10_total","warranty_15_total","warranty_20_total",
            "office_fee_total","office_fee_15_total","office_fee_20_total",
            "total_cost","total_cost_15","total_cost_20",
            "commission_amt","commission_amt_15","commission_amt_20",
            "profit_share","profit_share_15","profit_share_20",
            "pcs_profit","pcs_profit_15","pcs_profit_20",
            "profit_pct","profit_pct_15","profit_pct_20",
            "daily_profit","daily_profit_15","daily_profit_20",
        ]
        ghost_values = {k: calc_result.get(k) for k in ghost_fields if k in calc_result}

        apply_adjusted_spread_rates(
            wb_profit,
            merged_map.get("product") or product,
            merged_map.get("current_roof") or roof_type,
            merged_map.get("adjusted_coverage", 0.0),
        )
        write_fields_to_profit_summary(wb_profit, merged_map)
        write_ghost_values(wb_profit, ghost_values)
        ensure_profit_summary_validations(wb_profit)
        wb_profit.save(profit_output)
        _log_timing(f"profit summary write for {folder_name}", workbook_started)

        if create_email_draft:
            try:
                destination_root = get_copy_destination_for_submitter(submitted_by)
                email_folder_path = os.path.join(destination_root, folder_name) if destination_root else ""
                draft_warning = create_outlook_proposal_summary_draft(
                    customer_name=customer_name,
                    street_address=street_address,
                    submitted_by=submitted_by,
                    total_squares=merged_map.get("squares", total_squares),
                    flat_roof_squares=merged_map.get("flat_roof_squares", 0),
                    wall_squares=merged_map.get("wall_squares", 0),
                    roof_type=merged_map.get("current_roof", roof_type),
                    daily_profit=calc_result.get("daily_profit", 0),
                    proposal_note=merged_map.get("proposal_note", ""),
                    proposal_language=merged_map.get("proposal_language", ""),
                    folder_name=folder_name,
                    folder_link=build_proposal_folder_link(
                        email_folder_path,
                        submitted_by=submitted_by,
                        folder_name=folder_name,
                    ),
                )
                if draft_warning:
                    _notify_user(draft_warning, "warning")
            except Exception as exc:
                _notify_user(str(exc), "warning")

        # Append a line to the Proposal Tracking workbook
        if update_tracking:
            try:
                tracker_started = time.perf_counter()
                append_to_proposal_tracking(
                    created_date=formatted_date,
                    customer_name=customer_name,
                    street_address=street_address,
                    city=city,
                    state=state,
                    zip_code=zip_code,
                    product=product,
                    roof_type=roof_type,
                    total_squares=total_squares,
                    warranty_incl=warranty_incl,
                    submitted_by=submitted_by,
                    folder_name=folder_name,
                    proposal_folder=proposal_folder,
                    tp10=tp10,
                    tp15=tp15,
                    tp20=tp20,
                    lead_value=(mapped_data.get("lead") if mapped_data else ""),
                )
                _log_timing(f"proposal tracker update for {folder_name}", tracker_started)
            except Exception as e:
                try:
                    with open(APP_ERROR_LOG, "a", encoding="utf-8") as _f:
                        _f.write(f"[ERROR] tracking-append error for {folder_name}: {e}\n")
                except Exception:
                    pass

    except Exception as e:
        # Log like before (re-use same Desktop log pattern)
        _LOG_PATH = pathlib.Path(
            os.environ.get("PCS_XLWINGS_LOG_PATH", str(DEFAULT_DATA_DIR / "pcs_xlwings.log"))
        )
        try:
            with open(_LOG_PATH, "a", encoding="utf-8") as _f:
                _f.write(f"\n[OPENPYXL ERROR] {e}\n\n")
                traceback.print_exc(file=_f)
        except Exception:
            pass

    if copy_destination:
        try:
            copy_started = time.perf_counter()
            copy_proposal_to_submitter_destination(
                proposal_folder,
                folder_name,
                submitted_by,
                wait_for_pdfs=pdf_async,
            )
            _log_timing(f"destination copy for {folder_name}", copy_started)
        except Exception as e:
            with open(APP_ERROR_LOG, "a", encoding="utf-8") as f:
                f.write(f"[ERROR] Copy after create failed for {folder_name}: {e}\n")

    try:
        temp_copy_started = time.perf_counter()
        copy_proposal_to_temp_dir(proposal_folder, folder_name)
        _log_timing(f"temp proposal sync for {folder_name}", temp_copy_started)
    except Exception as e:
        try:
            with open(APP_ERROR_LOG, "a", encoding="utf-8") as f:
                f.write(f"[ERROR] Temp sync after create/save failed for {folder_name}: {e}\n")
        except Exception:
            pass

    _log_timing(f"create_proposal_from_fields total for {folder_name}", create_started)

    return folder_name

# Resolve template/static paths for both dev and frozen app (PyInstaller)
if getattr(sys, 'frozen', False):
    BASE_DIR = sys._MEIPASS  # type: ignore[attr-defined]
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, 'templates')
STATIC_PATH = os.path.join(BASE_DIR, 'static')
_BUNDLED_PROPOSAL_SUMMARY_TEMPLATE_PATH = os.path.join(
    BASE_DIR,
    "resources",
    "1. Proposal Summary Template.emltpl",
)
if os.path.exists(_BUNDLED_PROPOSAL_SUMMARY_TEMPLATE_PATH):
    PROPOSAL_SUMMARY_TEMPLATE_PATH = _BUNDLED_PROPOSAL_SUMMARY_TEMPLATE_PATH

app = Flask(__name__, template_folder=TEMPLATE_PATH, static_folder=STATIC_PATH)
app.secret_key = flask_secret_key()


@app.context_processor
def application_identity():
    tenant = None
    if MULTI_TENANT_ENABLED:
        try:
            tenant = current_tenant_context()
        except TenantAuthenticationError:
            pass
    return {
        "app_display_name": APP_DISPLAY_NAME,
        "app_variant": APP_VARIANT,
        "app_is_beta": APP_IS_BETA,
        "multi_tenant_enabled": MULTI_TENANT_ENABLED,
        "proposal_storage_mode": PROPOSAL_STORAGE_MODE,
        "active_tenant": tenant,
    }


_AUTHENTICATION_EXEMPT_ENDPOINTS = {
    "static", "tenant_login", "desktop_session_heartbeat", "desktop_session_closed",
    "application_settings",
}


@app.before_request
def require_tenant_session():
    if not MULTI_TENANT_ENABLED or request.endpoint in _AUTHENTICATION_EXEMPT_ENDPOINTS:
        return None
    try:
        current_tenant_context()
    except TenantAuthenticationError:
        return redirect(url_for("tenant_login", next=request.full_path.rstrip("?")))
    return None


@app.route("/sign-in", methods=["GET", "POST"])
def tenant_login():
    if request.method == "POST":
        try:
            context = tenant_sign_in(request.form.get("email"), request.form.get("password"))
        except TenantAuthenticationError as exc:
            flash(str(exc), "danger")
        else:
            flash(f"Signed in to {context.tenant_name}.", "success")
            destination = str(request.form.get("next") or "")
            if not destination.startswith("/") or destination.startswith("//"):
                destination = url_for("landing_page")
            return redirect(destination)
    return render_template("tenant_login.html", next=request.args.get("next", ""))


@app.post("/sign-out")
def tenant_logout():
    tenant_sign_out()
    return redirect(url_for("tenant_login"))

DESKTOP_LIFECYCLE_ENABLED = os.environ.get("PCS_PROPOSAL_DESKTOP_LIFECYCLE", "0").strip().lower() in {
    "1",
    "true",
    "yes",
    "on",
}
DESKTOP_STARTUP_GRACE_SECONDS = 60.0
DESKTOP_CLOSE_GRACE_SECONDS = 5.0
_desktop_lifecycle_lock = threading.Lock()
_desktop_server_started_at = time.monotonic()
_desktop_last_heartbeat = None
_desktop_close_requested_at = None
DESKTOP_BACKGROUND_WORK_ACTIVE = threading.Event()


def _desktop_session_should_stop(
    now,
    *,
    server_started_at,
    last_heartbeat,
    close_requested_at,
    startup_grace=DESKTOP_STARTUP_GRACE_SECONDS,
    close_grace=DESKTOP_CLOSE_GRACE_SECONDS,
):
    if close_requested_at is not None:
        return now - close_requested_at >= close_grace
    return last_heartbeat is None and now - server_started_at >= startup_grace


def _desktop_lifecycle_watchdog():
    while True:
        time.sleep(1)
        now = time.monotonic()
        with _desktop_lifecycle_lock:
            should_stop = _desktop_session_should_stop(
                now,
                server_started_at=_desktop_server_started_at,
                last_heartbeat=_desktop_last_heartbeat,
                close_requested_at=_desktop_close_requested_at,
            )
        if should_stop and not DESKTOP_BACKGROUND_WORK_ACTIVE.is_set():
            _safe_debug("[INFO] PCS desktop browser session ended; stopping the local server.")
            os._exit(0)


if DESKTOP_LIFECYCLE_ENABLED:
    threading.Thread(
        target=_desktop_lifecycle_watchdog,
        name="pcs-desktop-lifecycle",
        daemon=True,
    ).start()


@app.post("/api/desktop-session/heartbeat")
def desktop_session_heartbeat():
    global _desktop_last_heartbeat, _desktop_close_requested_at
    if not DESKTOP_LIFECYCLE_ENABLED:
        return ("", 404)
    with _desktop_lifecycle_lock:
        _desktop_last_heartbeat = time.monotonic()
        _desktop_close_requested_at = None
    return ("", 204)


@app.post("/api/desktop-session/closed")
def desktop_session_closed():
    global _desktop_close_requested_at
    if not DESKTOP_LIFECYCLE_ENABLED:
        return ("", 404)
    with _desktop_lifecycle_lock:
        _desktop_close_requested_at = time.monotonic()
    return ("", 204)


_DESKTOP_LIFECYCLE_SCRIPT = """
<script id="pcs-desktop-lifecycle">
(() => {
    let heartbeatTimer = null;
    let internalNavigationPending = false;
    let internalNavigationResetTimer = null;

    const heartbeat = () => fetch('/api/desktop-session/heartbeat', {
        method: 'POST',
        cache: 'no-store',
        keepalive: true,
    }).catch(() => {});

    const startHeartbeat = () => {
        if (heartbeatTimer !== null) window.clearInterval(heartbeatTimer);
        heartbeat();
        heartbeatTimer = window.setInterval(heartbeat, 2000);
    };

    const markInternalNavigation = () => {
        internalNavigationPending = true;
        if (internalNavigationResetTimer !== null) {
            window.clearTimeout(internalNavigationResetTimer);
        }
        internalNavigationResetTimer = window.setTimeout(() => {
            internalNavigationPending = false;
            internalNavigationResetTimer = null;
        }, 2000);
    };

    document.addEventListener('click', (event) => {
        const element = event.target instanceof Element ? event.target : null;
        const anchor = element ? element.closest('a[href]') : null;
        if (anchor && !anchor.hasAttribute('download') && anchor.target !== '_blank') {
            try {
                const destination = new URL(anchor.href, window.location.href);
                if (destination.origin === window.location.origin) markInternalNavigation();
            } catch (_) {}
            return;
        }
        if (element && element.closest('[data-href]')) markInternalNavigation();
    }, true);

    document.addEventListener('submit', markInternalNavigation, true);

    window.addEventListener('pageshow', () => {
        internalNavigationPending = false;
        if (internalNavigationResetTimer !== null) {
            window.clearTimeout(internalNavigationResetTimer);
            internalNavigationResetTimer = null;
        }
        startHeartbeat();
    });

    window.addEventListener('pagehide', () => {
        if (heartbeatTimer !== null) {
            window.clearInterval(heartbeatTimer);
            heartbeatTimer = null;
        }
        if (!internalNavigationPending) {
            navigator.sendBeacon('/api/desktop-session/closed', '');
        }
    });

    startHeartbeat();
})();
</script>
"""


@app.after_request
def _inject_desktop_lifecycle(response):
    if not DESKTOP_LIFECYCLE_ENABLED or response.mimetype != "text/html":
        return response
    body = response.get_data(as_text=True)
    if "</body>" in body and 'id="pcs-desktop-lifecycle"' not in body:
        response.set_data(body.replace("</body>", f"{_DESKTOP_LIFECYCLE_SCRIPT}</body>", 1))
    return response

# ---- Error logging to the application folder so packaged app errors are visible ----
import logging
from logging.handlers import RotatingFileHandler

# Ensure log file exists
try:
    os.makedirs(os.path.dirname(APP_ERROR_LOG), exist_ok=True)
    with open(APP_ERROR_LOG, "a", encoding="utf-8") as _f:
        _f.write("\n=== App start ===\n")
except Exception:
    pass

try:
    _err_handler = RotatingFileHandler(APP_ERROR_LOG, maxBytes=2_000_000, backupCount=3, encoding="utf-8")
    _err_handler.setLevel(logging.ERROR)
    app.logger.addHandler(_err_handler)
    app.logger.setLevel(logging.ERROR)
except Exception:
    # Fallback: ignore logging setup errors
    pass

@app.errorhandler(Exception)
def _log_and_respond(e):
    # Write full traceback to the app error log.
    try:
        import traceback
        with open(APP_ERROR_LOG, "a", encoding="utf-8") as _f:
            _f.write("\n=== Uncaught Exception ===\n")
            traceback.print_exc(file=_f)
    except Exception:
        pass
    # Minimal production-like response (keeps 500 behavior consistent)
    return ("Internal Server Error", 500)

# Make sure our handler is registered even if FLASK_ENV differs
try:
    app.register_error_handler(Exception, _log_and_respond)
except Exception:
    pass
app.config['PROPAGATE_EXCEPTIONS'] = False

# Lightweight request logging to the app error log.
@app.before_request
def _log_request_min():
    try:
        with open(APP_ERROR_LOG, "a", encoding="utf-8") as _f:
            _f.write(f"\n[REQ] {request.method} {request.path}\n")
            if request.method == 'POST':
                logged_form = dict(request.form)
                for sensitive_name in ("google_maps_api_key", "api_key", "password", "secret", "token"):
                    if sensitive_name in logged_form:
                        logged_form[sensitive_name] = "[REDACTED]"
                _f.write(f"[FORM] {logged_form}\n")
    except Exception:
        pass

# ---- Jinja filters for number/currency blank if 0 ----
@app.template_filter("num_blank0")
def jinja_num_blank0(val, decimals=0):
    try:
        if val is None:
            return ""
        # NaN check
        if isinstance(val, float) and math.isnan(val):
            return ""
        if float(val) == 0:
            return ""
        fmt = "{:,.%df}" % int(decimals)
        return fmt.format(float(val))
    except Exception:
        return ""

@app.template_filter("currency_blank0")
def jinja_currency_blank0(val, decimals=0):
    s = jinja_num_blank0(val, decimals)
    return f"${s}" if s else ""

# Excel-style rounding (ROUND_HALF_UP) to match Excel's ROUND behavior
def excel_round(value, digits=0):
    try:
        q = Decimal('1') if digits == 0 else Decimal(f'1e-{digits}')
        return float(Decimal(str(value)).quantize(q, rounding=ROUND_HALF_UP))
    except Exception:
        # Fallback: return original value if rounding fails
        return value


def _delete_old_artifacts(proposal_folder: str, remove_pdf: bool = True):
    """Remove generated files before regenerating."""
    patterns = [
        os.path.join(proposal_folder, "Gaco S42 Proposal - *.docx"),
        os.path.join(proposal_folder, "Uniflex Proposal - *.docx"),
        os.path.join(proposal_folder, "Profit Summary - *.xlsx"),
    ]
    if remove_pdf:
        patterns.insert(2, os.path.join(proposal_folder, "Gaco S42 Proposal - *.pdf"))
        patterns.insert(3, os.path.join(proposal_folder, "Uniflex Proposal - *.pdf"))
    for patt in patterns:
        for path in glob.glob(patt):
            try:
                os.remove(path)
            except Exception as e:
                print(f"Warning: could not remove {path}: {e}")


def _is_generated_proposal_artifact(filename: str) -> bool:
    lower_name = str(filename or "").lower()
    return (
        (lower_name.startswith("gaco s42 proposal - ") and lower_name.endswith((".docx", ".pdf")))
        or (lower_name.startswith("uniflex proposal - ") and lower_name.endswith((".docx", ".pdf")))
        or (lower_name.startswith("profit summary - ") and lower_name.endswith(".xlsx"))
    )


def _dated_archive_folder(proposal_folder: str) -> str:
    archive_root = os.path.join(proposal_folder, "Archive")
    os.makedirs(archive_root, exist_ok=True)

    archive_name = datetime.datetime.now().strftime("%m-%d-%Y %H-%M")
    archive_folder = os.path.join(archive_root, archive_name)
    if not os.path.exists(archive_folder):
        return archive_folder

    counter = 2
    while True:
        candidate = os.path.join(archive_root, f"{archive_name} ({counter})")
        if not os.path.exists(candidate):
            return candidate
        counter += 1


def _archive_existing_artifacts(proposal_folder: str):
    """Copy the current proposal folder files to a dated Archive snapshot before regenerating."""
    archive_folder = _dated_archive_folder(proposal_folder)
    os.makedirs(archive_folder, exist_ok=True)

    for entry in os.scandir(proposal_folder):
        if not entry.is_file():
            continue

        source_path = entry.path
        target_path = os.path.join(archive_folder, entry.name)
        try:
            shutil.copy2(source_path, target_path)
        except Exception as e:
            print(f"Warning: could not archive {source_path}: {e}")

    for entry in os.scandir(proposal_folder):
        if not entry.is_file():
            continue
        if not _is_generated_proposal_artifact(entry.name):
            continue
        try:
            os.remove(entry.path)
        except Exception as e:
            print(f"Warning: could not remove archived artifact {entry.path}: {e}")

def _libreoffice_convert_sync(doc_path: str, outdir: str, timeout: int = 180):
    """
    Convert a DOCX to PDF using LibreOffice headless.
    Blocks until done (or raises on failure).
    """
    os.makedirs(outdir, exist_ok=True)
    cmd = [
        LIBREOFFICE_PATH,
        "--headless",
        "--convert-to", "pdf",
        "--outdir", outdir,
        doc_path,
    ]
    try:
        completed = subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=timeout,
            check=True,
            text=True,
        )
        if completed.returncode != 0:
            raise RuntimeError(f"LibreOffice returned {completed.returncode}: {completed.stderr.strip()}")
    except FileNotFoundError:
        raise FileNotFoundError(
            f"LibreOffice not found at {LIBREOFFICE_PATH}. "
            "Install it from libreoffice.org and update LIBREOFFICE_PATH if needed."
        )
    except subprocess.CalledProcessError as e:
        raise RuntimeError(f"LibreOffice conversion failed: {e.stderr or e.stdout}")


def _convert_to_pdf(doc_path: str, outdir: str, use_libreoffice: bool = True, async_mode: bool = True):
    """
    Dispatch PDF conversion. If LibreOffice is available and requested, use it;
    otherwise fall back to docx2pdf (Word). Optionally run async so the UI returns immediately.
    """
    # Prefer LibreOffice when packaged (frozen) to avoid spawning our own app binary via docx2pdf
    try:
        is_frozen = bool(getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"))
    except Exception:
        is_frozen = False
    if is_frozen and os.path.exists(LIBREOFFICE_PATH):
        use_libreoffice = True

    def _worker():
        started = time.perf_counter()
        try:
            pdf_output = os.path.join(
                outdir,
                f"{os.path.splitext(os.path.basename(doc_path))[0]}.pdf",
            )
            if os.path.exists(pdf_output):
                try:
                    os.remove(pdf_output)
                except Exception:
                    pass
            if use_libreoffice and os.path.exists(LIBREOFFICE_PATH):
                _libreoffice_convert_sync(doc_path, outdir)
            else:
                # Fallback to Word/docx2pdf (may pop Word)
                convert(doc_path, outdir)
            _log_timing(f"pdf conversion for {os.path.basename(doc_path)}", started)
        except Exception as e:
            print(f"PDF conversion failed: {e}")

    if async_mode:
        threading.Thread(target=_worker, daemon=True).start()
    else:
        _worker()

# ---- Central Excel mapping & writer ----
EXCEL_CELL_MAP = {
    # Header (known)
    "customer_name": "C1",
    "pcs_or_roofer_ind": "H1",
    "street_address": "L1",
    "city": "R1",
    "state": "T1",
    "zip_code": "V1",
    "squares": "E3",
    "current_roof": "E5",
    "product": "H3",
    "warranty_incl": "H5",
    "submitted_by": "H7",
    "include_travel": None,
    "price_per_sq_10": "M3",
    "price_per_sq_15": "M5",
    "price_per_sq_20": "M7",
    "labor_days": "E7",
    "total_price_10": "P3",
    "total_price_15": "P5",
    "total_price_20": "P7",
    "silicone_units_10": "C11",
    "silicone_units_15": "H11",
    "silicone_units_20": "N11",
    "gaco_patch_units": "C12",
    "bleed_trap_units": "C13",
    "gaco_e5320_units": "C14",
    "sw_1flash_units": "C15",
    "sw_bleed_block_units": "C16",
    "drainage_mat_units": "C17",
    "foam_units": "C18",
    "silicone_price": "D11",
    "gaco_patch_price": "D12",
    "bleed_trap_price": "D13",
    "gaco_e5320_price": "D14",
    "sw_1flash_price": "D15",
    "sw_bleed_block_price": "D16",
    "drainage_mat_price": "D17",
    "foam_price": "D18",
    "rfc_labor_price": "D19",
    "pcs_labor_price": "D21",
    "scarifying_total": "E20",
    "travel_total": "E22",
    "repair_costs_total": "E23",
    "adjusted_coverage": None, 
    "office_fee_pct": None,     
    "lead": "C35",
    "proposal_note": "C36",     
    "proposal_language": "C37",
    # Calculated output cells (explicitly written so Excel file holds values, not formulas)
    "silicone_total": "E11",
    "silicone_15_total": "K11",
    "silicone_20_total": "P11",
    "gaco_patch_total": "E12",
    "bleed_trap_total": "E13",
    "gaco_e5320_total": "E14",
    "sw_1flash_total": "E15",
    "sw_bleed_block_total": "E16",
    "drainage_mat_total": "E17",
    "foam_total": "E18",
    "rfc_labor_total": "E19",
    "pcs_labor_total": "E21",
    "warranty_10_total": "E24",
    "office_fee_total": "E33",
    "office_fee_15_total": "K33",
    "office_fee_20_total": "P33",
    "total_cost": "E27",
    "total_cost_15": "K27",
    "total_cost_20": "P27",
    "pcs_profit": "E29",
    "pcs_profit_15": "K29",
    "pcs_profit_20": "P29",
    "profit_pct": "E30",
    "profit_pct_15": "K30",
    "profit_pct_20": "P30",
    "daily_profit": "E31",
    "daily_profit_15": "K31",
    "daily_profit_20": "P31",
    "profit_share": "E32",
    "profit_share_15": "K32",
    "profit_share_20": "P32",
    "commission_amt": "E25",
    "commission_amt_15": "K25",
    "commission_amt_20": "P25",
}

UNIT_PRICE_FIELDS = {
    "silicone_price",
    "gaco_patch_price",
    "bleed_trap_price",
    "gaco_e5320_price",
    "sw_1flash_price",
    "sw_bleed_block_price",
    "drainage_mat_price",
    "foam_price",
    "rfc_labor_price",
    "pcs_labor_price",
}

PROFIT_SUMMARY_V2_ROW14_MARKER = "Gaco E5320"

def _shift_cell_row(cell_addr: str, row_delta: int, min_row: int) -> str:
    if not cell_addr:
        return cell_addr
    match = re.match(r"^([A-Z]+)(\d+)$", str(cell_addr).strip().upper())
    if not match:
        return cell_addr
    col, row_text = match.groups()
    row = int(row_text)
    if row < min_row:
        return cell_addr
    return f"{col}{row + row_delta}"

def _legacy_profit_summary_cell_map() -> dict:
    """Map fields for Profit Summary files created before the row-14 material insert."""
    return {
        field: _shift_cell_row(cell, -1, 15) if cell else cell
        for field, cell in EXCEL_CELL_MAP.items()
    }

EXCEL_CELL_MAP_LEGACY = _legacy_profit_summary_cell_map()
EXCEL_CELL_MAP_LEGACY.update({
    "office_fee_total": "E24",
    "office_fee_15_total": "K24",
    "office_fee_20_total": "P24",
    "commission_amt": "E32",
    "commission_amt_15": "K32",
    "commission_amt_20": "P32",
    "lead": "C37",
    "proposal_note": "C40",
    "proposal_language": "C41",
})
for _legacy_new_field in ("gaco_e5320_units", "gaco_e5320_price", "gaco_e5320_total"):
    EXCEL_CELL_MAP_LEGACY[_legacy_new_field] = None

def detect_profit_summary_cell_map(ws) -> dict:
    """Choose the correct import/read map using the row-14 Gaco E5320 marker."""
    marker = re.sub(r"\s+", " ", PROFIT_SUMMARY_V2_ROW14_MARKER).strip().casefold()
    try:
        max_col = min(ws.max_column or 1, 40)
        row_values = [ws.cell(row=14, column=col).value for col in range(1, max_col + 1)]
    except Exception:
        row_values = []
    for value in row_values:
        normalized = re.sub(r"\s+", " ", str(value or "")).strip().casefold()
        if marker in normalized:
            return EXCEL_CELL_MAP
    return EXCEL_CELL_MAP_LEGACY

HIDDEN_SHEET_NAME = "Hidden Sheet"
HIDDEN_SHEET_CELL_MAP = {
    "flat_roof_squares": "A5",
    "wall_squares": "A6",
    "include_travel": "A7",
    "adjusted_coverage": "A8",
}
CALC_ONLY_FIELDS = {
    "adjusted_coverage",
    "office_fee_pct",
    "previous_include_travel",
    "previous_calc_travel_total",
    "previous_adjusted_coverage",
    "previous_silicone_units_10",
    "previous_roof_type",
    "previous_product",
    "previous_squares",
    "previous_pcs_or_roofer_ind",
}

def include_travel_from_travel_total(value) -> str:
    """Use the actual E22 travel value as the persisted travel indicator."""
    if value is None:
        return "No"
    if isinstance(value, str):
        cleaned = value.replace("$", "").replace(",", "").strip()
        if cleaned == "":
            return "No"
        try:
            return "Yes" if float(cleaned) != 0 else "No"
        except Exception:
            return "Yes"
    try:
        return "Yes" if float(value) != 0 else "No"
    except Exception:
        return "No"

def get_hidden_sheet(wb, create=False):
    try:
        return wb[HIDDEN_SHEET_NAME]
    except Exception:
        pass

    for ws in getattr(wb, "worksheets", []):
        try:
            if str(ws.title).strip().lower() == HIDDEN_SHEET_NAME.lower():
                return ws
        except Exception:
            continue

    if not create:
        return None

    ws = wb.create_sheet(HIDDEN_SHEET_NAME)
    try:
        ws.sheet_state = "hidden"
    except Exception:
        pass
    return ws

def write_hidden_sheet_values(wb, data: dict):
    ws = get_hidden_sheet(wb, create=True)
    if ws is None:
        return

    for field, cell in HIDDEN_SHEET_CELL_MAP.items():
        if field not in data:
            continue
        value = data[field]
        if value is None:
            continue
        try:
            ws[cell].value = value
        except AttributeError as exc:
            _safe_debug(f"[WARN] Skipping write for {field} at {cell}: {exc}")


def _adjustment_value(adjusted_coverage):
    try:
        adjustment = float(adjusted_coverage or 0.0)
    except Exception:
        adjustment = 0.0
    if isinstance(adjustment, float) and math.isnan(adjustment):
        return 0.0
    return adjustment


def adjusted_coverage_rates(product, roof_type, adjusted_coverage):
    base_rates = coverage_amounts.get(str(product or "").strip(), {}).get(str(roof_type or "").strip(), {})
    adjustment = _adjustment_value(adjusted_coverage)
    return {
        year: float(base_rates.get(year, 0.0) or 0.0) + adjustment
        for year in (10, 15, 20)
    }


def apply_adjusted_spread_rates(wb_profit, product, roof_type, adjusted_coverage):
    adjustment = _adjustment_value(adjusted_coverage)
    if adjustment == 0.0:
        return

    try:
        ws = wb_profit["Data"]
    except Exception:
        return

    product_name = str(product or "").strip()
    product_key = product_name.lower()
    if product_key == "gaco":
        row_start, row_end = 4, 9
    elif product_key == "uniflex":
        row_start, row_end = 13, 18
    else:
        return

    rates = adjusted_coverage_rates(product_name, roof_type, adjustment)
    roof_key = str(roof_type or "").strip().lower()
    for row in range(row_start, row_end + 1):
        if str(ws.cell(row=row, column=2).value or "").strip().lower() != roof_key:
            continue
        for col, year in ((3, 10), (4, 15), (5, 20)):
            ws.cell(row=row, column=col).value = rates[year]
        return


def infer_adjusted_spread_rate(wb_profit, product, roof_type):
    try:
        ws = wb_profit["Data"]
    except Exception:
        return 0.0

    product_key = str(product or "").strip()
    roof_key = str(roof_type or "").strip()
    base_rates = coverage_amounts.get(product_key, {}).get(roof_key, {})
    if not base_rates:
        return 0.0

    product_key_lc = product_key.lower()
    if product_key_lc == "gaco":
        row_start, row_end = 4, 9
    elif product_key_lc == "uniflex":
        row_start, row_end = 13, 18
    else:
        return 0.0

    roof_key_lc = roof_key.lower()
    for row in range(row_start, row_end + 1):
        if str(ws.cell(row=row, column=2).value or "").strip().lower() != roof_key_lc:
            continue
        diffs = []
        for col, year in ((3, 10), (4, 15), (5, 20)):
            try:
                actual = float(ws.cell(row=row, column=col).value or 0.0)
                base = float(base_rates.get(year, 0.0))
            except Exception:
                continue
            diffs.append(actual - base)
        if not diffs:
            return 0.0
        adjustment = diffs[0]
        if any(abs(diff - adjustment) > 1e-6 for diff in diffs):
            _safe_debug(
                "[WARN] Data sheet spread rates differ by term for "
                f"{product_key}/{roof_key}: diffs={diffs}; using 10-year diff"
            )
        if abs(adjustment) <= 1e-9:
            return 0.0
        return round(adjustment, 6)

    return 0.0


def read_hidden_sheet_values(wb, default=0):
    ws = get_hidden_sheet(wb, create=False)
    values = {}
    for field, cell in HIDDEN_SHEET_CELL_MAP.items():
        value = default
        if ws is not None:
            try:
                raw = ws[cell].value
                if raw is not None:
                    value = raw
            except Exception:
                pass
        values[field] = value
    return values

# --- Ghost-cell support: store evaluated numbers alongside formula cells for reliable UI display ---
# Mapping rule: C→AA, D→AB, E→AC, H→AH, K→AK, M→AM, N→AN, P→AP
import re as _re_for_ghost

_GHOST_COL_MAP = {"C": "AA", "D": "AB", "E": "AC", "H": "AH", "K": "AK", "M": "AM", "N": "AN", "P": "AP"}

def _ghost_addr(cell_addr: str) -> str | None:
    """Translate 'C16' -> 'AA16', 'D16' -> 'AB16', etc."""
    if not cell_addr:
        return None
    m = _re_for_ghost.match(r"^([A-Z]+)(\d+)$", str(cell_addr))
    if not m:
        return None
    col, row = m.groups()
    if col in _GHOST_COL_MAP:
        return f"{_GHOST_COL_MAP[col]}{row}"
    return None

from openpyxl.styles import Font

def _set_worksheet_value(ws, cell_addr, value):
    """Set a value, resolving merged-cell addresses to their writable anchor."""
    target = ws[cell_addr]
    try:
        target.value = value
        return
    except AttributeError:
        pass

    try:
        merged_ranges = getattr(getattr(ws, "merged_cells", None), "ranges", [])
        for merged_range in merged_ranges:
            if cell_addr in merged_range:
                ws.cell(merged_range.min_row, merged_range.min_col).value = value
                return
    except Exception:
        pass

    raise AttributeError(f"Unable to write to merged cell {cell_addr}")

def write_ghost_values(wb_profit, values: dict):
    """Write plain evaluated values (no formulas) to ghost cells and set them to white font."""
    try:
        ws = wb_profit.worksheets[0]
    except Exception:
        ws = wb_profit.active

    white_font = Font(color="FFFFFF")  # white text

    for field, val in (values or {}).items():
        cell = EXCEL_CELL_MAP.get(field)
        if not cell:
            continue
        gcell = _ghost_addr(cell)
        if not gcell:
            continue
        try:
            _set_worksheet_value(ws, gcell, val)
            ws[gcell].font = white_font  # make ghost value invisible
        except Exception:
            pass

def ensure_profit_summary_validations(wb_profit):
    """
    Ensure required Profit Summary dropdowns exist after openpyxl write cycles.
    """
    try:
        ws = wb_profit.worksheets[0]
    except Exception:
        ws = wb_profit.active

    target_cell = "E5"
    has_validation = False

    try:
        for dv in ws.data_validations.dataValidation:
            try:
                for rng in getattr(dv, "ranges", []):
                    if target_cell in rng:
                        has_validation = True
                        break
            except Exception:
                continue
            if has_validation:
                break
    except Exception:
        has_validation = False

    if not has_validation:
        roof_dropdown = DataValidation(
            type="list",
            formula1=f"\"{','.join(roof_types)}\"",
            allow_blank=True,
        )
        roof_dropdown.errorTitle = "Invalid Roof Type"
        roof_dropdown.error = "Please select a roof type from the dropdown list."
        roof_dropdown.promptTitle = "Current Roof"
        roof_dropdown.prompt = "Select the current roof type."
        ws.add_data_validation(roof_dropdown)
        roof_dropdown.add(target_cell)


def sync_profit_summary_data_constants(wb_profit):
    """Keep the Profit Summary Data sheet aligned with program constants."""
    try:
        ws = wb_profit["Data"]
    except Exception:
        return
    constants = {
        "J3": PCS_BASE_LABOR_RATE,
        "K8": GACO_S42_BASE_PRICE,
        "K9": GACO_PATCH_BASE_PRICE,
        "K10": GACO_E5320_PRICE,
        "K11": BLEED_TRAP_BASE_PRICE,
        "K12": DRAINAGE_MAT_BASE_PRICE,
        "N8": UNIFLEX_BASE_PRICE,
        "N9": SW_1FLASH_BASE_PRICE,
        "N10": SW_BLEED_BLOCK_BASE_PRICE,
        "N15": GACO_FOAM_BASE_PRICE,
        "N16": UNIFLEX_FOAM_BASE_PRICE,
        "N18": RFC_LABOR_RATE,
    }
    for cell, value in constants.items():
        ws[cell] = value


def write_fields_to_profit_summary(wb_profit, data: dict):
    """
    Writes values from `data` to the first sheet of an openpyxl workbook `wb_profit`
    based on EXCEL_CELL_MAP. Fields with mapping None are skipped.
    """
    try:
        ws = wb_profit.worksheets[0]
    except Exception:
        ws = wb_profit.active
    sync_profit_summary_data_constants(wb_profit)
    for field, cell in EXCEL_CELL_MAP.items():
        if not cell:
            continue
        # Only write when a concrete value is provided; don’t clobber formulas with None
        if field not in data:
            continue
        value = data[field]
        if value is None:
            continue
        # Coerce Decimals to float for Excel
        try:
            from decimal import Decimal
            if isinstance(value, Decimal):
                value = float(value)
        except Exception:
            pass
        _set_worksheet_value(ws, cell, value)
    try:
        _set_worksheet_value(ws, "I14", PS_F_I14.replace("\n", ""))
        _set_worksheet_value(ws, "O14", PS_F_O14.replace("\n", ""))
        _set_worksheet_value(ws, "T14", PS_F_T14.replace("\n", ""))
        _set_worksheet_value(ws, "K14", '=IF(H14<>"",H14*I14,0)')
        _set_worksheet_value(ws, "P14", '=IF(N14<>"",N14*O14,0)')
        _set_worksheet_value(ws, "U14", '=IF(S14<>"",S14*T14,0)')
        _set_worksheet_value(ws, "K25", PS_F_K25.replace("\n", ""))
        _set_worksheet_value(ws, "P25", PS_F_P25.replace("\n", ""))
        _set_worksheet_value(ws, "U25", PS_F_U25.replace("\n", ""))
        _set_worksheet_value(ws, "K27", PS_F_K27.replace("\n", ""))
        _set_worksheet_value(ws, "P27", PS_F_P27.replace("\n", ""))
        _set_worksheet_value(ws, "U27", PS_F_U27.replace("\n", ""))
        _set_worksheet_value(ws, "K29", PS_F_K29.replace("\n", ""))
        _set_worksheet_value(ws, "P29", PS_F_P29.replace("\n", ""))
        _set_worksheet_value(ws, "K30", PS_F_K30.replace("\n", ""))
        _set_worksheet_value(ws, "P30", PS_F_P30.replace("\n", ""))
        _set_worksheet_value(ws, "K31", PS_F_K31.replace("\n", ""))
        _set_worksheet_value(ws, "P31", PS_F_P31.replace("\n", ""))
        _set_worksheet_value(ws, "K32", PS_F_K32.replace("\n", ""))
        _set_worksheet_value(ws, "P32", PS_F_P32.replace("\n", ""))
        _set_worksheet_value(ws, "K33", PS_F_K33.replace("\n", ""))
        _set_worksheet_value(ws, "P33", PS_F_P33.replace("\n", ""))
        _set_worksheet_value(ws, "U33", PS_F_U33.replace("\n", ""))
    except Exception:
        pass
    write_hidden_sheet_values(wb_profit, data)


def sync_cost_rollups_from_display_components(data: dict) -> dict:
    """Keep cost/profit rollups aligned with the line-item totals shown in proposal detail."""
    if not isinstance(data, dict):
        return data

    def _num(v, default=0.0):
        try:
            if v is None:
                return default
            if isinstance(v, str):
                s = v.replace("$", "").replace(",", "").strip()
                if s == "":
                    return default
                return float(s)
            if isinstance(v, float) and math.isnan(v):
                return default
            return float(v)
        except Exception:
            return default

    shared_component_fields = (
        "gaco_patch_total",
        "bleed_trap_total",
        "gaco_e5320_total",
        "sw_1flash_total",
        "sw_bleed_block_total",
        "drainage_mat_total",
        "foam_total",
        "rfc_labor_total",
        "pcs_labor_total",
        "scarifying_total",
        "travel_total",
        "repair_costs_total",
    )
    labor_days = _num(data.get("labor_days"), 0.0)

    year_specs = (
        ("10", "silicone_total", "warranty_10_total", "commission_amt", "total_price_10", "total_cost", "profit_share", "pcs_profit", "profit_pct", "daily_profit"),
        ("15", "silicone_15_total", "warranty_15_total", "commission_amt_15", "total_price_15", "total_cost_15", "profit_share_15", "pcs_profit_15", "profit_pct_15", "daily_profit_15"),
        ("20", "silicone_20_total", "warranty_20_total", "commission_amt_20", "total_price_20", "total_cost_20", "profit_share_20", "pcs_profit_20", "profit_pct_20", "daily_profit_20"),
    )

    for _, silicone_field, warranty_field, commission_field, price_field, cost_field, share_field, pcs_field, pct_field, daily_field in year_specs:
        total_cost = (
            _num(data.get(silicone_field), 0.0)
            + sum(_num(data.get(field), 0.0) for field in shared_component_fields)
            + _num(data.get(warranty_field), 0.0)
            + _num(data.get(commission_field), 0.0)
        )
        data[cost_field] = total_cost

        total_price = _num(data.get(price_field), 0.0)
        profit_share_amt = excel_round(PROFIT_SHARE_PCT * (total_price - total_cost), 0)
        pcs_profit = total_price - total_cost - profit_share_amt
        data[share_field] = profit_share_amt
        data[pcs_field] = pcs_profit
        data[pct_field] = excel_round(pcs_profit / total_price, 2) if total_price else 0
        data[daily_field] = excel_round(pcs_profit / labor_days, 0) if labor_days else 0
    return data


# ---- Read existing Excel and build display data with fallbacks ----
def read_profit_summary_for_display(folder_path: str) -> dict | None:
    """Read Profit Summary.xlsx and build a display dict.
    If a cell contains a formula (no cached value yet), compute a fallback via calculation_routine
    so the proposal detail screen shows numbers immediately.
    """
    excel_path = find_profit_summary_file(folder_path)
    if not excel_path or not os.path.exists(excel_path):
        return None

    # Load two views: cached values (data_only=True) and raw (to detect formulas)
    try:
        wb_vals = load_workbook(excel_path, data_only=True)
        wb_raw  = load_workbook(excel_path, data_only=False)
    except Exception:
        return None
    ws_vals = wb_vals.worksheets[0]
    ws_raw  = wb_raw.worksheets[0]
    cell_map = detect_profit_summary_cell_map(ws_vals)

    def _get_ghost_val(field, default=None):
        cell = cell_map.get(field)
        if not cell:
            return default
        gcell = _ghost_addr(cell)
        if not gcell:
            return default
        try:
            gv = ws_vals[gcell].value
            return gv if gv is not None else default
        except Exception:
            return default

    def _get_cell_val(field, default=None):
        cell = cell_map.get(field)
        if not cell:
            return default
        try:
            v = ws_vals[cell].value
            return v if v is not None else default
        except Exception:
            return default

    def _raw_cell_is_formula(field) -> bool:
        cell = cell_map.get(field)
        if not cell:
            return False
        try:
            return getattr(ws_raw[cell], "data_type", None) == "f"
        except Exception:
            return False

    def _get_calc_input_val(field, default=0.0, treat_zero_as_missing=False):
        value = _get_cell_val(field, None)
        missing = value is None
        if not missing and isinstance(value, str) and value.strip() == "":
            missing = True
        if not missing and treat_zero_as_missing:
            try:
                if isinstance(value, float) and math.isnan(value):
                    missing = True
                else:
                    missing = float(value) == 0.0
            except Exception:
                missing = False
        if missing:
            return default
        return value

    # Pull inputs needed for calc (mirror of POST parse basics)
    def _num(v, d=0.0):
        try:
            if v is None:
                return d
            if isinstance(v, str):
                s = v.replace('$','').replace(',','').strip()
                if s == '':
                    return d
                return float(s)
            return float(v)
        except Exception:
            return d

    def _int(v, d=0):
        try:
            return int(_num(v, d))
        except Exception:
            return d

    hidden_values = read_hidden_sheet_values(wb_vals, default=0)
    squares = _num(_get_cell_val('squares'), 0.0)
    product = str(_get_cell_val('product') or '')
    roof_type = str(_get_cell_val('current_roof') or '')
    labor_days = _int(_get_cell_val('labor_days'), None)
    warranty_incl = str(_get_cell_val('warranty_incl') or 'No')
    travel_total = _num(_get_cell_val('travel_total'), 0.0)
    include_travel = include_travel_from_travel_total(travel_total)
    price_per_sq_10 = _num(_get_cell_val('price_per_sq_10'), 0.0)
    commission_pct = _num(_get_cell_val('commission_pct'), 0.0)
    submitted_by = str(_get_cell_val('submitted_by') or '')
    office_fee_pct = _num(_get_cell_val('office_fee_pct'), 0.0)
    pcs_or_roofer_ind = str(_get_cell_val('pcs_or_roofer_ind') or '')
    adjusted_coverage = infer_adjusted_spread_rate(wb_vals, product, roof_type)
    hidden_values["adjusted_coverage"] = adjusted_coverage

    # Units & prices
    silicone_units_10 = _num(_get_calc_input_val('silicone_units_10', None), None)
    silicone_price    = _num(_get_calc_input_val('silicone_price', None), None)
    gaco_patch_units  = _num(_get_calc_input_val('gaco_patch_units', None), None)
    gaco_patch_price  = _num(_get_calc_input_val('gaco_patch_price', None), None)
    bleed_trap_units  = _num(_get_calc_input_val('bleed_trap_units', None), None)
    bleed_trap_price  = _num(_get_calc_input_val('bleed_trap_price', None), None)
    gaco_e5320_units  = _num(_get_calc_input_val('gaco_e5320_units', 0.0), 0.0)
    gaco_e5320_price  = _num(_get_calc_input_val('gaco_e5320_price', GACO_E5320_PRICE), GACO_E5320_PRICE)
    sw_1flash_units   = _num(_get_calc_input_val('sw_1flash_units', None), None)
    sw_1flash_price   = _num(_get_calc_input_val('sw_1flash_price', None), None)
    sw_bleed_block_units = _num(_get_calc_input_val('sw_bleed_block_units', None), None)
    sw_bleed_block_price = _num(_get_calc_input_val('sw_bleed_block_price', None), None)
    drainage_mat_units   = _num(_get_calc_input_val('drainage_mat_units', None), None)
    drainage_mat_price   = _num(_get_calc_input_val('drainage_mat_price', None), None)
    foam_units           = _num(_get_calc_input_val('foam_units', None), None)
    foam_price           = _num(_get_calc_input_val('foam_price', None), None)
    rfc_labor_price      = _num(_get_calc_input_val('rfc_labor_price', None), None)
    pcs_labor_price      = _num(_get_calc_input_val('pcs_labor_price', None), None)
    scarifying_total     = _num(_get_cell_val('scarifying_total'), 0.0)
    repair_costs_total     = _num(_get_cell_val('repair_costs_total'), 0.0)

    if _raw_cell_is_formula("silicone_price"):
        silicone_price = None
    if _raw_cell_is_formula("gaco_patch_price"):
        gaco_patch_price = None
    if _raw_cell_is_formula("bleed_trap_price"):
        bleed_trap_price = None
    if _raw_cell_is_formula("gaco_e5320_price"):
        gaco_e5320_price = None
    if _raw_cell_is_formula("sw_1flash_price"):
        sw_1flash_price = None
    if _raw_cell_is_formula("sw_bleed_block_price"):
        sw_bleed_block_price = None
    if _raw_cell_is_formula("drainage_mat_price"):
        drainage_mat_price = None
    if _raw_cell_is_formula("foam_price"):
        foam_price = None
    if _raw_cell_is_formula("rfc_labor_price"):
        rfc_labor_price = None
    if _raw_cell_is_formula("pcs_labor_price"):
        pcs_labor_price = None

    # Call calculation_routine to get authoritative computed values
    calc_res = calculation_routine(
        squares=squares,
        product=product,
        roof_type=roof_type,
        labor_days=labor_days,
        warranty_incl=warranty_incl,
        include_travel=include_travel,
        price_per_sq_10=price_per_sq_10,
        commission_pct=commission_pct,
        submitted_by=submitted_by,
        previous_submitted_by=submitted_by,
        office_fee_pct=office_fee_pct,
        adjusted_coverage=adjusted_coverage,
        silicone_units_10=silicone_units_10,
        silicone_price=silicone_price,
        gaco_patch_units=gaco_patch_units,
        gaco_patch_price=gaco_patch_price,
        sw_1flash_units=sw_1flash_units,
        sw_1flash_price=sw_1flash_price,
        bleed_trap_units=bleed_trap_units,
        bleed_trap_price=bleed_trap_price,
        gaco_e5320_units=gaco_e5320_units,
        gaco_e5320_price=gaco_e5320_price,
        sw_bleed_block_units=sw_bleed_block_units,
        sw_bleed_block_price=sw_bleed_block_price,
        drainage_mat_units=drainage_mat_units,
        drainage_mat_price=drainage_mat_price,
        foam_units=foam_units,
        foam_price=foam_price,
        rfc_labor_price=rfc_labor_price,
        pcs_labor_price=pcs_labor_price,
        scarifying_total=scarifying_total,
        travel_total=travel_total,
        repair_costs_total=repair_costs_total,
        previous_squares=squares,
        previous_roof_type=roof_type,
        previous_product=product,
        previous_adjusted_coverage=adjusted_coverage,
        previous_silicone_units_10=silicone_units_10,
        proposal_note=str(_get_cell_val('proposal_note') or ''),
        pcs_or_roofer_ind=pcs_or_roofer_ind,
        previous_pcs_or_roofer_ind=pcs_or_roofer_ind,
        previous_include_travel=include_travel,
        previous_calc_travel_total=0,
    )

    # Build base data dict from cached values in the sheet
    data = {}
    for field, cell in cell_map.items():
        if not cell:
            continue
        try:
            v = ws_vals[cell].value
        except Exception:
            v = None
        data[field] = v
    data.update(hidden_values)
    data["include_travel"] = include_travel
    data["calc_travel_total"] = calc_res.get("calc_travel_total", 0)
    data["previous_calc_travel_total"] = data["calc_travel_total"]
    data["previous_include_travel"] = data.get("include_travel") or "No"

    # Fallback fields likely to be formulas with missing cached results
    _fallback_fields = [
        # Inputs that may be stored as formulas (or blank when N/A) but must show a number on screen
        "labor_days",
        "silicone_units_10","silicone_units_15","silicone_units_20","gaco_patch_units","bleed_trap_units","gaco_e5320_units",
        "sw_1flash_units","sw_bleed_block_units","drainage_mat_units","foam_units",

        # Unit prices that may be formulas or overridden numbers
        "silicone_price","gaco_patch_price","bleed_trap_price","gaco_e5320_price",
        "sw_1flash_price","sw_bleed_block_price","drainage_mat_price",
        "foam_price","rfc_labor_price","pcs_labor_price",

        # Price per square and totals
        "price_per_sq_10","price_per_sq_15","price_per_sq_20",
        "total_price_10","total_price_15","total_price_20",

        # Cost/fee totals and downstream profit metrics
        "total_cost","total_cost_15","total_cost_20","warranty_10_total","office_fee_total","office_fee_15_total","office_fee_20_total",
        "silicone_total","silicone_15_total","silicone_20_total","gaco_patch_total","bleed_trap_total","gaco_e5320_total",
        "sw_1flash_total","sw_bleed_block_total","drainage_mat_total",
        "foam_total","rfc_labor_total","pcs_labor_total",
        "commission_amt","commission_amt_15","commission_amt_20",
        "profit_share","profit_share_15","profit_share_20",
        "daily_profit","daily_profit_15","daily_profit_20",
        "profit_pct","profit_pct_15","profit_pct_20",
        "pcs_profit","pcs_profit_15","pcs_profit_20"
    ]

    def _is_blank(v):
        if v is None:
            return True
        if isinstance(v, str) and v.strip() == "":
            return True
        return False

    def _is_blank_or_zero(v):
        if _is_blank(v):
            return True
        try:
            if isinstance(v, float) and math.isnan(v):
                return True
            return float(v) == 0.0
        except Exception:
            return False

    for f in _fallback_fields:
        cell = cell_map.get(f)
        if not cell:
            continue
        raw_cell = ws_raw[cell]
        val = data.get(f)
        if f in UNIT_PRICE_FIELDS and getattr(raw_cell, 'data_type', None) == 'f' and f in calc_res:
            data[f] = calc_res[f]
            continue
        # If raw cell is a formula and cached value is blank/0 -> use computed fallback
        if getattr(raw_cell, 'data_type', None) == 'f' and _is_blank_or_zero(val):
            if f in calc_res:
                data[f] = calc_res[f]
            else:
                gval = _get_ghost_val(f, None)
                if not _is_blank_or_zero(gval):
                    data[f] = gval
        # Additionally, if there's no formula but value is None/blank, also fill
        elif _is_blank(val) and f in calc_res:
            data[f] = calc_res[f]

    # These are readonly derived outputs on the details screen.
    # Prefer freshly calculated readonly values. Hidden ghost values are only a
    # fallback because they can be stale after spread-rate changes.
    _derived_display_fields = [
        "silicone_units_15","silicone_units_20",
        "price_per_sq_15","price_per_sq_20",
        "total_price_10","total_price_15","total_price_20",
        "warranty_10_total","office_fee_total","office_fee_15_total","office_fee_20_total",
        "total_cost","total_cost_15","total_cost_20",
        "silicone_total","silicone_15_total","silicone_20_total","gaco_patch_total","bleed_trap_total","gaco_e5320_total",
        "sw_1flash_total","sw_bleed_block_total","drainage_mat_total",
        "foam_total","rfc_labor_total","pcs_labor_total",
        "commission_amt","commission_amt_15","commission_amt_20",
        "profit_share","profit_share_15","profit_share_20",
        "daily_profit","daily_profit_15","daily_profit_20",
        "profit_pct","profit_pct_15","profit_pct_20",
        "pcs_profit","pcs_profit_15","pcs_profit_20",
    ]
    for f in _derived_display_fields:
        if f in calc_res:
            data[f] = calc_res[f]
            continue
        gval = _get_ghost_val(f, None)
        if not _is_blank_or_zero(gval):
            data[f] = gval
        elif f in calc_res:
            data[f] = calc_res[f]

    sync_cost_rollups_from_display_components(data)
    return data

# Merge helper to update a template data dict with display fallbacks for an existing folder
def merge_display_fallbacks(
    data: dict,
    folder_path: str,
    folder_name: str,
    prefer_saved_derived: bool = True,
) -> dict:
    try:
        if folder_name and folder_name not in ("NEW", "__blank__") and os.path.isdir(folder_path):
            display_data = read_profit_summary_for_display(folder_path)
            if display_data:
                always_sync_fields = {
                    "price_per_sq_15", "price_per_sq_20",
                    "total_price_10", "total_price_15", "total_price_20",
                    "warranty_10_total", "office_fee_total", "office_fee_15_total", "office_fee_20_total",
                    "total_cost", "total_cost_15", "total_cost_20",
                    "silicone_total", "silicone_15_total", "silicone_20_total", "gaco_patch_total", "bleed_trap_total", "gaco_e5320_total",
                    "sw_1flash_total", "sw_bleed_block_total", "drainage_mat_total",
                    "foam_total", "rfc_labor_total", "pcs_labor_total",
                    "commission_amt", "commission_amt_15", "commission_amt_20",
                    "profit_share", "profit_share_15", "profit_share_20",
                    "daily_profit", "daily_profit_15", "daily_profit_20",
                    "profit_pct", "profit_pct_15", "profit_pct_20",
                    "pcs_profit", "pcs_profit_15", "pcs_profit_20",
                }
                unit_price_fields = globals().get("UNIT_PRICE_FIELDS", {
                    "silicone_price",
                    "gaco_patch_price",
                    "bleed_trap_price",
                    "gaco_e5320_price",
                    "sw_1flash_price",
                    "sw_bleed_block_price",
                    "drainage_mat_price",
                    "foam_price",
                    "rfc_labor_price",
                    "pcs_labor_price",
                })
                always_sync_fields.update(unit_price_fields)

                if prefer_saved_derived:
                    for k in always_sync_fields:
                        if k in display_data and display_data[k] is not None:
                            data[k] = display_data[k]

                # Only backfill computed display fields when the current value is missing.
                # Never overwrite user-edited driver/header fields during POST recalcs.
                fallback_fields = {
                    "labor_days",
                    "silicone_units_10", "silicone_units_15", "silicone_units_20", "gaco_patch_units", "bleed_trap_units", "gaco_e5320_units",
                    "sw_1flash_units", "sw_bleed_block_units", "drainage_mat_units", "foam_units",
                    "silicone_price", "gaco_patch_price", "bleed_trap_price", "gaco_e5320_price",
                    "sw_1flash_price", "sw_bleed_block_price", "drainage_mat_price",
                    "foam_price", "rfc_labor_price", "pcs_labor_price",
                    "price_per_sq_10", "price_per_sq_15", "price_per_sq_20",
                    "total_price_10", "total_price_15", "total_price_20",
                    "total_cost", "total_cost_15", "total_cost_20",
                    "warranty_10_total", "office_fee_total", "office_fee_15_total", "office_fee_20_total",
                    "silicone_total", "silicone_15_total", "silicone_20_total", "gaco_patch_total", "bleed_trap_total", "gaco_e5320_total",
                    "sw_1flash_total", "sw_bleed_block_total", "drainage_mat_total",
                    "foam_total", "rfc_labor_total", "pcs_labor_total",
                    "commission_amt", "commission_amt_15", "commission_amt_20",
                    "profit_share", "profit_share_15", "profit_share_20",
                    "daily_profit", "daily_profit_15", "daily_profit_20",
                    "profit_pct", "profit_pct_15", "profit_pct_20",
                    "pcs_profit", "pcs_profit_15", "pcs_profit_20",
                }

                def _is_missing(v):
                    if v is None:
                        return True
                    if isinstance(v, float) and math.isnan(v):
                        return True
                    if isinstance(v, str) and v.strip() == "":
                        return True
                    if not prefer_saved_derived:
                        return False
                    try:
                        return float(v) == 0.0
                    except Exception:
                        return False

                for k in fallback_fields:
                    if k in display_data and display_data[k] is not None and _is_missing(data.get(k)):
                        data[k] = display_data[k]
    except Exception:
        pass
    return data

# ---- Blank defaults for starting without Excel ----
def make_blank_data():
    return {
        "flat_roof_squares": 0,
        "wall_squares": 0,
        "squares": 0,
        "product": "",                   # force user to choose
        "current_roof": "",               # force user to choose
        "warranty_incl": "No",
        "include_travel": "No",
        "calc_travel_total": 0,
        "previous_include_travel": "No",
        "previous_calc_travel_total": 0,
        "labor_days": 0,
        "commission_pct": 0,
        "submitted_by": "",               # force user to choose
        "price_per_sq_10": 0,
        "price_per_sq_15": 0,
        "price_per_sq_20": 0,
        "total_price_10": 0,
        "total_price_15": 0,
        "total_price_20": 0,
        "silicone_units_10": 0,
        "silicone_units_15": 0,
        "silicone_units_20": 0,
        "silicone_price": 0,
        "gaco_patch_units": 0,
        "gaco_patch_price": 0,
        "bleed_trap_units": 0,
        "bleed_trap_price": 0,
        "gaco_e5320_units": 0,
        "gaco_e5320_price": 0,
        "gaco_e5320_total": 0,
        "sw_1flash_units": 0,
        "sw_1flash_price": 0,
        "sw_bleed_block_units": 0,
        "sw_bleed_block_price": 0,
        "drainage_mat_units": 0,
        "drainage_mat_price": 0,
        "foam_units": 0,
        "foam_price": 0,
        "rfc_labor_price": 0,
        "pcs_labor_price": 0,
        "scarifying_total": 0,
        "travel_total": 0,
        "repair_costs_total": 0,
        "warranty_10_total": 0,
        "warranty_15_total": 0,
        "warranty_20_total": 0,
        "office_fee_total": 0,
        "office_fee_15_total": 0,
        "office_fee_20_total": 0,
        "total_cost": 0,
        "total_cost_15": 0,
        "total_cost_20": 0,
        "pcs_labor_total": 0,
        "rfc_labor_total": 0,
        "pcs_profit": 0,
        "pcs_profit_15": 0,
        "pcs_profit_20": 0,
        "profit_pct": 0,
        "profit_pct_15": 0,
        "profit_pct_20": 0,
        "daily_profit": 0,
        "daily_profit_15": 0,
        "daily_profit_20": 0,
        "profit_share": 0,
        "profit_share_15": 0,
        "profit_share_20": 0,
        "commission_amt": 0,
        "commission_amt_15": 0,
        "commission_amt_20": 0,
        "coverage_10": 0,
        "coverage_15": 0,
        "coverage_20": 0,
        "adjusted_coverage": 0,
        "office_fee_pct": None,  # None so calc uses Submitted By default
        "previous_squares": 0,
        "previous_roof_type": "",         
        "previous_product": "",
        "previous_warranty_incl": "No",
        "previous_adjusted_coverage": 0,
        "previous_submitted_by": "",
        "proposal_note": "",
        "street_address": "",
        "city": "",
        "state": "",
        "zip_code": "",
        "lead": "",
        "pcs_or_roofer_ind": "",
        "previous_pcs_or_roofer_ind": "",
    }


def proposal_customer_organization_names() -> list[str]:
    """Return active tenant organization names for proposal autocomplete."""
    try:
        organizations = get_contact_store().list_organizations()
    except (ContactStoreError, TenantAuthenticationError) as exc:
        _safe_debug(f"[WARN] Could not load proposal customer organizations: {exc}")
        return []
    names_by_key = {}
    for organization in organizations:
        name = " ".join(str(organization.get("name") or "").split())
        if name:
            names_by_key.setdefault(name.casefold(), name)
    return sorted(names_by_key.values(), key=str.casefold)



def calculation_routine(
    squares,
    product,
    roof_type,
    labor_days,
    warranty_incl,
    include_travel,
    price_per_sq_10,
    commission_pct,
    submitted_by,
    previous_submitted_by,
    office_fee_pct,
    adjusted_coverage,
    silicone_units_10,
    silicone_price,
    gaco_patch_units,
    gaco_patch_price,
    sw_1flash_units,
    sw_1flash_price,
    bleed_trap_units,
    bleed_trap_price,
    sw_bleed_block_units,
    sw_bleed_block_price,
    drainage_mat_units,
    drainage_mat_price,
    foam_units,
    foam_price,
    rfc_labor_price,
    pcs_labor_price,
    scarifying_total,
    travel_total,
    repair_costs_total,
    previous_squares,
    previous_roof_type,
    previous_product,
    previous_adjusted_coverage,
    previous_silicone_units_10,
    proposal_note,
    pcs_or_roofer_ind=None,
    previous_pcs_or_roofer_ind=None,
    previous_include_travel="No",
    previous_calc_travel_total=None,
    gaco_e5320_units=0.0,
    gaco_e5320_price=None,
):
    # === Select correct pricing arrays based on PCS/Roofer indicator ===
    _safe_debug(f"[TRACE] Entering pricing selector: incoming pcs_or_roofer_ind={pcs_or_roofer_ind!r}")
    try:
        ind = (pcs_or_roofer_ind or "").strip()
    except Exception:
        ind = ""
    _safe_debug(f"[DEBUG] pcs_or_roofer_ind raw={pcs_or_roofer_ind!r}, normalized={ind!r}")
    try:
        prev_ind = (previous_pcs_or_roofer_ind or "").strip()
    except Exception:
        prev_ind = ""
    _safe_debug(f"[DEBUG] previous_pcs_or_roofer_ind raw={previous_pcs_or_roofer_ind!r}, normalized={prev_ind!r}")

    if ind == "PCS Direct":
        pricing10_local = pcs_pricing10
        pricing15_local = pcs_pricing15
        pricing20_local = pcs_pricing20
    else:
        pricing10_local = roofer_pricing10
        pricing15_local = roofer_pricing15
        pricing20_local = roofer_pricing20
    # Labor days logic (preserve posted value; compute baseline; force recalc on key changes)
    if roof_type in ["Ballasted 60 mil", "Ballasted 45 mil"]:
        calc_labor_days = math.ceil(squares / 30)
    elif roof_type == "Rock/Foam/Coat":
        calc_labor_days = math.ceil(squares / 75)
    else:
        calc_labor_days = math.ceil(squares / 45)

    # Trigger a reset to the baseline when roof type or squares changed
    labor_days_recalc = (previous_roof_type != roof_type) or (previous_squares != squares)

    def _is_blank_zero_or_nan_int(v):
        if v is None:
            return True
        try:
            if isinstance(v, float) and math.isnan(v):
                return True
            s = str(v).strip()
            if s == "":
                return True
            return int(float(s)) == 0
        except Exception:
            return True

    if labor_days_recalc:
        # Force reset to baseline when core drivers changed
        labor_days = calc_labor_days
    else:
        # Otherwise, if user left it blank/0, use the baseline
        if _is_blank_zero_or_nan_int(labor_days):
            labor_days = calc_labor_days

    # Override flag: true when the posted value differs from the calculated baseline
    try:
        ov_labor_days = int(labor_days) != int(calc_labor_days)
    except Exception:
        ov_labor_days = False
    _safe_debug(f"[DEBUG] calc_labor_days={calc_labor_days}, labor_days={labor_days}, ov_labor_days={ov_labor_days}, labor_days_recalc={labor_days_recalc}")

    # Set price_per_sq_* with safe defaults. Allow user override only for 10-yr price.
    # 15/20 are always derived from the pricing tables based on roof_type.
    def _is_blank_zero_or_nan(v):
        if v is None:
            return True
        if isinstance(v, float) and math.isnan(v):
            return True
        try:
            return float(v) == 0.0
        except (TypeError, ValueError):
            return True

    try:
        roof_type_index = roof_types.index(roof_type)
        base_pps10 = pricing10_local[roof_type_index]
        base_pps15 = pricing15_local[roof_type_index]
        base_pps20 = pricing20_local[roof_type_index]
    except ValueError:
        # Unknown roof type: fall back to zeros to avoid UnboundLocalError
        base_pps10 = 0
        base_pps15 = 0
        base_pps20 = 0

    # If the roof type OR PCS/Roofer indicator changed, reset 10/15/20 to base.
    if (previous_roof_type != roof_type) or (ind != prev_ind):
        price_per_sq_10 = base_pps10
        price_per_sq_15 = base_pps15
        price_per_sq_20 = base_pps20
    else:
        # Respect user-entered 10-yr price when provided; otherwise use base.
        user_pps10 = price_per_sq_10 if not _is_blank_zero_or_nan(price_per_sq_10) else base_pps10
        # Apply the same delta from base_10 to 15 and 20 so they "recalculate" in line with the override
        delta10 = 0.0
        try:
            delta10 = float(user_pps10) - float(base_pps10)
        except Exception:
            delta10 = 0.0
        price_per_sq_10 = user_pps10
        price_per_sq_15 = float(base_pps15) + delta10
        price_per_sq_20 = float(base_pps20) + delta10

    # Look up coverage factors, adjusted from the application constants.
    coverage_factors = adjusted_coverage_rates(product, roof_type, adjusted_coverage)
    coverage_10 = coverage_factors.get(10, 0)
    coverage_15 = coverage_factors.get(15, 0)
    coverage_20 = coverage_factors.get(20, 0)

    # Calculate current coverage-based units
    calc_units_10 = (squares / 5) * coverage_10
    calc_units_15 = (squares / 5) * coverage_15
    calc_units_20 = (squares / 5) * coverage_20

    # Baseline (calculated) 10-yr silicone units, rounded up
    try:
        calc_silicone_units_10 = int(math.ceil(float(calc_units_10 or 0)))
    except Exception:
        calc_silicone_units_10 = 0

    # Silicone units logic
    def _norm_adj(v):
        if v is None or (isinstance(v, float) and math.isnan(v)):
            return 0.0
        try:
            return float(v)
        except (TypeError, ValueError):
            return 0.0

    def _almost_equal(a, b, tol=1e-6):
        try:
            return abs(float(a) - float(b)) <= tol
        except Exception:
            return False

    adjusted_coverage_changed = _norm_adj(adjusted_coverage) != _norm_adj(previous_adjusted_coverage)
    spread_driver_changed = (
        (previous_product != product)
        or (previous_roof_type != roof_type)
        or (squares != previous_squares)
        or adjusted_coverage_changed
    )

    # Detect manual change of silicone units in THIS submit.
    # If a spread driver changed, the 10-year units may be stale because product,
    # roof type, square count, or spread rate changed. Recalculate instead of
    # treating the posted old units as a manual override.
    user_changed_units = (
        silicone_units_10 is not None
        and not (isinstance(silicone_units_10, float) and math.isnan(silicone_units_10))
        and not _almost_equal(silicone_units_10, previous_silicone_units_10)
        and not spread_driver_changed
    )

    # If the user manually entered silicone units, adjusted_coverage is ignored/reset
    if user_changed_units:
        adjusted_coverage = 0.0
        base_cov_factors = coverage_amounts.get(product, {}).get(roof_type, {})
        coverage_10 = base_cov_factors.get(10, 0)
        coverage_15 = base_cov_factors.get(15, 0)
        coverage_20 = base_cov_factors.get(20, 0)
        calc_units_10 = (squares / 5) * coverage_10
        calc_units_15 = (squares / 5) * coverage_15
        calc_units_20 = (squares / 5) * coverage_20

    # Recalc when any spread driver changes. Manual 10-year unit edits only apply
    # when the spread drivers themselves are unchanged.
    recalc_trigger = spread_driver_changed

    if recalc_trigger:
        silicone_units_10 = calc_units_10
    else:
        def _is_blank_zero_or_nan(v):
            if v is None:
                return True
            if isinstance(v, float) and math.isnan(v):
                return True
            try:
                return float(v) == 0.0
            except (TypeError, ValueError):
                return True
        if _is_blank_zero_or_nan(silicone_units_10):
            silicone_units_10 = calc_units_10

    silicone_units_15 = calc_units_15
    silicone_units_20 = calc_units_20

    # If user overrode 10-yr units, derive 15/20 from 10 using coverage ratios
    if user_changed_units:
        if coverage_10:
            silicone_units_15 = silicone_units_10 * (coverage_15 / coverage_10)
            silicone_units_20 = silicone_units_10 * (coverage_20 / coverage_10)
        else:
            silicone_units_15 = silicone_units_10
            silicone_units_20 = silicone_units_10

    # Normalize silicone units to whole numbers by **rounding up** (ceiling)
    try:
        silicone_units_10 = math.ceil(float(silicone_units_10 or 0))
    except Exception:
        silicone_units_10 = 0
    try:
        silicone_units_15 = math.ceil(float(silicone_units_15 or 0))
    except Exception:
        silicone_units_15 = 0
    try:
        silicone_units_20 = math.ceil(float(silicone_units_20 or 0))
    except Exception:
        silicone_units_20 = 0
    
    # Determine if user has overridden silicone units (10-year) relative to baseline
    try:
        ov_silicone_units_10 = int(silicone_units_10) != int(calc_silicone_units_10)
    except Exception:
        ov_silicone_units_10 = False
    _safe_debug(f"[DEBUG] calc_sil_units_10={calc_silicone_units_10}, units_10={silicone_units_10}, ov_sil_units_10={ov_silicone_units_10}, user_changed_units={user_changed_units}, recalc_trigger={recalc_trigger}")

    # Silicone price logic:
    base_silicone_price = (
        GACO_S42_BASE_PRICE if product == "Gaco" else (
            UNIFLEX_BASE_PRICE if product == "Uniflex" else silicone_price
        )
    )
    if previous_product != product:
        silicone_price = base_silicone_price
    else:
        if silicone_price is None or (isinstance(silicone_price, float) and math.isnan(silicone_price)):
            silicone_price = base_silicone_price

    # Gaco patch units logic (units depend on product, roof type, and squares)
    if product == "Gaco":
        base_gaco_patch_units = (
            math.ceil(squares * 0.03)
            if roof_type == "Rock/Foam/Coat"
            else math.ceil(squares / 10)
        )
    else:
        base_gaco_patch_units = 0

    gaco_patch_recalc = (
        (previous_product != product)
        or (previous_roof_type != roof_type)
        or (previous_squares != squares)
    )

    if gaco_patch_recalc:
        gaco_patch_units = base_gaco_patch_units
    else:
        if gaco_patch_units is None or (isinstance(gaco_patch_units, float) and math.isnan(gaco_patch_units)):
            gaco_patch_units = base_gaco_patch_units

    # Gaco patch price logic
    base_gaco_patch_price = GACO_PATCH_BASE_PRICE if product == "Gaco" else 0

    if previous_product != product:
        gaco_patch_price = base_gaco_patch_price
    else:
        if gaco_patch_price is None or (isinstance(gaco_patch_price, float) and math.isnan(gaco_patch_price)):
            gaco_patch_price = base_gaco_patch_price
    
    # Bleed Trap logic (units & price)
    if product == "Gaco" and roof_type == "Mod Bit":
        base_bleed_units = math.ceil(squares / 5)
        base_bleed_price = BLEED_TRAP_BASE_PRICE
    else:
        base_bleed_units = 0
        base_bleed_price = 0

    bleed_recalc_trigger = (
        (previous_product != product)
        or (previous_roof_type != roof_type)
        or (previous_squares != squares)
    )

    if bleed_recalc_trigger:
        bleed_trap_units = base_bleed_units
        bleed_trap_price = base_bleed_price
    else:
        if bleed_trap_units is None or (isinstance(bleed_trap_units, float) and math.isnan(bleed_trap_units)):
            bleed_trap_units = base_bleed_units
        if bleed_trap_price is None or (isinstance(bleed_trap_price, float) and math.isnan(bleed_trap_price)):
            bleed_trap_price = base_bleed_price

    # Gaco E5320 logic. Units are user-entered; price defaults to baseline when units are present.
    try:
        gaco_e5320_units = 0.0 if gaco_e5320_units is None else float(gaco_e5320_units)
    except Exception:
        gaco_e5320_units = 0.0
    try:
        gaco_e5320_price = 0.0 if gaco_e5320_price is None else float(gaco_e5320_price)
    except Exception:
        gaco_e5320_price = 0.0

    # SW 1-Flash logic (units & price)
    if product == "Uniflex":
        base_sw_1flash_price = SW_1FLASH_BASE_PRICE
        if roof_type in ["TPO/EPDM", "Mod Bit", "Rock/Foam/Coat"]:
            base_sw_1flash_units = math.ceil(squares / 20)  # rounded up, no decimals
        else:
            base_sw_1flash_units = math.ceil(squares / 10)  # rounded up, no decimals
    else:
        base_sw_1flash_price = 0
        base_sw_1flash_units = 0

    sw1_recalc_trigger = (
        (previous_product != product)
        or (previous_roof_type != roof_type)
        or (previous_squares != squares)
    )

    if sw1_recalc_trigger:
        sw_1flash_units = base_sw_1flash_units
        sw_1flash_price = base_sw_1flash_price
    else:
        if sw_1flash_units is None or (isinstance(sw_1flash_units, float) and math.isnan(sw_1flash_units)):
            sw_1flash_units = base_sw_1flash_units
        if sw_1flash_price is None or (isinstance(sw_1flash_price, float) and math.isnan(sw_1flash_price)):
            sw_1flash_price = base_sw_1flash_price

    # SW Bleed Block logic (units & price)
    if product == "Uniflex" and roof_type == "Mod Bit":
        base_sw_bleed_block_units = math.ceil(squares / 5)
        base_sw_bleed_block_price = SW_BLEED_BLOCK_BASE_PRICE
    else:
        base_sw_bleed_block_units = 0
        base_sw_bleed_block_price = 0

    sw_bleed_block_recalc = (
        (previous_product != product)
        or (previous_roof_type != roof_type)
        or (previous_squares != squares)
        or (previous_adjusted_coverage != adjusted_coverage)
    )

    if sw_bleed_block_recalc:
        sw_bleed_block_units = base_sw_bleed_block_units
        sw_bleed_block_price = base_sw_bleed_block_price
    else:
        if sw_bleed_block_units is None or (isinstance(sw_bleed_block_units, float) and math.isnan(sw_bleed_block_units)):
            sw_bleed_block_units = base_sw_bleed_block_units
        if sw_bleed_block_price is None or (isinstance(sw_bleed_block_price, float) and math.isnan(sw_bleed_block_price)):
            sw_bleed_block_price = base_sw_bleed_block_price

    # Drainage Mat logic (units & price)
    if roof_type in ["Ballasted 60 mil", "Ballasted 45 mil"]:
        base_drainage_units = math.ceil(squares / 18)
        base_drainage_price = DRAINAGE_MAT_BASE_PRICE
    else:
        base_drainage_units = 0
        base_drainage_price = 0

    drainage_recalc = (previous_roof_type != roof_type) or (previous_squares != squares)

    if drainage_recalc:
        drainage_mat_units = base_drainage_units
        drainage_mat_price = base_drainage_price
    else:
        if drainage_mat_units is None or (isinstance(drainage_mat_units, float) and math.isnan(drainage_mat_units)):
            drainage_mat_units = base_drainage_units
        if drainage_mat_price is None or (isinstance(drainage_mat_price, float) and math.isnan(drainage_mat_price)):
            drainage_mat_price = base_drainage_price

    # Foam logic (units & price)
    if roof_type == "Rock/Foam/Coat":
        base_foam_units = math.ceil(squares / 25)  # rounded up, no decimals
        if product == "Gaco":
            base_foam_price = GACO_FOAM_BASE_PRICE
        elif product == "Uniflex":
            base_foam_price = UNIFLEX_FOAM_BASE_PRICE
        else:
            base_foam_price = 0
    else:
        base_foam_units = 0
        base_foam_price = 0

    # Recalc foam when roof type OR product OR squares changes so base price updates correctly
    foam_recalc = (
        (previous_roof_type != roof_type)
        or (previous_product != product)
        or (previous_squares != squares)
    )

    if foam_recalc:
        foam_units = base_foam_units
        foam_price = base_foam_price
    else:
        if _is_blank_zero_or_nan(foam_units):
            foam_units = base_foam_units
        if _is_blank_zero_or_nan(foam_price):
            foam_price = base_foam_price

    # RFC labor price logic (aka rfc_price)
    base_rfc_price = RFC_LABOR_RATE if roof_type == "Rock/Foam/Coat" else 0

    rfc_recalc = (previous_roof_type != roof_type)

    if rfc_recalc:
        rfc_labor_price = base_rfc_price
    else:
        if _is_blank_zero_or_nan(rfc_labor_price):
            rfc_labor_price = base_rfc_price

    # PCS labor price logic
    base_pcs_labor_price = PCS_BASE_LABOR_RATE
    if (
        pcs_labor_price is None
        or (isinstance(pcs_labor_price, float) and math.isnan(pcs_labor_price))
        or pcs_labor_price == 0
    ):
        pcs_labor_price = base_pcs_labor_price

    # --- Enforce unit/price coupling rules ---
    def _is_blank_zero_or_nan_num(v):
        if v is None:
            return True
        try:
            if isinstance(v, float) and math.isnan(v):
                return True
            return float(v) == 0.0
        except Exception:
            return True

    def _normalize_unit_price(units_val, price_val, base_price_val):
        """Return normalized (units, price) per rule: if units<=0 -> price=0; if units>0 and price blank/0 -> base price."""
        try:
            u = 0.0 if units_val is None else float(units_val)
        except Exception:
            u = 0.0
        # No units -> zero price
        if u <= 0:
            return u, 0.0
        # Has units -> ensure price
        if _is_blank_zero_or_nan_num(price_val):
            return u, float(base_price_val or 0.0)
        try:
            return u, float(price_val)
        except Exception:
            return u, float(base_price_val or 0.0)

    # Apply to each line-item pair. Automatic quantities remain conditional on
    # product/roof type, but a user-entered quantity must always receive the
    # item's catalog price when its price is blank or zero.
    default_foam_price = (
        GACO_FOAM_BASE_PRICE
        if product == "Gaco"
        else (UNIFLEX_FOAM_BASE_PRICE if product == "Uniflex" else 0)
    )
    silicone_units_10, silicone_price = _normalize_unit_price(silicone_units_10, silicone_price, base_silicone_price)
    gaco_patch_units, gaco_patch_price = _normalize_unit_price(gaco_patch_units, gaco_patch_price, base_gaco_patch_price)
    bleed_trap_units, bleed_trap_price = _normalize_unit_price(bleed_trap_units, bleed_trap_price, BLEED_TRAP_BASE_PRICE)
    gaco_e5320_units, gaco_e5320_price = _normalize_unit_price(gaco_e5320_units, gaco_e5320_price, GACO_E5320_PRICE)
    sw_1flash_units, sw_1flash_price = _normalize_unit_price(sw_1flash_units, sw_1flash_price, SW_1FLASH_BASE_PRICE)
    sw_bleed_block_units, sw_bleed_block_price = _normalize_unit_price(sw_bleed_block_units, sw_bleed_block_price, SW_BLEED_BLOCK_BASE_PRICE)
    drainage_mat_units, drainage_mat_price = _normalize_unit_price(drainage_mat_units, drainage_mat_price, DRAINAGE_MAT_BASE_PRICE)
    foam_units, foam_price = _normalize_unit_price(foam_units, foam_price, default_foam_price)

    def _price_overridden(actual_price, base_price):
        try:
            return abs(float(actual_price or 0.0) - float(base_price or 0.0)) > 0.01
        except Exception:
            return False

    # Ensure all units and per-unit prices are whole numbers before multiplying
    silicone_total       = excel_round(silicone_units_10, 0)        * excel_round(silicone_price, 0)
    silicone_15_total    = excel_round(silicone_units_15, 0)        * excel_round(silicone_price, 0)
    silicone_20_total    = excel_round(silicone_units_20, 0)        * excel_round(silicone_price, 0)
    gaco_patch_total     = excel_round(gaco_patch_units, 0)         * excel_round(gaco_patch_price, 0)
    bleed_trap_total     = excel_round(bleed_trap_units, 0)         * excel_round(bleed_trap_price, 0)
    gaco_e5320_total     = excel_round(gaco_e5320_units, 0)         * excel_round(gaco_e5320_price, 0)
    sw_bleed_block_total = excel_round(sw_bleed_block_units, 0)     * excel_round(sw_bleed_block_price, 0)
    sw_1flash_total      = excel_round(sw_1flash_units, 0)          * excel_round(sw_1flash_price, 0)
    drainage_mat_total   = excel_round(drainage_mat_units, 0)       * excel_round(drainage_mat_price, 0)
    foam_total           = excel_round(foam_units, 0)               * excel_round(foam_price, 0)

    # Labor totals (RFC: total follows RFC price and squares: if either is 0/blank, total is 0)
    # RFC total follows RFC price and squares: if either is 0/blank, total is 0
    try:
        _rfc_price_num = float(rfc_labor_price or 0)
        _sq_num = float(squares or 0)
    except Exception:
        _rfc_price_num, _sq_num = 0.0, 0.0
    if _rfc_price_num <= 0 or _sq_num <= 0:
        rfc_labor_total = 0.0
    else:
        rfc_labor_total = _rfc_price_num * _sq_num
    pcs_labor_total = pcs_labor_price * labor_days

    calc_travel_total = 0
    include_travel_yes = (include_travel or "No").strip().lower() == "yes"
    previous_include_travel_yes = (previous_include_travel or "No").strip().lower() == "yes"
    if include_travel_yes:
        travel_misc_total = TRAVEL_MISC_250 if labor_days <= 2 else TRAVEL_MISC_500
        calc_travel_total = (
            TRAVEL_GAS_PER_JOB
            + (TRAVEL_ROOMS_PER_NIGHT * TRAVEL_HOTEL_PER_NIGHT * max(labor_days - 1, 0))
            + (TRAVEL_FOOD_PER_DAY * labor_days)
            + travel_misc_total
        )
        try:
            current_travel_total = float(travel_total or 0)
        except Exception:
            current_travel_total = 0.0
        has_manual_travel_override = (
            current_travel_total > 0
            and abs(current_travel_total - float(calc_travel_total or 0)) > 0.01
        )
        if not has_manual_travel_override:
            travel_total = calc_travel_total
    elif previous_include_travel_yes:
        travel_total = 0

    # --- Warranty total logic ---
    if product == "Gaco" and (warranty_incl or "No").strip().lower() == "yes":
        warranty_10_total = 500
        warranty_15_total = 500
        warranty_20_total = 500
    else:
        warranty_10_total = 0
        warranty_15_total = 0
        warranty_20_total = 0

    # --- Office Fee % effective value ---
    def _blank_or_nan(v):
        try:
            if v is None:
                return True
            if isinstance(v, float) and math.isnan(v):
                return True
            return float(v) == 0.0
        except Exception:
            return True

    # Re-evaluate Office Fee %
    if submitted_by != previous_submitted_by:
        office_fee_pct = office_fee_pct_for_submitter(submitted_by)
    else:
        if _blank_or_nan(office_fee_pct):
            office_fee_pct = office_fee_pct_for_submitter(submitted_by)
        else:
            office_fee_pct = float(office_fee_pct)

    effective_office_fee_pct = office_fee_pct

    # Total Price logic (moved after warranty totals are set)
    subtotal_price_10 = (
        (squares * price_per_sq_10)
        + warranty_10_total
        + (travel_total or 0)
        + (repair_costs_total or 0)
    )
    subtotal_price_15 = (
        (squares * price_per_sq_15)
        + warranty_15_total
        + (travel_total or 0)
        + (repair_costs_total or 0)
    )
    subtotal_price_20 = (
        (squares * price_per_sq_20)
        + warranty_20_total
        + (travel_total or 0)
        + (repair_costs_total or 0)
    )

    office_fee_total = excel_round(subtotal_price_10 * effective_office_fee_pct, 0)
    office_fee_15_total = excel_round(subtotal_price_15 * effective_office_fee_pct, 0)
    office_fee_20_total = excel_round(subtotal_price_20 * effective_office_fee_pct, 0)

    total_price_10 = subtotal_price_10 + office_fee_total
    total_price_15 = subtotal_price_15 + office_fee_15_total
    total_price_20 = subtotal_price_20 + office_fee_20_total

    # --- Commission percent & amount ---
    commission_pct = commission_pct_for_submitter(submitted_by)
    commission_amt = excel_round(commission_pct * (total_price_10 - foam_total - rfc_labor_total - scarifying_total - travel_total - repair_costs_total - office_fee_total), 0)
    commission_amt_15 = excel_round(commission_pct * (total_price_15 - foam_total - rfc_labor_total - scarifying_total - travel_total - repair_costs_total - office_fee_15_total), 0)
    commission_amt_20 = excel_round(commission_pct * (total_price_20 - foam_total - rfc_labor_total - scarifying_total - travel_total - repair_costs_total - office_fee_20_total), 0)

    total_cost = sum([
        silicone_total,
        gaco_patch_total,
        bleed_trap_total,
        gaco_e5320_total,
        sw_1flash_total,
        sw_bleed_block_total,
        drainage_mat_total,
        foam_total,
        rfc_labor_total,
        pcs_labor_total,
        scarifying_total,
        travel_total,
        repair_costs_total,
        warranty_10_total,
        commission_amt
    ])
    total_cost_15 = sum([
        silicone_15_total,
        gaco_patch_total,
        bleed_trap_total,
        gaco_e5320_total,
        sw_1flash_total,
        sw_bleed_block_total,
        drainage_mat_total,
        foam_total,
        rfc_labor_total,
        pcs_labor_total,
        scarifying_total,
        travel_total,
        repair_costs_total,
        warranty_15_total,
        commission_amt_15
    ])
    total_cost_20 = sum([
        silicone_20_total,
        gaco_patch_total,
        bleed_trap_total,
        gaco_e5320_total,
        sw_1flash_total,
        sw_bleed_block_total,
        drainage_mat_total,
        foam_total,
        rfc_labor_total,
        pcs_labor_total,
        scarifying_total,
        travel_total,
        repair_costs_total,
        warranty_20_total,
        commission_amt_20
    ])

    # Profit share calculation: use the configured percentage regardless of submitted_by.
    profit_share_amt = excel_round(PROFIT_SHARE_PCT * (total_price_10 - total_cost), 0)
    pcs_profit = total_price_10 - total_cost - profit_share_amt
    profit_pct = excel_round(pcs_profit / total_price_10, 2) if total_price_10 else 0
    daily_profit = excel_round(pcs_profit / labor_days, 0) if labor_days else 0
    profit_share_15 = excel_round(PROFIT_SHARE_PCT * (total_price_15 - total_cost_15), 0)
    pcs_profit_15 = total_price_15 - total_cost_15 - profit_share_15
    profit_pct_15 = excel_round(pcs_profit_15 / total_price_15, 2) if total_price_15 else 0
    daily_profit_15 = excel_round(pcs_profit_15 / labor_days, 0) if labor_days else 0
    profit_share_20 = excel_round(PROFIT_SHARE_PCT * (total_price_20 - total_cost_20), 0)
    pcs_profit_20 = total_price_20 - total_cost_20 - profit_share_20
    profit_pct_20 = excel_round(pcs_profit_20 / total_price_20, 2) if total_price_20 else 0
    daily_profit_20 = excel_round(pcs_profit_20 / labor_days, 0) if labor_days else 0

    result = {
        "labor_days": labor_days,
        "calc_labor_days": calc_labor_days,
        "ov_labor_days": ov_labor_days,
        "submitted_by": submitted_by,
        "price_per_sq_10": price_per_sq_10,
        "price_per_sq_15": price_per_sq_15,
        "price_per_sq_20": price_per_sq_20,
        "total_price_10": total_price_10,
        "total_price_15": total_price_15,
        "total_price_20": total_price_20,
        "silicone_units_10": silicone_units_10,
        "calc_silicone_units_10": calc_silicone_units_10,
        "ov_silicone_units_10": ov_silicone_units_10,
        "silicone_price": silicone_price,
        "calc_silicone_price": base_silicone_price,
        "ov_silicone_price": _price_overridden(silicone_price, base_silicone_price),
        "silicone_total": silicone_total,
        "silicone_15_total": silicone_15_total,
        "silicone_20_total": silicone_20_total,
        "gaco_patch_units": gaco_patch_units,
        "gaco_patch_price": gaco_patch_price,
        "calc_gaco_patch_price": base_gaco_patch_price,
        "ov_gaco_patch_price": _price_overridden(gaco_patch_price, base_gaco_patch_price),
        "gaco_patch_total": gaco_patch_total,
        "bleed_trap_units": bleed_trap_units,
        "bleed_trap_price": bleed_trap_price,
        "calc_bleed_trap_price": base_bleed_price,
        "ov_bleed_trap_price": _price_overridden(bleed_trap_price, base_bleed_price),
        "bleed_trap_total": bleed_trap_total,
        "gaco_e5320_units": gaco_e5320_units,
        "gaco_e5320_price": gaco_e5320_price,
        "calc_gaco_e5320_price": GACO_E5320_PRICE,
        "ov_gaco_e5320_price": _price_overridden(gaco_e5320_price, GACO_E5320_PRICE),
        "gaco_e5320_total": gaco_e5320_total,
        "sw_1flash_units": sw_1flash_units,
        "sw_1flash_price": sw_1flash_price,
        "calc_sw_1flash_price": base_sw_1flash_price,
        "ov_sw_1flash_price": _price_overridden(sw_1flash_price, base_sw_1flash_price),
        "sw_1flash_total": sw_1flash_total,
        "sw_bleed_block_units": sw_bleed_block_units,
        "sw_bleed_block_price": sw_bleed_block_price,
        "calc_sw_bleed_block_price": base_sw_bleed_block_price,
        "ov_sw_bleed_block_price": _price_overridden(sw_bleed_block_price, base_sw_bleed_block_price),
        "sw_bleed_block_total": sw_bleed_block_total,
        "drainage_mat_units": drainage_mat_units,
        "drainage_mat_price": drainage_mat_price,
        "calc_drainage_mat_price": base_drainage_price,
        "ov_drainage_mat_price": _price_overridden(drainage_mat_price, base_drainage_price),
        "drainage_mat_total": drainage_mat_total,
        "foam_units": foam_units,
        "foam_price": foam_price,
        "calc_foam_price": base_foam_price,
        "ov_foam_price": _price_overridden(foam_price, base_foam_price),
        "foam_total": foam_total,
        "rfc_labor_price": rfc_labor_price,
        "pcs_labor_price": pcs_labor_price,
        "calc_rfc_labor_price": base_rfc_price,
        "ov_rfc_labor_price": _price_overridden(rfc_labor_price, base_rfc_price),
        "calc_pcs_labor_price": base_pcs_labor_price,
        "ov_pcs_labor_price": _price_overridden(pcs_labor_price, base_pcs_labor_price),
        "rfc_labor_total": rfc_labor_total,
        "pcs_labor_total": pcs_labor_total,
        "scarifying_total": scarifying_total,
        "travel_total": travel_total,
        "calc_travel_total": calc_travel_total,
        "repair_costs_total": repair_costs_total,
        "office_fee_total": office_fee_total,
        "office_fee_15_total": office_fee_15_total,
        "office_fee_20_total": office_fee_20_total,
        "pcs_profit": pcs_profit,
        "pcs_profit_15": pcs_profit_15,
        "pcs_profit_20": pcs_profit_20,
        "profit_pct": profit_pct,
        "profit_pct_15": profit_pct_15,
        "profit_pct_20": profit_pct_20,
        "daily_profit": daily_profit,
        "daily_profit_15": daily_profit_15,
        "daily_profit_20": daily_profit_20,
        "profit_share": profit_share_amt,
        "profit_share_15": profit_share_15,
        "profit_share_20": profit_share_20,
        "warranty_10_total": warranty_10_total,
        "warranty_15_total": warranty_15_total,
        "warranty_20_total": warranty_20_total,
        "coverage_10": coverage_10,
        "coverage_15": coverage_15,
        "coverage_20": coverage_20,
        "silicone_units_15": silicone_units_15,
        "silicone_units_20": silicone_units_20,
        "commission_amt": commission_amt,
        "commission_amt_15": commission_amt_15,
        "commission_amt_20": commission_amt_20,
        "commission_pct": commission_pct,
        "total_cost": total_cost,
        "total_cost_15": total_cost_15,
        "total_cost_20": total_cost_20,
        "warranty_incl": warranty_incl,
        "include_travel": include_travel,
        "previous_include_travel": include_travel,
        "previous_calc_travel_total": calc_travel_total,
        "office_fee_pct": effective_office_fee_pct,
        "adjusted_coverage": adjusted_coverage,
        "previous_submitted_by": submitted_by,
        "previous_roof_type": roof_type,
        "previous_squares": squares,
        "previous_product": product,
        "previous_adjusted_coverage": adjusted_coverage,
        "previous_silicone_units_10": silicone_units_10,
        "proposal_note": proposal_note,
    }
    return result


@app.route('/')
def landing_page():
    return render_template('landing.html')


CONTACT_ORGANIZATION_TYPES = (
    "Property Management",
    "Roofing Company",
    "Roofing Contractor",
    "Property Owner",
    "General Contractor",
    "Real Estate",
    "Distributor",
    "Manufacturer",
    "Consultant",
    "Vendor",
    "Unknown",
    "Other",
)


def _contact_form_values():
    values = {
        "first_name": request.form.get("first_name", "").strip(),
        "last_name": request.form.get("last_name", "").strip(),
        "business_email": request.form.get("business_email", "").strip(),
        "business_phone": request.form.get("business_phone", "").strip(),
        "mobile_phone": request.form.get("mobile_phone", "").strip(),
        "main_office_address_line_1": request.form.get("main_office_address_line_1", "").strip(),
        "main_office_address_line_2": request.form.get("main_office_address_line_2", "").strip(),
        "main_office_city": request.form.get("main_office_city", "").strip(),
        "main_office_state": request.form.get("main_office_state", "").strip(),
        "main_office_zip_code": request.form.get("main_office_zip_code", "").strip(),
        "branch_address_line_1": request.form.get("branch_address_line_1", "").strip(),
        "branch_address_line_2": request.form.get("branch_address_line_2", "").strip(),
        "branch_city": request.form.get("branch_city", "").strip(),
        "branch_state": request.form.get("branch_state", "").strip(),
        "branch_zip_code": request.form.get("branch_zip_code", "").strip(),
        "title": request.form.get("title", "").strip(),
        "linkedin_url": request.form.get("linkedin_url", "").strip(),
        "contact_notes": request.form.get("contact_notes", "").strip(),
        "relationship_notes": request.form.get("relationship_notes", "").strip(),
        "do_not_contact": request.form.get("do_not_contact") == "on",
        "organization_id": request.form.get("organization_id", "").strip(),
        "organization_name": request.form.get("organization_name", "").strip(),
        "organization_type": request.form.get("organization_type", "Other").strip(),
    }
    if not values["first_name"] and not values["last_name"]:
        raise ValueError("Enter a first name or last name.")
    email = values["business_email"]
    if email and (parseaddr(email)[1] != email or "@" not in email):
        raise ValueError("Enter a valid email address.")
    for field, label in (
        ("main_office_state", "Main office state"),
        ("branch_state", "Branch state"),
    ):
        if values[field] and len(values[field]) != 2:
            raise ValueError(f"{label} must use a two-letter abbreviation.")
    return values


def _resolve_contact_organization(store, values):
    organization_id = values.get("organization_id", "").strip()
    organization_name = values.get("organization_name", "").strip()
    organization_type = values.get("organization_type", "Other").strip()
    if organization_id:
        if not organization_name:
            raise ValueError("Enter an organization name.")
        if organization_type not in CONTACT_ORGANIZATION_TYPES:
            raise ValueError("Select a valid organization type.")
        store.update_organization(organization_id, organization_name, organization_type, values)
    elif organization_name:
        if organization_type not in CONTACT_ORGANIZATION_TYPES:
            raise ValueError("Select a valid organization type.")
        existing = store.find_organization_by_name(organization_name)
        if existing:
            values["organization_id"] = existing["id"]
            store.update_organization(
                existing["id"], organization_name, organization_type, values
            )
        else:
            values["organization_id"] = store.create_organization(
                organization_name, organization_type, values
            )
    else:
        values["organization_id"] = store.resolve_organization_from_email(
            values.get("business_email", "")
        )


def _contact_assignment_context(source):
    proposal_id = str(source.get("attach_to_proposal", "") or "").strip()
    proposal_name = " ".join(str(source.get("proposal_name", "") or "").split())
    if not proposal_id:
        return "", proposal_name
    try:
        proposal_id = str(uuid.UUID(proposal_id))
    except ValueError as exc:
        raise ValueError("That proposal could not be selected.") from exc
    return proposal_id, proposal_name


def _attach_contact_record(proposal_id, contact_record):
    relationship_id = str((contact_record or {}).get("id") or "").strip()
    if not relationship_id:
        raise ContactStoreError("The contact's organization relationship could not be found.")
    return get_proposal_tracking_store().assign_or_create_primary_contact(
        proposal_id,
        organization_contact_id=relationship_id,
    )


_CONTACT_DETAIL_RETURN_SESSION_KEY = "proposal_contact_detail_return"


def _remember_contact_detail_return(proposal_id, source):
    if str(source.get("return_to_detail", "")).strip() != "1":
        session.pop(_CONTACT_DETAIL_RETURN_SESSION_KEY, None)
        return
    folder_name = os.path.basename(
        str(source.get("proposal_folder_name", "") or "").strip()
    )
    session[_CONTACT_DETAIL_RETURN_SESSION_KEY] = {
        "proposal_id": str(proposal_id),
        "folder_name": folder_name,
        "customer_was_blank": (
            str(source.get("customer_was_blank", "")).strip() == "1"
        ),
    }


def _contact_assignment_success_redirect(proposal_id, contact_result):
    context = session.get(_CONTACT_DETAIL_RETURN_SESSION_KEY) or {}
    if str(context.get("proposal_id") or "") != str(proposal_id):
        return redirect(url_for("proposal_list"))

    customer_was_blank = bool(context.get("customer_was_blank"))
    organization_name = " ".join(
        str((contact_result or {}).get("organization") or "").split()
    )
    if customer_was_blank and organization_name:
        get_proposal_tracking_store().update_proposal_customer_name(
            str(proposal_id),
            organization_name,
        )

    session.pop(_CONTACT_DETAIL_RETURN_SESSION_KEY, None)
    folder_name = str(context.get("folder_name") or "").strip() or "__blank__"
    return redirect(url_for(
        "proposal_details_query",
        folder_name=folder_name,
        proposal_id=str(proposal_id),
        read_only="No",
        customer_was_blank=(
            "1" if customer_was_blank and not organization_name else None
        ),
    ))


@app.get('/contacts')
def contact_management():
    search = request.args.get("q", "").strip()
    status = request.args.get("status", "active").strip().lower()
    if status not in {"active", "archived", "all"}:
        status = "active"
    edit_id = request.args.get("edit", "").strip()
    contacts = []
    organizations = []
    selected_contact = None
    configuration_error = ""
    attach_to_proposal = ""
    proposal_name = ""
    try:
        attach_to_proposal, proposal_name = _contact_assignment_context(request.args)
        if attach_to_proposal:
            _remember_contact_detail_return(attach_to_proposal, request.args)
    except ValueError as exc:
        flash(str(exc), "danger")
    try:
        store = get_contact_store()
        contacts = store.list_contacts(search=search, status=status)
        organizations = store.list_organizations()
        if edit_id:
            selected_contact = store.get_contact(edit_id)
            if selected_contact is None:
                flash("That contact could not be found.", "danger")
    except (ContactConfigurationError, ContactStoreError) as exc:
        configuration_error = str(exc)
    return render_template(
        'contact_management.html',
        contacts=contacts,
        organizations=organizations,
        selected_contact=selected_contact,
        search=search,
        status=status,
        configuration_error=configuration_error,
        organization_types=CONTACT_ORGANIZATION_TYPES,
        attach_to_proposal=attach_to_proposal,
        proposal_name=proposal_name,
    )


@app.post('/contacts')
def create_contact():
    attach_to_proposal = ""
    proposal_name = ""
    try:
        attach_to_proposal, proposal_name = _contact_assignment_context(request.form)
        store = get_contact_store()
        values = _contact_form_values()
        action = request.form.get("duplicate_action", "").strip().lower()
        if action not in {"", "keep", "replace"}:
            raise ValueError("Select a valid duplicate-contact action.")
        if (
            values["organization_name"]
            and values["organization_type"] not in CONTACT_ORGANIZATION_TYPES
        ):
            raise ValueError("Select a valid organization type.")

        duplicates = store.find_duplicate_contacts(values)
        if duplicates and not action:
            return render_template(
                "contact_management.html",
                contacts=store.list_contacts(status="active"),
                organizations=store.list_organizations(),
                selected_contact=None,
                search="",
                status="active",
                configuration_error="",
                organization_types=CONTACT_ORGANIZATION_TYPES,
                duplicate_matches=duplicates,
                pending_values=values,
                attach_to_proposal=attach_to_proposal,
                proposal_name=proposal_name,
            )

        if action == "replace":
            duplicate_contact_id = request.form.get("duplicate_contact_id", "").strip()
            valid_duplicate_ids = {
                (row.get("contact") or {}).get("id")
                for row in duplicates
            }
            if duplicate_contact_id not in valid_duplicate_ids:
                raise ValueError("Select an existing duplicate contact to replace.")
            _resolve_contact_organization(store, values)
            store.update_contact(duplicate_contact_id, values)
            if attach_to_proposal:
                result = _attach_contact_record(
                    attach_to_proposal,
                    store.get_contact(duplicate_contact_id),
                )
                flash(f"{result.get('name') or 'Contact'} updated and attached.", "success")
                return _contact_assignment_success_redirect(
                    attach_to_proposal, result
                )
            flash("Existing contact replaced with the submitted information.", "success")
            return redirect(url_for("contact_management", edit=duplicate_contact_id))

        _resolve_contact_organization(store, values)
        contact_id = store.create_contact(values)
        if attach_to_proposal:
            result = _attach_contact_record(
                attach_to_proposal,
                store.get_contact(contact_id),
            )
            flash(f"{result.get('name') or 'Contact'} added and attached.", "success")
            return _contact_assignment_success_redirect(attach_to_proposal, result)
    except (ValueError, ContactStoreError, TenantAuthenticationError) as exc:
        flash(str(exc), "danger")
        if attach_to_proposal:
            return redirect(url_for(
                "contact_management",
                attach_to_proposal=attach_to_proposal,
                proposal_name=proposal_name,
            ))
    else:
        flash("Contact added." if action != "keep" else "Contact kept as a separate record.", "success")
    return redirect(url_for("contact_management"))


@app.post('/proposals/<uuid:proposal_id>/contacts/<uuid:organization_contact_id>/attach')
def attach_proposal_contact(proposal_id, organization_contact_id):
    proposal_name = " ".join(request.form.get("proposal_name", "").split())
    try:
        result = get_proposal_tracking_store().assign_or_create_primary_contact(
            str(proposal_id),
            organization_contact_id=str(organization_contact_id),
        )
    except (ContactStoreError, TenantAuthenticationError, ValueError) as exc:
        flash(str(exc), "danger")
        return redirect(url_for(
            "contact_management",
            attach_to_proposal=str(proposal_id),
            proposal_name=proposal_name,
        ))
    flash(f"{result.get('name') or 'Contact'} attached to {proposal_name or 'the proposal'}.", "success")
    return _contact_assignment_success_redirect(str(proposal_id), result)


@app.post('/contacts/<uuid:contact_id>/edit')
def edit_contact(contact_id):
    attach_to_proposal = ""
    proposal_name = ""
    try:
        attach_to_proposal, proposal_name = _contact_assignment_context(request.form)
        store = get_contact_store()
        values = _contact_form_values()
        _resolve_contact_organization(store, values)
        store.update_contact(str(contact_id), values)
        if attach_to_proposal:
            result = _attach_contact_record(
                attach_to_proposal,
                store.get_contact(str(contact_id)),
            )
            flash(f"{result.get('name') or 'Contact'} updated and attached.", "success")
            return _contact_assignment_success_redirect(attach_to_proposal, result)
    except (ValueError, ContactStoreError, TenantAuthenticationError) as exc:
        flash(str(exc), "danger")
        return redirect(url_for(
            "contact_management",
            edit=str(contact_id),
            attach_to_proposal=attach_to_proposal or None,
            proposal_name=proposal_name or None,
        ))
    flash("Contact updated.", "success")
    return redirect(url_for("contact_management"))


@app.post('/contacts/<uuid:contact_id>/delete')
def delete_contact(contact_id):
    try:
        get_contact_store().archive_contact(str(contact_id))
    except ContactStoreError as exc:
        flash(str(exc), "danger")
    else:
        flash("Contact removed from the active list. Its history was preserved.", "success")
    return redirect(url_for("contact_management"))


@app.route('/proposals')
def proposal_list():
    requested_filter = (
        request.args.get('filter')
        or request.args.get('status')
        or 'all'
    ).strip().lower()
    filter_aliases = {
        'open': 'all',
        'under': 'under_contract',
        'under-contract': 'under_contract',
        'contract': 'under_contract',
        'draft_unsent': 'draft',
        'draft-unsent': 'draft',
        'unsent': 'draft',
        'not_sent': 'draft',
        'not-sent': 'draft',
    }
    selected_filter = filter_aliases.get(requested_filter, requested_filter)
    filter_statuses = {
        'all': {'draft', 'sent', 'under_contract', 'finished', 'dead'},
        'draft': {'draft'},
        'sent': {'sent'},
        'under_contract': {'under_contract'},
        'finished': {'finished'},
        'dead': {'dead'},
    }
    if selected_filter not in filter_statuses:
        selected_filter = 'all'

    recent_cutoff = datetime.datetime.now() - datetime.timedelta(days=7)
    try:
        store = get_proposal_tracking_store()
        proposals = store.list_management_proposals(
            filter_statuses[selected_filter]
        )
    except (ContactStoreError, TenantAuthenticationError) as exc:
        flash(str(exc), "danger")
        proposals = []

    proposals = [
        proposal for proposal in proposals
        if proposal.get("status") in filter_statuses[selected_filter]
    ]
    for proposal in proposals:
        last_modified = proposal.get("last_modified")
        proposal["is_recent"] = bool(
            last_modified and last_modified >= recent_cutoff
        )

    return render_template(
        'proposal_list.html',
        proposal_list=proposals,
        selected_filter=selected_filter,
    )


@app.post('/api/proposals/<uuid:proposal_id>/primary-contact')
def update_proposal_primary_contact(proposal_id):
    payload = request.get_json(silent=True)
    if not isinstance(payload, dict):
        return jsonify({"error": "Enter contact information and try again."}), 400
    try:
        result = get_proposal_tracking_store().assign_or_create_primary_contact(
            str(proposal_id),
            organization_contact_id=payload.get("organization_contact_id", ""),
            contact_name=payload.get("contact_name", ""),
            email=payload.get("email", ""),
            organization_name=payload.get("organization_name", ""),
        )
    except ProposalContactOrganizationRequired as exc:
        return jsonify({
            "error": str(exc),
            "organization_required": True,
            "domain": exc.domain,
        }), 409
    except (ContactStoreError, TenantAuthenticationError, ValueError) as exc:
        return jsonify({"error": str(exc)}), 400
    return jsonify({"contact": result})


@app.route('/blast-emails')
def blast_email_management():
    return render_template('blast_email_management.html')


def _roof_local_worker_enabled():
    value = os.environ.get("ROOF_INTELLIGENCE_LOCAL_WORKER", "0").strip().lower()
    return value in {"1", "true", "yes", "on"}


def _roof_report_editing_enabled():
    return load_cutover_flags().editing_enabled


def _roof_worker_readiness_error():
    if not os.path.isdir(ROOF_INTELLIGENCE_PROJECT_DIR):
        return "The PilotPoint IQ project folder is not available on this Mac."
    python_path = os.path.join(ROOF_INTELLIGENCE_PROJECT_DIR, ".venv", "bin", "python")
    if not os.path.isfile(python_path):
        return "The PilotPoint IQ Python environment is not installed or is incomplete."
    if not os.path.isfile(ROOF_INTELLIGENCE_SCRIPT):
        return "The PCS Roof Intelligence adapter is missing from the desktop app."
    if not os.path.isfile(ROOF_INTELLIGENCE_AREA_SCRIPT):
        return "The PCS area-batch adapter is missing from the desktop app."
    return None


def _roof_error_details(message):
    clean_message = " ".join(str(message or "Unable to complete the Roof Intelligence report.").split())
    lowered = clean_message.lower()
    if "footprint discrepancy needs attention" in lowered:
        return "footprint_discrepancy", clean_message[:500], False
    if "no supported county parcel match" in lowered:
        return "unsupported_county", clean_message[:500], False
    if "no parcel match" in lowered or "parcel" in lowered and "not found" in lowered:
        return "parcel_not_found", "No matching county parcel was found for that address.", False
    if "no building" in lowered or "building footprint" in lowered:
        return "building_not_found", "A building footprint could not be matched to the property parcel.", False
    if "coordinate systems" in lowered or "county parcel discovery failed" in lowered:
        return (
            "gis_service_unavailable",
            "County GIS services were temporarily unavailable while locating properties. Please retry the report.",
            True,
        )
    if "aerial" in lowered or "imagery" in lowered:
        return "imagery_unavailable", "A usable aerial image could not be retrieved for the property.", True
    if "timed out" in lowered or "timeout" in lowered:
        return "worker_timeout", "Report processing exceeded the local worker time limit.", True
    if "openai" in lowered or "gemini" in lowered or "ai " in lowered:
        return "ai_analysis_failed", "The AI roof assessment could not be completed.", True
    return "internal_processing_error", clean_message[:500], True


def _footprint_error_context(payload, fallback_address=""):
    message = str(payload.get("error") or "")
    match = re.search(
        r"for\s+(.+?)\s+parcel\s+([A-Za-z0-9-]+):.*?footprint\s+([\d,.]+)\s+sq\s*ft\s+versus\s+county\s+GIS\s+footprint\s+([\d,.]+)\s+sq\s*ft\s+\(([\d.]+)%",
        message,
        re.IGNORECASE,
    )
    context = {
        "error_code": "footprint_discrepancy",
        "address": fallback_address,
        "error": message,
    }
    if match:
        context.update(
            {
                "county": match.group(1),
                "parcel": match.group(2),
                "footprint_validation": {
                    "status": "discrepancy",
                    "primary_sqft": float(match.group(3).replace(",", "")),
                    "secondary_sqft": float(match.group(4).replace(",", "")),
                    "difference_pct": float(match.group(5)),
                    "primary_label": "Supabase Microsoft",
                    "secondary_label": "County GIS",
                },
            }
        )
    else:
        assessor_match = re.search(
            r"for\s+(.+?)\s+parcel\s+([A-Za-z0-9-]+):\s*selected\s+footprint\s+([\d,.]+)\s+sq\s*ft\s+versus\s+explicit\s+county\s+assessor\s+footprint\s+([\d,.]+)\s+sq\s*ft\s+\(([\d.]+)%",
            message,
            re.IGNORECASE,
        )
        if assessor_match:
            context.update(
                {
                    "county": assessor_match.group(1),
                    "parcel": assessor_match.group(2),
                    "footprint_validation": {
                        "status": "discrepancy",
                        "primary_sqft": float(assessor_match.group(3).replace(",", "")),
                        "secondary_sqft": float(assessor_match.group(4).replace(",", "")),
                        "difference_pct": float(assessor_match.group(5)),
                        "primary_label": "Selected building",
                        "secondary_label": "County assessor",
                    },
                }
            )
    canonical_match = re.search(r"Canonical footprint\s+(\d+)\s+is pending review", message, re.IGNORECASE)
    if canonical_match:
        context["canonical_id"] = int(canonical_match.group(1))
    return context


def _run_local_individual_roof_job(job_id, user_key=None):
    """Process an individual job that was atomically claimed by the local worker."""
    store = get_job_store()
    claimed = store.get_job(job_id)
    trusted_user_key = user_key or (claimed or {}).get("user_key")
    job = store.get_job(job_id, trusted_user_key) if trusted_user_key else None
    if not job or job.get("job_type") != "individual_address" or job.get("status") != "running":
        return

    try:
        readiness_error = _roof_worker_readiness_error()
        if readiness_error:
            raise RuntimeError(readiness_error)
        python_path = os.path.join(ROOF_INTELLIGENCE_PROJECT_DIR, ".venv", "bin", "python")

        store.update_job(job_id, stage="processing_report")
        command = [
            python_path,
            ROOF_INTELLIGENCE_SCRIPT,
            "--address",
            job["input"]["property_address"],
            "--project-dir",
            ROOF_INTELLIGENCE_PROJECT_DIR,
            "--county",
            "auto",
            "--use-ai",
            "--allow-ai-fallback",
        ]
        report_output_dir, image_output_dir = _tenant_report_output_paths(trusted_user_key)
        command.extend(["--output-dir", report_output_dir, "--image-dir", image_output_dir])
        override = job["input"].get("footprint_override") or {}
        if override.get("selected_source"):
            command.extend([
                "--footprint-source", str(override["selected_source"]),
                "--footprint-override-reason", str(override.get("reason") or ""),
            ])
        if _roof_report_editing_enabled():
            area_override = store.get_active_square_footage_override(
                address=job["input"]["property_address"],
                user_key=trusted_user_key,
            )
            if area_override:
                command.extend(["--roof-area-override", str(area_override["numeric_value"])])
        completed = subprocess.run(
            command,
            cwd=ROOF_INTELLIGENCE_PROJECT_DIR,
            capture_output=True,
            text=True,
            timeout=int(os.environ.get("ROOF_INTELLIGENCE_LOCAL_TIMEOUT", "900")),
            check=False,
        )
        payload = {}
        for line in reversed((completed.stdout or "").splitlines()):
            line = line.strip()
            if line.startswith("{") and line.endswith("}"):
                try:
                    payload = json.loads(line)
                except json.JSONDecodeError:
                    payload = {}
                break
        if completed.returncode != 0 or payload.get("error"):
            error = payload.get("error") or (completed.stderr or completed.stdout or "Unable to generate the report.").strip()
            code, message, retryable = _roof_error_details(error)
            details = _footprint_error_context(payload, job["input"].get("property_address", "")) if code == "footprint_discrepancy" else payload
            store.fail_job(job_id, code, message, retryable=retryable, error_details=details)
            return
        if not payload.get("report_path") or not os.path.isfile(payload["report_path"]):
            raise RuntimeError("The worker completed without producing a report PDF.")

        store.update_job(job_id, stage="saving_report")
        store.complete_individual_job(job_id, payload)

        temporary_image = str(payload.get("aerial_image_file") or "")
        if temporary_image and os.path.isfile(temporary_image):
            try:
                os.remove(temporary_image)
            except OSError:
                pass
    except subprocess.TimeoutExpired:
        store.fail_job(
            job_id,
            "worker_timeout",
            "Report processing exceeded the local worker time limit.",
            retryable=True,
        )
    except Exception as exc:
        code, message, retryable = _roof_error_details(exc)
        store.fail_job(job_id, code, message, retryable=retryable)


def _last_json_payload(completed):
    payload = {}
    for line in reversed((completed.stdout or "").splitlines()):
        line = line.strip()
        if line.startswith("{") and line.endswith("}"):
            try:
                return json.loads(line)
            except json.JSONDecodeError:
                continue
    return payload


def _canonical_footprint_reviews(limit=20):
    if not _roof_local_worker_enabled():
        return []
    command = [
        os.path.join(ROOF_INTELLIGENCE_PROJECT_DIR, ".venv", "bin", "python"),
        os.path.join(ROOF_INTELLIGENCE_PROJECT_DIR, "scripts", "review_canonical_footprints.py"),
        "--limit", str(limit),
    ]
    try:
        completed = subprocess.run(
            command, cwd=ROOF_INTELLIGENCE_PROJECT_DIR, capture_output=True,
            text=True, timeout=30, check=False,
        )
        if completed.returncode == 0:
            payload = json.loads(completed.stdout or "[]")
            return payload if isinstance(payload, list) else []
    except Exception:
        pass
    return []


def _resolve_canonical_footprint(canonical_id, selected_source, reason):
    source = "microsoft" if selected_source == "supabase" else selected_source
    command = [
        os.path.join(ROOF_INTELLIGENCE_PROJECT_DIR, ".venv", "bin", "python"),
        os.path.join(ROOF_INTELLIGENCE_PROJECT_DIR, "scripts", "review_canonical_footprints.py"),
        "--resolve", str(int(canonical_id)), "--source", source,
        "--reason", str(reason), "--reviewer", _roof_intelligence_user_key(),
    ]
    completed = subprocess.run(
        command, cwd=ROOF_INTELLIGENCE_PROJECT_DIR, capture_output=True,
        text=True, timeout=30, check=False,
    )
    if completed.returncode != 0:
        raise RuntimeError((completed.stderr or completed.stdout or "Unable to resolve canonical footprint.").strip())
    return json.loads(completed.stdout)


class RoofWorkerFailure(RuntimeError):
    def __init__(self, payload):
        self.payload = payload or {}
        super().__init__(self.payload.get("error") or "Unable to generate the report.")


def _discover_local_area_candidates(job):
    job_input = job.get("input", {})
    bounds = job_input.get("bounds") or {}
    python_path = os.path.join(ROOF_INTELLIGENCE_PROJECT_DIR, ".venv", "bin", "python")
    command = [
        python_path,
        ROOF_INTELLIGENCE_AREA_SCRIPT,
        "--project-dir", ROOF_INTELLIGENCE_PROJECT_DIR,
        "--north", str(bounds.get("north", "")),
        "--south", str(bounds.get("south", "")),
        "--east", str(bounds.get("east", "")),
        "--west", str(bounds.get("west", "")),
        "--minimum-roof-size", str(job.get("minimum_roof_size") or 1),
        "--max-candidates", str(int(os.environ.get("ROOF_INTELLIGENCE_AREA_MAX_CANDIDATES", "2000"))),
    ]
    selection_type = job_input.get("selection_type") or "rectangle"
    command.extend(["--selection-type", selection_type])
    if selection_type == "radius":
        center = job_input.get("center") or {}
        command.extend([
            "--center-lat", str(center.get("lat", "")),
            "--center-lng", str(center.get("lng", "")),
            "--radius-miles", str(job_input.get("radius_miles", "")),
        ])
    attempts = max(1, int(os.environ.get("ROOF_INTELLIGENCE_AREA_DISCOVERY_ATTEMPTS", "2")))
    for attempt in range(1, attempts + 1):
        completed = subprocess.run(
            command,
            cwd=ROOF_INTELLIGENCE_PROJECT_DIR,
            capture_output=True,
            text=True,
            timeout=int(os.environ.get("ROOF_INTELLIGENCE_AREA_DISCOVERY_TIMEOUT", "900")),
            check=False,
        )
        payload = _last_json_payload(completed)
        if completed.returncode == 0 and not payload.get("error"):
            return list(payload.get("candidates") or []), list(payload.get("warnings") or [])

        error = str(
            payload.get("error")
            or (completed.stderr or completed.stdout or "Unable to discover properties in the selected area.").strip()
        )
        transient_crs_failure = (
            "coordinate systems" in error.lower()
            or "county parcel discovery failed" in error.lower()
        )
        if not transient_crs_failure or attempt >= attempts:
            raise RuntimeError(error)
        _safe_debug(f"[ROOF AREA] GIS discovery attempt {attempt} failed; retrying.")
        time.sleep(float(os.environ.get("ROOF_INTELLIGENCE_AREA_RETRY_DELAY", "2")))

    raise RuntimeError("Unable to discover properties in the selected area.")


def _run_local_candidate_report(candidate):
    user_key = str(candidate.get("_tenant_user_key") or "local-user")
    python_path = os.path.join(ROOF_INTELLIGENCE_PROJECT_DIR, ".venv", "bin", "python")
    command = [
        python_path,
        ROOF_INTELLIGENCE_SCRIPT,
        "--address", candidate["address"],
        "--project-dir", ROOF_INTELLIGENCE_PROJECT_DIR,
        "--county", candidate.get("county_profile") or "auto",
        "--allow-pending-footprint-review",
        "--use-ai",
        "--allow-ai-fallback",
    ]
    report_output_dir, image_output_dir = _tenant_report_output_paths(user_key)
    command.extend(["--output-dir", report_output_dir, "--image-dir", image_output_dir])
    if candidate.get("parcel"):
        command.extend(["--parcel-id", str(candidate["parcel"])])
    if _roof_report_editing_enabled():
        area_override = get_job_store().get_active_square_footage_override(
            address=candidate.get("address"),
            county=candidate.get("county"),
            parcel_number=candidate.get("parcel"),
            user_key=user_key,
        )
        if area_override:
            command.extend(["--roof-area-override", str(area_override["numeric_value"])])
    override = candidate.get("footprint_override") or {}
    if override.get("selected_source"):
        command.extend([
            "--footprint-source", str(override["selected_source"]),
            "--footprint-override-reason", str(override.get("reason") or ""),
        ])
    completed = subprocess.run(
        command,
        cwd=ROOF_INTELLIGENCE_PROJECT_DIR,
        capture_output=True,
        text=True,
        timeout=int(os.environ.get("ROOF_INTELLIGENCE_LOCAL_TIMEOUT", "900")),
        check=False,
    )
    payload = _last_json_payload(completed)
    if completed.returncode != 0 or payload.get("error"):
        if not payload:
            payload = {"error": (completed.stderr or completed.stdout or "Unable to generate the report.").strip()}
        raise RoofWorkerFailure(payload)
    if not payload.get("report_path") or not os.path.isfile(payload["report_path"]):
        raise RuntimeError("The worker completed without producing a report PDF.")
    return payload


def _roof_type_matches_selection(result, selected_types):
    if set(selected_types or []) == set(SUPPORTED_ROOF_TYPES):
        return True
    value = str(result.get("roof_type") or "").strip().lower()
    aliases = {
        "TPO": ("tpo",),
        "PVC": ("pvc",),
        "EPDM": ("epdm",),
        "Modified Bitumen": ("modified bitumen", "mod bit"),
        "Ballasted": ("ballasted",),
        "Tar and Gravel": ("tar and gravel", "built-up", "bur"),
        "Metal": ("metal",),
    }
    return any(
        term in value
        for roof_type in selected_types or []
        for term in aliases.get(roof_type, (roof_type.lower(),))
    )


def _remove_area_temporary_files(result, *, remove_report=False):
    paths = [str(result.get("aerial_image_file") or "")]
    if remove_report:
        paths.append(str(result.get("report_path") or ""))
    for path in paths:
        if path and os.path.isfile(path):
            try:
                os.remove(path)
            except OSError:
                pass


def _run_local_area_roof_job(job_id, user_key=None):
    store = get_job_store()
    claimed = store.get_job(job_id)
    trusted_user_key = user_key or (claimed or {}).get("user_key")
    job = store.get_job(job_id, trusted_user_key) if trusted_user_key else None
    if not job or job.get("job_type") != "zip_batch" or job.get("status") != "running":
        return
    try:
        readiness_error = _roof_worker_readiness_error()
        if readiness_error:
            raise RuntimeError(readiness_error)
        items = store.list_area_items(job_id)
        if not items:
            store.update_job(job_id, stage="discovering_properties")
            candidates, warnings = _discover_local_area_candidates(job)
            if warnings:
                _safe_debug("[ROOF AREA] " + " | ".join(warnings))
            items = store.prepare_area_candidates(job_id, candidates)

        while True:
            current = store.get_job(job_id, trusted_user_key)
            if not current or current.get("status") == "cancelled":
                return
            item = store.claim_next_area_item(job_id)
            if not item:
                break
            try:
                candidate_input = dict(item["input"])
                candidate_input["_tenant_user_key"] = trusted_user_key
                result = _run_local_candidate_report(candidate_input)
                if not _roof_type_matches_selection(result, current.get("roof_types") or []):
                    _remove_area_temporary_files(result, remove_report=True)
                    store.skip_area_item(
                        job_id,
                        item["id"],
                        "roof_type_excluded",
                        f"Detected roof type '{result.get('roof_type') or 'Unknown'}' was not selected.",
                    )
                    continue
                store.complete_area_item(job_id, item["id"], result)
                _remove_area_temporary_files(result)
            except subprocess.TimeoutExpired:
                store.fail_area_item(job_id, item["id"], "worker_timeout", "Report processing exceeded the local worker time limit.")
            except RoofWorkerFailure as exc:
                code, message, _ = _roof_error_details(exc)
                details = _footprint_error_context(exc.payload, item["input"].get("address", "")) if code == "footprint_discrepancy" else exc.payload
                store.fail_area_item(job_id, item["id"], code, message, error_details=details)
            except Exception as exc:
                code, message, _ = _roof_error_details(exc)
                store.fail_area_item(job_id, item["id"], code, message)

        store.finish_area_job(job_id)
    except subprocess.TimeoutExpired:
        store.fail_job(
            job_id,
            "area_discovery_timeout",
            "Property discovery exceeded the local worker time limit.",
            retryable=True,
        )
    except Exception as exc:
        code, message, retryable = _roof_error_details(exc)
        store.fail_job(job_id, code, message, retryable=retryable, stage="area_batch_failed")


_roof_worker_wake = threading.Event()
_roof_worker_lock = threading.Lock()
_roof_worker_thread = None
_county_health_lock = threading.Lock()
_county_health_process = None


def _county_health_check_running():
    global _county_health_process
    with _county_health_lock:
        if _county_health_process is None:
            return False
        if _county_health_process.poll() is None:
            return True
        _county_health_process = None
        return False


def _start_manual_county_health_check():
    global _county_health_process
    if not _roof_local_worker_enabled():
        raise RuntimeError("Manual county health checks require the local PilotPoint worker.")
    readiness_error = _roof_worker_readiness_error()
    if readiness_error:
        raise RuntimeError(readiness_error)

    python_path = os.path.join(ROOF_INTELLIGENCE_PROJECT_DIR, ".venv", "bin", "python")
    script_path = os.path.join(ROOF_INTELLIGENCE_PROJECT_DIR, "county_discovery_health.py")
    if not os.path.isfile(script_path):
        raise RuntimeError("The PilotPoint county health-check service is not available.")

    output_path = os.path.join(
        ROOF_INTELLIGENCE_PROJECT_DIR,
        "data",
        "health",
        "county-discovery-latest.json",
    )
    command = [
        python_path,
        script_path,
        "--strict-discrepancies",
        "--all-samples",
        "--notify-pcs",
        "--output",
        output_path,
    ]
    os.makedirs(DEFAULT_DATA_DIR, exist_ok=True)
    log_path = os.path.join(str(DEFAULT_DATA_DIR), "county-health-manual.log")

    with _county_health_lock:
        if _county_health_process is not None and _county_health_process.poll() is None:
            return False
        with open(log_path, "a", encoding="utf-8") as log_handle:
            _county_health_process = subprocess.Popen(
                command,
                cwd=ROOF_INTELLIGENCE_PROJECT_DIR,
                stdin=subprocess.DEVNULL,
                stdout=log_handle,
                stderr=subprocess.STDOUT,
                env=os.environ.copy(),
                start_new_session=True,
            )
    return True


def _roof_worker_loop():
    store = get_job_store()
    store.recover_interrupted_individual_jobs()
    store.recover_interrupted_area_jobs()
    while True:
        job = store.claim_next_individual_job()
        if job:
            DESKTOP_BACKGROUND_WORK_ACTIVE.set()
            _run_local_individual_roof_job(job["id"], job.get("user_key"))
            continue

        job = store.claim_next_area_job()
        if job:
            DESKTOP_BACKGROUND_WORK_ACTIVE.set()
            _run_local_area_roof_job(job["id"], job.get("user_key"))
            continue

        DESKTOP_BACKGROUND_WORK_ACTIVE.clear()
        # Close the small signal/claim race without busy-waiting.
        if store.has_queued_individual_jobs():
            DESKTOP_BACKGROUND_WORK_ACTIVE.set()
            continue
        if store.has_queued_area_jobs():
            DESKTOP_BACKGROUND_WORK_ACTIVE.set()
            continue
        _roof_worker_wake.wait(timeout=2.0)
        _roof_worker_wake.clear()


def _ensure_roof_worker_started():
    global _roof_worker_thread
    if not _roof_local_worker_enabled():
        return False
    with _roof_worker_lock:
        if _roof_worker_thread is None or not _roof_worker_thread.is_alive():
            DESKTOP_BACKGROUND_WORK_ACTIVE.set()
            _roof_worker_thread = threading.Thread(
                target=_roof_worker_loop,
                name="pcs-roof-intelligence-worker",
                daemon=True,
            )
            _roof_worker_thread.start()
    _roof_worker_wake.set()
    return True


_ensure_roof_worker_started()


def _roof_job_payload(store, job):
    if not job:
        return None
    result = dict(job)
    report = store.get_report_for_job(job["id"])
    if report:
        report["view_url"] = url_for("download_roof_intelligence_report", report_id=report["id"])
        if _roof_report_editing_enabled() and store.list_report_revisions(report["id"], job.get("user_key")):
            report["review_url"] = url_for("review_roof_intelligence_report", report_id=report["id"])
        result["report"] = report
    else:
        result["report"] = None
    reports = store.get_reports_for_job(job["id"])
    for batch_report in reports:
        batch_report["view_url"] = url_for("download_roof_intelligence_report", report_id=batch_report["id"])
        if _roof_report_editing_enabled() and store.list_report_revisions(batch_report["id"], job.get("user_key")):
            batch_report["review_url"] = url_for(
                "review_roof_intelligence_report", report_id=batch_report["id"]
            )
    result["reports"] = reports
    items = store.list_area_items(job["id"]) if job.get("job_type") == "zip_batch" else []
    result["items"] = items
    result["failed_items"] = [item for item in items if item.get("status") == "failed"]
    result["status_url"] = url_for("roof_intelligence_job_status", job_id=job["id"])
    result["page_url"] = url_for("roof_intelligence", job_id=job["id"])
    return result


@app.route('/roof-intelligence')
def roof_intelligence():
    store = get_job_store()
    requested_job_id = request.args.get("job_id", "").strip()
    active_job = store.get_job(requested_job_id, _roof_intelligence_user_key()) if requested_job_id else None
    recent_notifications = store.list_notifications(_roof_intelligence_user_key(), limit=30)
    selected_notifications = [
        item for item in recent_notifications
        if not active_job or item.get("job_id") in {None, active_job["id"]}
    ][:8]
    return render_template(
        'roof_intelligence.html',
        active_job=_roof_job_payload(store, active_job),
        recent_jobs=store.list_jobs(_roof_intelligence_user_key(), limit=12),
        notifications=selected_notifications,
        county_health=store.list_latest_county_health(limit=20),
        county_health_running=_county_health_check_running(),
        canonical_reviews=_canonical_footprint_reviews(limit=20),
        supported_roof_types=SUPPORTED_ROOF_TYPES,
        local_worker_enabled=_roof_local_worker_enabled(),
        google_maps_configured=bool(google_maps_api_key()),
        report_editing_enabled=_roof_report_editing_enabled(),
    )


@app.post('/roof-intelligence/county-health/run')
def run_roof_intelligence_county_health():
    try:
        started = _start_manual_county_health_check()
    except (OSError, RuntimeError, ValueError) as exc:
        flash(str(exc), "danger")
    else:
        if started:
            flash(
                "County health check started. Results will appear here when it finishes.",
                "success",
            )
        else:
            flash("A county health check is already running.", "success")
    return redirect(url_for("roof_intelligence"))


@app.get('/api/roof-intelligence/county-health/status')
def roof_intelligence_county_health_status():
    return jsonify({"running": _county_health_check_running()})


@app.route('/settings', methods=['GET', 'POST'])
def application_settings():
    if request.method == 'POST':
        action = request.form.get("action", "save")
        if action == "remove_google_maps_key":
            remove_google_maps_api_key()
            flash("The local Google Maps API key was removed.", "success")
            return redirect(url_for("application_settings"))
        if action == "remove_supabase_configuration":
            remove_supabase_configuration()
            tenant_sign_out()
            flash("The local Supabase configuration was removed.", "success")
            return redirect(url_for("application_settings"))
        try:
            if action == "save_supabase_configuration":
                submitted_url = request.form.get("supabase_url", "").strip().rstrip("/")
                submitted_key = request.form.get("supabase_publishable_key", "").strip()
                save_supabase_configuration(submitted_url, submitted_key)
                tenant_sign_out()
                flash("Supabase is configured. Sign in with your company account.", "success")
            elif action == "save_report_export_directory":
                saved_path = save_report_export_directory(
                    request.form.get("report_export_directory", "")
                )
                flash(f"Local report exports will be saved in {saved_path}.", "success")
            elif action == "create_report_folder":
                current_tenant_context()
                folder = TenantSettingsStore.from_current_session().create_report_folder(
                    request.form.get("report_folder_name", "")
                )
                flash(f"Created the protected report folder {folder['name']}.", "success")
            elif action == "set_default_report_folder":
                current_tenant_context()
                TenantSettingsStore.from_current_session().set_default_report_folder(
                    request.form.get("default_report_folder_id", "")
                )
                flash("The default protected report folder was updated.", "success")
            else:
                save_google_maps_api_key(request.form.get("google_maps_api_key", ""))
                flash("Google Maps is configured for Roof Intelligence.", "success")
        except (ValueError, ContactStoreError, TenantAuthenticationError) as exc:
            flash(str(exc), "danger")
        else:
            return redirect(url_for("application_settings"))
    key = google_maps_api_key()
    supabase_url, supabase_key = supabase_configuration()
    tenant = None
    report_folders = []
    tenant_settings = {}
    try:
        tenant = current_tenant_context()
        settings_store = TenantSettingsStore.from_current_session()
        report_folders = settings_store.list_report_folders()
        tenant_settings = settings_store.get_settings()
    except (TenantAuthenticationError, ContactStoreError):
        pass
    return render_template(
        'settings.html',
        google_maps_configured=bool(key),
        google_maps_key_suffix=key[-4:] if key else "",
        supabase_configured=bool(supabase_key),
        supabase_url=supabase_url,
        supabase_key_suffix=supabase_key[-4:] if supabase_key else "",
        tenant=tenant,
        report_folders=report_folders,
        tenant_settings=tenant_settings,
        report_export_directory=report_export_directory(),
    )


@app.get('/api/local-settings/google-maps')
def google_maps_browser_configuration():
    key = google_maps_api_key()
    if not key:
        return jsonify({"configured": False}), 404
    response = jsonify({"configured": True, "api_key": key})
    response.headers["Cache-Control"] = "no-store, private"
    response.headers["Pragma"] = "no-cache"
    return response


@app.post('/roof-intelligence/jobs/individual')
def create_individual_roof_intelligence_job():
    store = get_job_store()
    try:
        address = request.form.get("property_address", "")
        if _roof_local_worker_enabled():
            readiness_error = _roof_worker_readiness_error()
            if readiness_error:
                raise ValueError(readiness_error)
        job = store.create_individual_job(
            address,
            user_key=_roof_intelligence_user_key(),
        )
    except ValueError as exc:
        flash(str(exc), "danger")
        return redirect(url_for("roof_intelligence", mode="individual"))

    if _roof_local_worker_enabled():
        DESKTOP_BACKGROUND_WORK_ACTIVE.set()
        _ensure_roof_worker_started()
    return redirect(url_for("roof_intelligence", job_id=job["id"], mode="individual"))


@app.post('/roof-intelligence/jobs/area')
def create_area_roof_intelligence_job():
    store = get_job_store()
    try:
        job = store.create_area_job(
            request.form.get("bounds_north", ""),
            request.form.get("bounds_south", ""),
            request.form.get("bounds_east", ""),
            request.form.get("bounds_west", ""),
            minimum_roof_squares=request.form.get("minimum_roof_squares", "100"),
            roof_types=request.form.getlist("roof_types"),
            user_key=_roof_intelligence_user_key(),
            selection_type=request.form.get("selection_type", "rectangle"),
            center_lat=request.form.get("center_lat", ""),
            center_lng=request.form.get("center_lng", ""),
            center_address=request.form.get("center_address", ""),
            radius_miles=request.form.get("radius_miles", ""),
        )
    except ValueError as exc:
        flash(str(exc), "danger")
        return redirect(url_for("roof_intelligence", mode="area"))
    if _roof_local_worker_enabled():
        DESKTOP_BACKGROUND_WORK_ACTIVE.set()
        _ensure_roof_worker_started()
    return redirect(url_for("roof_intelligence", job_id=job["id"], mode="area"))


@app.get('/api/roof-intelligence/jobs/<job_id>')
def roof_intelligence_job_status(job_id):
    store = get_job_store()
    job = store.get_job(job_id, _roof_intelligence_user_key())
    if not job:
        return jsonify({"error": "Roof Intelligence job not found."}), 404
    return jsonify(_roof_job_payload(store, job))


@app.post('/roof-intelligence/jobs/<job_id>/cancel')
def cancel_roof_intelligence_job(job_id):
    store = get_job_store()
    job = store.cancel_job(job_id, _roof_intelligence_user_key())
    if not job:
        flash("Roof Intelligence job not found.", "danger")
        return redirect(url_for("roof_intelligence"))
    _roof_worker_wake.set()
    flash("The Roof Intelligence job was cancelled.", "success")
    return redirect(url_for("roof_intelligence", job_id=job_id, mode="individual" if job["job_type"] == "individual_address" else "area"))


@app.post('/roof-intelligence/jobs/<job_id>/resolve-footprint')
def resolve_roof_footprint_discrepancy(job_id):
    store = get_job_store()
    try:
        job = store.resolve_footprint_discrepancy(
            job_id,
            request.form.get("selected_source", ""),
            request.form.get("reason", ""),
            user_key=_roof_intelligence_user_key(),
            item_id=request.form.get("item_id", "").strip() or None,
        )
    except (ValueError, KeyError) as exc:
        flash(str(exc), "danger")
        return redirect(url_for("roof_intelligence", job_id=job_id))
    if _roof_local_worker_enabled():
        DESKTOP_BACKGROUND_WORK_ACTIVE.set()
        _ensure_roof_worker_started()
    flash("The footprint resolution was recorded and the report was queued again.", "success")
    return redirect(
        url_for(
            "roof_intelligence",
            job_id=job_id,
            mode="individual" if job["job_type"] == "individual_address" else "area",
        )
    )


@app.post('/roof-intelligence/canonical-footprints/<int:canonical_id>/resolve')
def resolve_canonical_footprint_review(canonical_id):
    try:
        selected_source = request.form.get("selected_source", "")
        reason = request.form.get("reason", "")
        if selected_source not in {"supabase", "county"}:
            raise ValueError("Select either the Microsoft or county footprint.")
        if len(" ".join(reason.split())) < 10:
            raise ValueError("Enter a resolution reason of at least 10 characters.")
        _resolve_canonical_footprint(canonical_id, selected_source, reason)
    except (ValueError, RuntimeError) as exc:
        flash(str(exc), "danger")
    else:
        flash("The canonical footprint decision was recorded for future reports.", "success")
    return redirect(url_for("roof_intelligence"))


@app.post('/roof-intelligence/notifications/<notification_id>/read')
def mark_roof_intelligence_notification_read(notification_id):
    store = get_job_store()
    store.mark_notification_read(notification_id, _roof_intelligence_user_key())
    job_id = request.form.get("job_id", "").strip()
    return redirect(url_for("roof_intelligence", job_id=job_id) if job_id else url_for("roof_intelligence"))


@app.route('/roof-intelligence/reports/<report_id>')
def download_roof_intelligence_report(report_id):
    store = get_job_store()
    user_key = _roof_intelligence_user_key()
    report = store.get_report(report_id, user_key)
    if not report:
        return "Report was not found.", 404
    report_path = str(report.get("report_path") or "")
    if not report_path or not os.path.isfile(report_path):
        return "The local report file is no longer available.", 404

    requested_path = os.path.realpath(report_path)
    tenant_report_root, _ = _tenant_report_output_paths(user_key)
    allowed_roots = (
        os.path.realpath(ROOF_INTELLIGENCE_PROJECT_DIR),
        os.path.realpath(tenant_report_root),
        os.path.realpath(os.path.join(str(DEFAULT_DATA_DIR), "roof_intelligence_reports")),
    )
    if not any(requested_path.startswith(root + os.sep) for root in allowed_roots):
        return "Report path is not allowed.", 403
    return send_file(requested_path, mimetype="application/pdf", as_attachment=False)


@app.get('/roof-intelligence/reports/<report_id>/review')
def review_roof_intelligence_report(report_id):
    if not _roof_report_editing_enabled():
        return "Report review and editing are not enabled.", 404
    report = get_job_store().get_report_review(report_id, _roof_intelligence_user_key())
    if not report:
        return "Report was not found.", 404
    for revision in report["revisions"]:
        revision["view_url"] = url_for(
            "download_roof_intelligence_revision",
            report_id=report_id,
            revision_id=revision["id"],
        )
    return render_template("roof_intelligence_report_review.html", report=report)


@app.post('/roof-intelligence/reports/<report_id>/revisions')
def create_roof_intelligence_revision(report_id):
    if not _roof_report_editing_enabled():
        return "Report review and editing are not enabled.", 404
    store = get_job_store()
    user_key = _roof_intelligence_user_key()
    report = store.get_report_review(report_id, user_key)
    if not report or not report.get("latest_revision"):
        flash("This report does not contain an editable revision snapshot.", "danger")
        return redirect(url_for("roof_intelligence"))
    parent = report["latest_revision"]
    parent_snapshot = parent["snapshot"]
    analysis = parent_snapshot.get("analysis") or {}
    report_fields = parent_snapshot.get("report_fields") or {}
    reason = " ".join(request.form.get("change_reason", "").split())
    if len(reason) < 10:
        flash("Enter a change reason of at least 10 characters.", "danger")
        return redirect(url_for("review_roof_intelligence_report", report_id=report_id))

    try:
        roof_area = float(request.form.get("roof_area_sqft", ""))
        condition_score = float(request.form.get("roof_condition_score", ""))
    except ValueError:
        flash("Roof area and condition score must be numeric.", "danger")
        return redirect(url_for("review_roof_intelligence_report", report_id=report_id))
    if not math.isfinite(roof_area) or not math.isfinite(condition_score) or roof_area < 0 or not 0 <= condition_score <= 100:
        flash("Roof area must be nonnegative and condition score must be between 0 and 100.", "danger")
        return redirect(url_for("review_roof_intelligence_report", report_id=report_id))

    submitted = {
        "roof_area_sqft": roof_area,
        "roof_type": request.form.get("roof_type", "").strip(),
        "roof_system": request.form.get("roof_system", "").strip(),
        "roof_condition_score": condition_score,
        "report_summary": request.form.get("report_summary", ""),
        "recommendation": request.form.get("recommendation", ""),
    }
    if not submitted["roof_type"] or not submitted["roof_system"]:
        flash("Roof type and roof information are required.", "danger")
        return redirect(url_for("review_roof_intelligence_report", report_id=report_id))
    current = {
        "roof_area_sqft": float(report_fields.get("roof_area_sqft") or 0),
        "roof_type": str(analysis.get("roof_type") or ""),
        "roof_system": str(analysis.get("roof_system") or ""),
        "roof_condition_score": float(analysis.get("overall_score") or 0),
        "report_summary": str(analysis.get("summary") or ""),
        "recommendation": str(analysis.get("recommendation") or ""),
    }
    edits = {key: value for key, value in submitted.items() if value != current[key]}
    if not edits:
        flash("Change at least one report field before creating a revision.", "danger")
        return redirect(url_for("review_roof_intelligence_report", report_id=report_id))
    apply_to_future = request.form.get("apply_square_footage_to_future") == "1"
    submit_for_future_processing = request.form.get("submit_for_future_processing") == "1"
    if apply_to_future and "roof_area_sqft" not in edits:
        flash("The future-report option requires a square-footage change.", "danger")
        return redirect(url_for("review_roof_intelligence_report", report_id=report_id))

    python_path = os.path.join(ROOF_INTELLIGENCE_PROJECT_DIR, ".venv", "bin", "python")
    revision_script = os.path.join(ROOF_INTELLIGENCE_PROJECT_DIR, "roof_intelligence_revision_service.py")
    if not os.path.isfile(python_path) or not os.path.isfile(revision_script):
        flash("The PilotPoint revision service is not available.", "danger")
        return redirect(url_for("review_roof_intelligence_report", report_id=report_id))
    revision_number = int(parent["revision_number"]) + 1
    tenant_report_root, _ = _tenant_report_output_paths(_roof_intelligence_user_key())
    output_dir = os.path.join(
        tenant_report_root,
        report_id,
        f"revision-{revision_number}",
    )
    os.makedirs(output_dir, exist_ok=True)
    output_pdf = os.path.join(
        output_dir,
        roof_report_pdf_filename(
            report.get("address"),
            report.get("city"),
        ),
    )
    image_path = str((parent_snapshot.get("imagery") or {}).get("local_report_image_path") or "")

    try:
        with tempfile.TemporaryDirectory(prefix="pcs-roof-revision-") as temporary_dir:
            parent_path = os.path.join(temporary_dir, "parent.json")
            edits_path = os.path.join(temporary_dir, "edits.json")
            snapshot_path = os.path.join(temporary_dir, "revision.json")
            with open(parent_path, "w", encoding="utf-8") as handle:
                json.dump(parent_snapshot, handle)
            with open(edits_path, "w", encoding="utf-8") as handle:
                json.dump(edits, handle)
            command = [
                python_path,
                revision_script,
                parent_path,
                edits_path,
                output_pdf,
                snapshot_path,
                "--created-by",
                _roof_intelligence_user_key(),
                "--change-reason",
                reason,
            ]
            if image_path and os.path.isfile(image_path):
                command.extend(["--report-image", image_path])
            if apply_to_future:
                command.append("--apply-square-footage-to-future")
            if submit_for_future_processing:
                feedback_directory = os.path.join(
                    os.path.dirname(tenant_report_root),
                    "roof_processing_feedback",
                )
                command.extend(
                    [
                        "--submit-for-future-processing",
                        "--feedback-directory",
                        feedback_directory,
                    ]
                )
            completed = subprocess.run(
                command,
                cwd=ROOF_INTELLIGENCE_PROJECT_DIR,
                capture_output=True,
                text=True,
                timeout=180,
                check=False,
            )
            if completed.returncode != 0:
                raise RuntimeError((completed.stderr or completed.stdout or "Revision generation failed.").strip())
            service_result = json.loads(completed.stdout or "{}")
            with open(snapshot_path, "r", encoding="utf-8") as handle:
                revised_snapshot = json.load(handle)
            processing_feedback = None
            if submit_for_future_processing:
                feedback_path = str(service_result.get("processing_feedback_path") or "")
                if not feedback_path or not os.path.isfile(feedback_path):
                    raise RuntimeError("The future-processing feedback record was not created.")
                with open(feedback_path, "r", encoding="utf-8") as handle:
                    processing_feedback = json.load(handle)
        pdf_size = os.path.getsize(output_pdf)
        digest = hashlib.sha256()
        with open(output_pdf, "rb") as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                digest.update(chunk)
        store.save_ready_report_revision(
            report_id,
            parent["id"],
            revised_snapshot,
            report_path=output_pdf,
            pdf_size=pdf_size,
            pdf_checksum=digest.hexdigest(),
            created_by=_roof_intelligence_user_key(),
            change_reason=reason,
            edits=edits,
            apply_square_footage_to_future=apply_to_future,
            processing_feedback=processing_feedback,
            user_key=user_key,
        )
    except (OSError, RuntimeError, ValueError, KeyError, subprocess.SubprocessError) as exc:
        flash(f"The revised report could not be generated: {' '.join(str(exc).split())[:400]}", "danger")
        return redirect(url_for("review_roof_intelligence_report", report_id=report_id))
    message = f"Revision {revision_number} is ready."
    if submit_for_future_processing:
        message += " Your correction was submitted for future-processing review."
    flash(message, "success")
    return redirect(url_for("review_roof_intelligence_report", report_id=report_id))


@app.get('/roof-intelligence/reports/<report_id>/revisions/<revision_id>')
def download_roof_intelligence_revision(report_id, revision_id):
    if not _roof_report_editing_enabled():
        return "Report review and editing are not enabled.", 404
    user_key = _roof_intelligence_user_key()
    revision = get_job_store().get_report_revision(revision_id, user_key)
    if not revision or revision["report_id"] != report_id:
        return "Report revision was not found.", 404
    report_path = os.path.realpath(str(revision.get("report_path") or ""))
    tenant_report_root, _ = _tenant_report_output_paths(user_key)
    allowed_roots = (
        os.path.realpath(tenant_report_root),
        os.path.realpath(os.path.join(str(DEFAULT_DATA_DIR), "roof_intelligence_reports")),
        os.path.realpath(ROOF_INTELLIGENCE_PROJECT_DIR),
    )
    if not report_path or not os.path.isfile(report_path):
        return "The report revision file is no longer available.", 404
    if not any(report_path.startswith(root + os.sep) for root in allowed_roots):
        return "Report revision path is not allowed.", 403
    return send_file(report_path, mimetype="application/pdf", as_attachment=False)


@app.route('/proposal-tracker')
def proposal_tracker():
    entries, tracker_error = load_proposal_tracker_missing_entries()
    saved_count = request.args.get("saved_count", "").strip()
    save_error = request.args.get("save_error", "").strip()
    return render_template(
        'proposal_tracker.html',
        entries=entries,
        tracker_error=tracker_error,
        saved_count=saved_count,
        save_error=save_error,
    )


@app.route('/proposal-tracker/save', methods=['POST'])
def save_proposal_tracker():
    entries = []
    for row_number in request.form.getlist("row_number"):
        row_key = str(row_number or "").strip()
        if row_key.startswith("new_"):
            customer_name = request.form.get(f"customer_name_{row_key}", "")
            project_street_address = request.form.get(
                f"project_street_address_{row_key}", ""
            )
            customer = " - ".join(
                value for value in (
                    str(customer_name or "").strip(),
                    str(project_street_address or "").strip(),
                ) if value
            )
            if not any([
                str(customer or "").strip(),
                request.form.get(f"contact_{row_key}", ""),
                request.form.get(f"email_address_{row_key}", ""),
                request.form.get(f"lead_source_{row_key}", ""),
                request.form.get(f"submitted_by_{row_key}", ""),
                request.form.get(f"estimated_by_{row_key}", ""),
                request.form.get(f"estimate_date_{row_key}", ""),
                request.form.get(f"proposal_date_{row_key}", ""),
                request.form.get(f"follow_up_date_{row_key}", ""),
                request.form.get(f"status_{row_key}", ""),
            ]):
                continue
            entries.append({
                "is_new": True,
                "row_number": row_key,
                "customer": customer,
                "customer_name": customer_name,
                "project_street_address": project_street_address,
                "contact": request.form.get(f"contact_{row_key}", ""),
                "email_address": request.form.get(f"email_address_{row_key}", ""),
                "lead_source": request.form.get(f"lead_source_{row_key}", ""),
                "submitted_by": request.form.get(f"submitted_by_{row_key}", ""),
                "estimated_by": request.form.get(f"estimated_by_{row_key}", ""),
                "estimate_date": request.form.get(f"estimate_date_{row_key}", ""),
                "proposal_date": request.form.get(f"proposal_date_{row_key}", ""),
                "follow_up_date": request.form.get(f"follow_up_date_{row_key}", ""),
                "status": request.form.get(f"status_{row_key}", "draft"),
            })
            continue
        if not re.fullmatch(r"[A-Za-z0-9_-]+", row_key):
            continue
        entries.append({
            "is_new": False,
            "row_number": row_key,
            "customer": request.form.get(f"customer_{row_key}", ""),
            "contact": request.form.get(f"contact_{row_key}", ""),
            "email_address": request.form.get(f"email_address_{row_key}", ""),
            "lead_source": request.form.get(f"lead_source_{row_key}", ""),
            "submitted_by": request.form.get(f"submitted_by_{row_key}", ""),
            "estimated_by": request.form.get(f"estimated_by_{row_key}", ""),
            "estimate_date": request.form.get(f"estimate_date_{row_key}", ""),
            "proposal_date": request.form.get(f"proposal_date_{row_key}", ""),
            "follow_up_date": request.form.get(f"follow_up_date_{row_key}", ""),
            "status": request.form.get(f"status_{row_key}", ""),
        })

    try:
        saved_count = update_proposal_tracker_missing_entries(entries)
    except Exception as exc:
        _safe_debug(f"[ERROR] Proposal tracker save failed: {exc}")
        return redirect(url_for("proposal_tracker", save_error=str(exc)))

    return redirect(url_for("proposal_tracker", saved_count=saved_count))


def _coerce_tracker_date(value):
    if value is None:
        return None
    if isinstance(value, datetime.datetime):
        return value.date()
    if isinstance(value, datetime.date):
        return value
    if isinstance(value, (int, float)):
        try:
            converted = from_excel(value)
            if isinstance(converted, datetime.datetime):
                return converted.date()
            if isinstance(converted, datetime.date):
                return converted
        except Exception:
            return None
    if isinstance(value, str):
        cleaned = value.strip()
        if not cleaned or cleaned in {"-", "N/A", "n/a"}:
            return None
        for fmt in ("%m/%d/%Y", "%m/%d/%y", "%Y-%m-%d"):
            try:
                return datetime.datetime.strptime(cleaned, fmt).date()
            except ValueError:
                pass
    return None


def _format_tracker_date(value):
    parsed = _coerce_tracker_date(value)
    if parsed is None:
        return ""
    return f"{parsed.month}/{parsed.day}/{parsed.year}"


def _format_tracker_date_input(value):
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    parsed = _coerce_tracker_date(value)
    return "" if parsed is None else f"{parsed.month}/{parsed.day}/{parsed.year}"


_TRACKER_STATUS_LABELS = {
    "draft": "Draft Unsent",
    "sent": "Sent",
    "under_contract": "Under Contract",
    "finished": "Finished",
    "dead": "Dead",
}
_TRACKER_STATUS_ALIASES = {
    "draft unsent": "draft",
    "follow_up": "sent",
    "follow up": "sent",
    "won": "under_contract",
    "under contract": "under_contract",
    "lost": "dead",
    "withdrawn": "dead",
    "archived": "dead",
}


def _normalize_tracker_status(value, proposal_date=None):
    cleaned = " ".join(str(value or "").strip().casefold().replace("-", " ").split())
    cleaned = _TRACKER_STATUS_ALIASES.get(cleaned, cleaned.replace(" ", "_"))
    if cleaned in _TRACKER_STATUS_LABELS:
        return cleaned
    return "sent" if _coerce_tracker_date(proposal_date) is not None else "draft"


def _tracker_status_label(value, proposal_date=None):
    return _TRACKER_STATUS_LABELS[_normalize_tracker_status(value, proposal_date)]


def _default_follow_up_cutoff_date():
    return datetime.date.today() - datetime.timedelta(days=14)


def _parse_iso_date(value, fallback_date):
    try:
        return datetime.datetime.strptime(str(value or "").strip(), "%Y-%m-%d").date()
    except Exception:
        return fallback_date


def _parse_optional_iso_date(value):
    cleaned = str(value or "").strip()
    if not cleaned:
        return None
    return _parse_iso_date(cleaned, None)


def _tracker_cell_is_blank(value):
    return str(value or "").strip() == ""


def _copy_tracker_row_style(ws, source_row, target_row):
    if source_row < 1 or target_row < 1 or source_row == target_row:
        return
    max_column = max(ws.max_column, 11)
    for col_idx in range(1, max_column + 1):
        source_cell = ws.cell(row=source_row, column=col_idx)
        target_cell = ws.cell(row=target_row, column=col_idx)
        if source_cell.has_style:
            target_cell._style = _copy_style(source_cell._style)
        if source_cell.number_format:
            target_cell.number_format = source_cell.number_format
        if source_cell.font:
            target_cell.font = _copy_style(source_cell.font)
        if source_cell.fill:
            target_cell.fill = _copy_style(source_cell.fill)
        if source_cell.border:
            target_cell.border = _copy_style(source_cell.border)
        if source_cell.alignment:
            target_cell.alignment = _copy_style(source_cell.alignment)
        if source_cell.protection:
            target_cell.protection = _copy_style(source_cell.protection)


def _find_tracker_insert_row(ws, customer_name):
    new_key = str(customer_name or "").strip().casefold()
    first_data_row = 2
    last_row = max(ws.max_row, first_data_row - 1)
    customer_column = _proposal_tracker_column_map(ws)["customer"]
    for row_number in range(first_data_row, last_row + 1):
        existing_key = str(
            ws.cell(row=row_number, column=customer_column).value or ""
        ).strip().casefold()
        if existing_key and existing_key > new_key:
            return row_number
    return last_row + 1


def _write_tracker_entry_to_row(ws, row_number, entry):
    columns = _proposal_tracker_column_map(ws)
    values = {
        "customer": entry.get("customer"),
        "contact": entry.get("contact"),
        "email_address": entry.get("email_address"),
        "lead_source": entry.get("lead_source"),
        "submitted_by": entry.get("submitted_by"),
        "estimate_date": entry.get("estimate_date"),
        "proposal_date": entry.get("proposal_date"),
        "follow_up_date": entry.get("follow_up_date"),
        "status": _tracker_status_label(
            entry.get("status"), entry.get("proposal_date")
        ),
        "estimated_by": entry.get("estimated_by"),
    }
    for field_name, value in values.items():
        cell = ws.cell(row=row_number, column=columns[field_name])
        if field_name.endswith("_date"):
            cell.number_format = "General"
        cell.value = str(value or "").strip()


def _load_proposal_tracker_missing_entries_spreadsheet(tracker_path=PROPOSAL_TRACKER):
    entries = []

    try:
        with TRACKER_IO_LOCK:
            source_path = _proposal_tracker_source_path(tracker_path)
            if not source_path:
                raise RuntimeError(f"Proposal Tracking.xlsx not found: {tracker_path}")
            wb = load_workbook(source_path, data_only=True, read_only=True)
            ws = wb.active
            try:
                columns = _proposal_tracker_column_map(ws)
                for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                    customer = _proposal_tracker_row_value(row, columns["customer"])
                    contact = _proposal_tracker_row_value(row, columns["contact"])
                    email_address = _proposal_tracker_row_value(row, columns["email_address"])
                    lead = _proposal_tracker_row_value(row, columns["lead_source"])
                    submitted_by = _proposal_tracker_row_value(row, columns["submitted_by"])
                    estimate_dt_raw = _proposal_tracker_row_value(row, columns["estimate_date"])
                    proposal_dt_raw = _proposal_tracker_row_value(row, columns["proposal_date"])
                    follow_up_raw = _proposal_tracker_row_value(row, columns["follow_up_date"])
                    estimated_by = _proposal_tracker_row_value(row, columns["estimated_by"])
                    status = _proposal_tracker_row_value(row, columns["status"])
                    normalized_status = _normalize_tracker_status(status, proposal_dt_raw)

                    if not any([customer, contact, email_address, lead, submitted_by, proposal_dt_raw, follow_up_raw]):
                        continue
                    if normalized_status == "dead":
                        continue
                    if not any([
                        _tracker_cell_is_blank(contact),
                        _tracker_cell_is_blank(email_address),
                        _tracker_cell_is_blank(proposal_dt_raw),
                    ]):
                        continue

                    entries.append({
                        "row_number": row_idx,
                        "customer": str(customer or "").strip(),
                        "contact": str(contact or "").strip(),
                        "email_address": str(email_address or "").strip(),
                        "lead_source": str(lead or "").strip(),
                        "submitted_by": str(submitted_by or "").strip(),
                        "estimated_by": str(estimated_by or "").strip(),
                        "estimate_date_input": _format_tracker_date_input(estimate_dt_raw),
                        "proposal_date_input": _format_tracker_date_input(proposal_dt_raw),
                        "follow_up_date_input": _format_tracker_date_input(follow_up_raw),
                        "status": normalized_status,
                    })
            finally:
                wb.close()
    except Exception as exc:
        _safe_debug(f"[ERROR] Proposal tracker read failed: {exc}")
        return [], str(exc)

    entries.sort(key=lambda item: item["customer"].casefold())
    return entries, None


def load_proposal_tracker_missing_entries(tracker_path=PROPOSAL_TRACKER):
    if load_proposal_tracking_cutover_flags().reads_enabled:
        try:
            return get_proposal_tracking_store().list_missing_entries(), None
        except Exception as exc:
            _safe_debug(f"[ERROR] Supabase proposal tracker read failed: {exc}")
            return [], str(exc)
    return _load_proposal_tracker_missing_entries_spreadsheet(tracker_path)


def _update_proposal_tracker_missing_entries_spreadsheet(entries, tracker_path=PROPOSAL_TRACKER):
    if not entries:
        return 0

    source_path = _proposal_tracker_source_path(tracker_path)
    if not source_path:
        raise RuntimeError(f"Proposal Tracking.xlsx not found: {tracker_path}")

    wb = None
    temp_path = None
    try:
        with TRACKER_IO_LOCK:
            wb = load_workbook(source_path)
            ws = wb.active
            _ensure_proposal_tracker_status_column(ws)
            columns = _proposal_tracker_column_map(ws)
            updated_count = 0
            existing_entries = [entry for entry in entries if not entry.get("is_new")]
            new_entries = [
                entry for entry in entries
                if entry.get("is_new") and str(entry.get("customer") or "").strip()
            ]

            for entry in existing_entries:
                row_number = int(entry.get("row_number") or 0)
                if row_number < 2 or row_number > ws.max_row:
                    continue

                values = {
                    "contact": entry.get("contact"),
                    "email_address": entry.get("email_address"),
                    "lead_source": entry.get("lead_source"),
                    "submitted_by": entry.get("submitted_by"),
                    "estimate_date": entry.get("estimate_date"),
                    "proposal_date": entry.get("proposal_date"),
                    "follow_up_date": entry.get("follow_up_date"),
                    "status": _tracker_status_label(
                        entry.get("status"), entry.get("proposal_date")
                    ),
                    "estimated_by": entry.get("estimated_by"),
                }
                for field_name, value in values.items():
                    cell = ws.cell(row=row_number, column=columns[field_name])
                    if field_name.endswith("_date"):
                        cell.number_format = "General"
                    cell.value = str(value or "").strip()
                updated_count += 1

            for entry in sorted(new_entries, key=lambda item: str(item.get("customer") or "").casefold()):
                insert_row = _find_tracker_insert_row(ws, entry.get("customer"))
                if insert_row <= ws.max_row:
                    ws.insert_rows(insert_row)
                    style_source_row = insert_row + 1 if insert_row + 1 <= ws.max_row else insert_row - 1
                    _copy_tracker_row_style(ws, style_source_row, insert_row)
                elif ws.max_row >= 2:
                    _copy_tracker_row_style(ws, ws.max_row, insert_row)
                _write_tracker_entry_to_row(ws, insert_row, entry)
                updated_count += 1

            temp_path = _proposal_tracker_temp_path(tracker_path)
            wb.save(temp_path)
            wb.close()
            wb = None
            _replace_proposal_tracker_file(temp_path, tracker_path)
            temp_path = None
            return updated_count
    finally:
        try:
            if wb is not None:
                wb.close()
        except Exception:
            pass
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass


def update_proposal_tracker_missing_entries(entries, tracker_path=PROPOSAL_TRACKER):
    flags = load_proposal_tracking_cutover_flags()
    if flags.reads_enabled and not flags.writes_enabled:
        raise RuntimeError(
            "Supabase proposal-tracking reads require Supabase writes to be enabled."
        )
    supabase_error = None
    updated_count = 0
    if flags.writes_enabled:
        try:
            updated_count = get_proposal_tracking_store().update_entries(entries)
        except Exception as exc:
            supabase_error = exc
            _safe_debug(f"[ERROR] Supabase proposal tracker save failed: {exc}")
    if flags.spreadsheet_writes_active:
        spreadsheet_count = _update_proposal_tracker_missing_entries_spreadsheet(
            entries, tracker_path
        )
        if not flags.writes_enabled:
            updated_count = spreadsheet_count
    elif supabase_error is not None:
        raise supabase_error
    return updated_count


def _load_weekly_follow_up_entries_spreadsheet(tracker_path=PROPOSAL_TRACKER, cutoff_date=None):
    cutoff_date = cutoff_date or _default_follow_up_cutoff_date()
    entries = []

    try:
        with TRACKER_IO_LOCK:
            wb = load_workbook(tracker_path, data_only=True, read_only=True)
            ws = wb.active
            try:
                columns = _proposal_tracker_column_map(ws)
                for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                    customer = _proposal_tracker_row_value(row, columns["customer"])
                    contact = _proposal_tracker_row_value(row, columns["contact"])
                    email_address = _proposal_tracker_row_value(row, columns["email_address"])
                    submitted_by = _proposal_tracker_row_value(row, columns["submitted_by"])
                    proposal_dt_raw = _proposal_tracker_row_value(row, columns["proposal_date"])
                    follow_up_raw = _proposal_tracker_row_value(row, columns["follow_up_date"])
                    status = _normalize_tracker_status(
                        _proposal_tracker_row_value(row, columns["status"]),
                        proposal_dt_raw,
                    )
                    proposal_dt = _coerce_tracker_date(proposal_dt_raw)

                    if not any([customer, contact, email_address, submitted_by, proposal_dt_raw]):
                        continue
                    if str(follow_up_raw or "").strip():
                        continue
                    if status != "sent":
                        continue
                    if proposal_dt is None or proposal_dt > cutoff_date:
                        continue

                    entries.append({
                        "row_number": row_idx,
                        "customer": str(customer or "").strip(),
                        "contact": str(contact or "").strip(),
                        "email_address": str(email_address or "").strip(),
                        "proposal_date": proposal_dt,
                        "proposal_date_display": _format_tracker_date(proposal_dt),
                        "submitted_by": str(submitted_by or "").strip(),
                    })
            finally:
                wb.close()
    except Exception as exc:
        _safe_debug(f"[ERROR] Weekly follow-up tracker read failed: {exc}")
        return [], cutoff_date, str(exc)

    entries.sort(key=lambda item: (
        item["submitted_by"].casefold(),
        item["proposal_date"],
        item["customer"].casefold(),
    ))
    return entries, cutoff_date, None


def load_weekly_follow_up_entries(tracker_path=PROPOSAL_TRACKER, cutoff_date=None):
    cutoff_date = cutoff_date or _default_follow_up_cutoff_date()
    if load_proposal_tracking_cutover_flags().reads_enabled:
        try:
            entries = get_proposal_tracking_store().list_weekly_follow_ups(cutoff_date)
            return entries, cutoff_date, None
        except Exception as exc:
            _safe_debug(f"[ERROR] Supabase weekly follow-up read failed: {exc}")
            return [], cutoff_date, str(exc)
    return _load_weekly_follow_up_entries_spreadsheet(tracker_path, cutoff_date)


def _update_weekly_follow_up_dates_spreadsheet(row_numbers, follow_up_date=None, tracker_path=PROPOSAL_TRACKER):
    row_numbers = {
        int(row_number)
        for row_number in (row_numbers or [])
        if str(row_number).strip().isdigit() and int(row_number) >= 2
    }
    if not row_numbers:
        return 0

    follow_up_date = follow_up_date or datetime.date.today()
    source_path = _proposal_tracker_source_path(tracker_path)
    if not source_path:
        raise RuntimeError(f"Proposal Tracking.xlsx not found: {tracker_path}")

    wb = None
    temp_path = None
    try:
        with TRACKER_IO_LOCK:
            wb = load_workbook(source_path)
            ws = wb.active
            _ensure_proposal_tracker_status_column(ws)
            follow_up_column = _proposal_tracker_column_map(ws)["follow_up_date"]
            status_column = _proposal_tracker_column_map(ws)["status"]
            updated_count = 0
            for row_number in sorted(row_numbers):
                if row_number > ws.max_row:
                    continue
                cell = ws.cell(row=row_number, column=follow_up_column)
                cell.value = follow_up_date
                cell.number_format = "m/d/yyyy"
                ws.cell(row=row_number, column=status_column).value = "Sent"
                updated_count += 1

            temp_path = _proposal_tracker_temp_path(tracker_path)
            wb.save(temp_path)
            wb.close()
            wb = None
            _replace_proposal_tracker_file(temp_path, tracker_path)
            temp_path = None
            return updated_count
    finally:
        try:
            if wb is not None:
                wb.close()
        except Exception:
            pass
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass


def update_weekly_follow_up_dates(row_numbers, follow_up_date=None, tracker_path=PROPOSAL_TRACKER):
    flags = load_proposal_tracking_cutover_flags()
    if flags.reads_enabled and not flags.writes_enabled:
        raise RuntimeError(
            "Supabase proposal-tracking reads require Supabase writes to be enabled."
        )
    follow_up_date = follow_up_date or datetime.date.today()
    supabase_error = None
    updated_count = 0
    if flags.writes_enabled:
        try:
            updated_count = get_proposal_tracking_store().mark_follow_ups(
                row_numbers, follow_up_date
            )
        except Exception as exc:
            supabase_error = exc
            _safe_debug(f"[ERROR] Supabase follow-up update failed: {exc}")
    if flags.spreadsheet_writes_active:
        spreadsheet_count = _update_weekly_follow_up_dates_spreadsheet(
            row_numbers, follow_up_date, tracker_path
        )
        if not flags.writes_enabled:
            updated_count = spreadsheet_count
    elif supabase_error is not None:
        raise supabase_error
    return updated_count


def _build_weekly_follow_up_email_bodies(submitter, follow_ups, cutoff_date):
    submitter_name = str(submitter or "Unassigned").strip() or "Unassigned"
    cutoff_display = _format_tracker_date(cutoff_date)
    subject_text = f"{submitter_name} Follow-Up List - {cutoff_display}"

    plain_lines = [
        f"{submitter_name} Follow-Up List",
        f"Proposal dates on or before {cutoff_display}",
        "",
    ]
    for item in follow_ups:
        plain_lines.extend([
            str(item.get("customer") or ""),
            f"Contact: {item.get('contact') or ''}",
            f"Email: {item.get('email_address') or ''}",
            f"Proposal Date: {item.get('proposal_date_display') or ''}",
            "",
        ])
    plain_body = "\n".join(plain_lines).rstrip()

    row_html = []
    for item in follow_ups:
        email_address = str(item.get("email_address") or "").strip()
        email_cell = ""
        if email_address:
            email_cell = (
                f'<a href="mailto:{html.escape(email_address, quote=True)}">'
                f"{html.escape(email_address)}</a>"
            )
        row_html.append(
            "<tr>"
            f"<td>{html.escape(str(item.get('customer') or ''))}</td>"
            f"<td>{html.escape(str(item.get('contact') or ''))}</td>"
            f"<td>{email_cell}</td>"
            f"<td>{html.escape(str(item.get('proposal_date_display') or ''))}</td>"
            "</tr>"
        )

    html_body = (
        '<html><body style="font-family: Aptos, Arial, Helvetica, sans-serif; font-size: 12pt; color: #212121;">'
        f"<p>{html.escape(submitter_name)} Follow-Up List</p>"
        f"<p>Proposal dates on or before {html.escape(cutoff_display)}.</p>"
        '<table border="1" cellspacing="0" cellpadding="6" style="border-collapse: collapse; font-size: 11pt;">'
        '<thead><tr style="background: #eef3f8;">'
        "<th align=\"left\">Customer</th>"
        "<th align=\"left\">Contact</th>"
        "<th align=\"left\">Email Address</th>"
        "<th align=\"left\">Proposal Date</th>"
        "</tr></thead>"
        f"<tbody>{''.join(row_html)}</tbody>"
        "</table>"
        "</body></html>"
    )
    return subject_text, plain_body, html_body


def get_weekly_follow_up_recipients_for_submitter(submitted_by):
    submitter = str(submitted_by or "").strip()
    recipients_by_submitter = {
        "David": ["david@procoatingsystems.com"],
        "Mark": ["mark@procoatingsystems.com"],
        "Lydia": ["lydia@procoatingsystems.com"],
        "Randy": ["randy@procoatingsystems.com"],
    }
    return recipients_by_submitter.get(submitter, get_email_recipients_for_submitter(submitter))


def get_weekly_follow_up_bcc_recipients():
    return ["mark@procoatingsystems.com"]


def get_weekly_follow_up_sender_email():
    return OUTLOOK_SENDER_EMAIL


def _open_outlook_html_draft_for_submitter(subject_text, plain_text_body, html_body, submitted_by, recipients=None):
    if sys.platform != "darwin":
        return None

    recipients = recipients or get_email_recipients_for_submitter(submitted_by)
    sender_email = get_weekly_follow_up_sender_email()
    sender_label = "Vern"
    bcc_recipients = get_weekly_follow_up_bcc_recipients()

    if _is_running_new_outlook():
        status = _open_new_outlook_template_draft(
            subject_text,
            plain_text_body,
            html_body,
            recipients,
            bcc_recipients,
            sender_email,
        )
        return _build_outlook_draft_warning(status, sender_email)

    recipient_blob = "||".join(recipients)
    bcc_recipient_blob = "||".join(bcc_recipients)
    account_match_enabled = "1"
    script_lines = [
        "on run argv",
        "set subjectText to item 1 of argv",
        "set htmlBody to item 2 of argv",
        "set senderEmail to item 3 of argv",
        "set recipientBlob to item 4 of argv",
        "set senderLabel to item 5 of argv",
        "set accountMatchFlag to item 6 of argv",
        "set bccRecipientBlob to item 7 of argv",
        'set AppleScript\'s text item delimiters to "||"',
        "set recipientList to text items of recipientBlob",
        "set bccRecipientList to text items of bccRecipientBlob",
        "set availableAccounts to {}",
        "set matchStatus to \"fallback:account-match-disabled\"",
        'tell application "Microsoft Outlook"',
        "activate",
        "set targetAccount to missing value",
        "set accountList to {}",
        'if accountMatchFlag is equal to "1" then',
        "try",
        "set accountList to accountList & (exchange accounts)",
        "end try",
        "try",
        "set accountList to accountList & (imap accounts)",
        "end try",
        "try",
        "set accountList to accountList & (pop accounts)",
        "end try",
        "end if",
        "repeat with acct in accountList",
        "try",
        "set acctEmail to email address of acct as string",
        "on error",
        "set acctEmail to \"\"",
        "end try",
        "try",
        "set acctName to name of acct as string",
        "on error",
        "set acctName to \"\"",
        "end try",
        "if acctName is not \"\" then",
        "set end of availableAccounts to acctName",
        "else if acctEmail is not \"\" then",
        "set end of availableAccounts to acctEmail",
        "end if",
        "if targetAccount is missing value then",
        "if acctEmail is not \"\" then",
        "ignoring case",
        "if acctEmail is equal to senderEmail then set targetAccount to acct",
        "end ignoring",
        "end if",
        "end if",
        "if targetAccount is missing value then",
        "if senderLabel is not \"\" then",
        "if acctName is not \"\" then",
        "ignoring case",
        "if acctName is equal to senderLabel then set targetAccount to acct",
        "end ignoring",
        "end if",
        "end if",
        "end if",
        "if targetAccount is missing value then",
        "if senderLabel is not \"\" then",
        "if acctName is not \"\" then",
        "ignoring case",
        "if acctName contains senderLabel then set targetAccount to acct",
        "end ignoring",
        "end if",
        "end if",
        "end if",
        "end repeat",
        "set newMessage to make new outgoing message with properties {subject:subjectText, content:htmlBody}",
        "if targetAccount is not missing value then",
        "try",
        "set account of newMessage to targetAccount",
        "set matchStatus to \"matched\"",
        "try",
        "set sender of newMessage to {address:senderEmail}",
        "end try",
        "on error",
        "set matchStatus to \"fallback:account-set-failed\"",
        "end try",
        "else",
        'if accountMatchFlag is equal to "1" then',
        "if (count of accountList) is 0 then",
        "set matchStatus to \"fallback:no-scriptable-accounts\"",
        "else",
        'set AppleScript\'s text item delimiters to ", "',
        "set availableAccountText to availableAccounts as string",
        'set AppleScript\'s text item delimiters to "||"',
        "set matchStatus to \"fallback:account-not-found:\" & availableAccountText",
        "end if",
        "else",
        "set matchStatus to \"fallback:account-match-disabled\"",
        "end if",
        "end if",
        "repeat with recipientAddress in recipientList",
        "set cleanAddress to (recipientAddress as string)",
        'if cleanAddress is not "" then',
        "make new to recipient at end of to recipients of newMessage with properties {email address:{address:cleanAddress}}",
        "end if",
        "end repeat",
        "repeat with recipientAddress in bccRecipientList",
        "set cleanAddress to (recipientAddress as string)",
        'if cleanAddress is not "" then',
        "make new bcc recipient at end of bcc recipients of newMessage with properties {email address:{address:cleanAddress}}",
        "end if",
        "end repeat",
        "open newMessage",
        "return matchStatus",
        "end tell",
        "end run",
    ]
    cmd = [
        "osascript",
        *sum((["-e", line] for line in script_lines), []),
        str(subject_text or "").strip(),
        str(html_body or ""),
        sender_email,
        recipient_blob,
        sender_label,
        account_match_enabled,
        bcc_recipient_blob,
    ]
    try:
        result = subprocess.run(
            cmd,
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
    except subprocess.CalledProcessError as exc:
        stderr_text = (exc.stderr or "").strip()
        details = stderr_text or str(exc)
        raise RuntimeError(f"Outlook draft creation failed: {details}") from exc
    return _build_outlook_draft_warning(result.stdout, sender_email)


@app.route('/weekly-follow-ups')
def weekly_follow_ups():
    cutoff_date = _parse_iso_date(
        request.args.get("cutoff_date"),
        _default_follow_up_cutoff_date(),
    )
    follow_ups, cutoff_date, tracker_error = load_weekly_follow_up_entries(cutoff_date=cutoff_date)
    return render_template(
        'weekly_follow_ups.html',
        follow_ups=follow_ups,
        cutoff_date_display=_format_tracker_date(cutoff_date),
        cutoff_date_value=cutoff_date.isoformat(),
        tracker_error=tracker_error,
    )


@app.route('/generate-weekly-follow-up-emails', methods=['POST'])
def generate_weekly_follow_up_emails():
    data = request.get_json(silent=True) or {}
    cutoff_date = _parse_iso_date(
        data.get("cutoff_date"),
        _default_follow_up_cutoff_date(),
    )
    selected_rows = {
        str(row_number).strip()
        for row_number in (data.get("selected_rows") or [])
        if re.fullmatch(r"[A-Za-z0-9_-]+", str(row_number).strip())
    }
    if not selected_rows:
        return jsonify({"error": "Select at least one follow-up row."}), 400

    follow_ups, cutoff_date, tracker_error = load_weekly_follow_up_entries(cutoff_date=cutoff_date)
    if tracker_error:
        return jsonify({"error": f"Unable to read proposal tracking data: {tracker_error}"}), 500

    selected_follow_ups = [
        item
        for item in follow_ups
        if str(item.get("row_number") or "").strip() in selected_rows
    ]
    if not selected_follow_ups:
        return jsonify({"error": "No selected follow-up rows are available for the current date filter."}), 400

    grouped = {}
    for item in selected_follow_ups:
        submitter = str(item.get("submitted_by") or "Unassigned").strip() or "Unassigned"
        grouped.setdefault(submitter, []).append(item)

    warnings = []
    try:
        for submitter in sorted(grouped, key=lambda value: value.casefold()):
            items = grouped[submitter]
            subject_text, plain_body, html_body = _build_weekly_follow_up_email_bodies(
                submitter,
                items,
                cutoff_date,
            )
            warning = _open_outlook_html_draft_for_submitter(
                subject_text,
                plain_body,
                html_body,
                submitter,
                get_weekly_follow_up_recipients_for_submitter(submitter),
            )
            if warning:
                warnings.append(warning)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    draft_count = len(grouped)
    row_count = len(selected_follow_ups)
    try:
        updated_count = update_weekly_follow_up_dates(
            [item["row_number"] for item in selected_follow_ups],
            datetime.date.today(),
        )
    except Exception as exc:
        return jsonify({
            "error": (
                "Outlook draft email(s) were created, but proposal tracking data "
                f"could not be updated: {exc}"
            )
        }), 500

    return jsonify({
        "message": (
            f"Created {draft_count} Outlook draft email(s) for {row_count} follow-up row(s). "
            f"Updated {updated_count} tracker row(s)."
        ),
        "draft_count": draft_count,
        "row_count": row_count,
        "updated_count": updated_count,
        "warnings": warnings,
    })


def _escape_applescript_string(value):
    return str(value).replace("\\", "\\\\").replace('"', '\\"')


def _selected_proposal_file_paths_from_form():
    raw_paths = (request.form.get("selected_proposal_file_paths") or "").strip()
    if not raw_paths:
        return []
    try:
        parsed_paths = json.loads(raw_paths)
    except Exception:
        return []
    if not isinstance(parsed_paths, list):
        return []

    selected_paths = []
    seen_paths = set()
    for path in parsed_paths:
        normalized_path = os.path.abspath(os.path.expanduser(str(path or "").strip()))
        if not normalized_path or normalized_path in seen_paths:
            continue
        seen_paths.add(normalized_path)
        selected_paths.append(normalized_path)
    return selected_paths


def _unique_destination_file_path(destination_folder, filename):
    base_name, extension = os.path.splitext(filename)
    candidate_path = os.path.join(destination_folder, filename)
    counter = 2
    while os.path.exists(candidate_path):
        candidate_path = os.path.join(destination_folder, f"{base_name} ({counter}){extension}")
        counter += 1
    return candidate_path


def move_selected_proposal_files_to_folder(selected_paths, proposal_folder):
    if not selected_paths or not proposal_folder:
        return []

    os.makedirs(proposal_folder, exist_ok=True)
    moved_files = []
    destination_folder_abs = os.path.abspath(proposal_folder)

    for source_path in selected_paths:
        source_path_abs = os.path.abspath(os.path.expanduser(str(source_path or "").strip()))
        if not source_path_abs:
            continue
        if not os.path.isfile(source_path_abs):
            _safe_debug(f"[WARN] Selected proposal file does not exist or is not a file: {source_path_abs}")
            continue
        source_parent_abs = os.path.abspath(os.path.dirname(source_path_abs))
        if source_parent_abs == destination_folder_abs:
            moved_files.append(source_path_abs)
            continue

        destination_path = _unique_destination_file_path(
            destination_folder_abs,
            os.path.basename(source_path_abs),
        )
        try:
            shutil.move(source_path_abs, destination_path)
            moved_files.append(destination_path)
        except Exception as exc:
            _safe_debug(f"[WARN] Could not move selected proposal file {source_path_abs}: {exc}")

    return moved_files


@app.route('/choose-email-template', methods=['POST'])
def choose_email_template():
    if sys.platform != "darwin":
        return jsonify({"error": "The native email template chooser is only available on macOS."}), 400

    if not os.path.isdir(EMAIL_TEMPLATE_DIR):
        return jsonify({"error": f"Email template folder not found: {EMAIL_TEMPLATE_DIR}"}), 404

    escaped_template_dir = _escape_applescript_string(EMAIL_TEMPLATE_DIR)
    script = "\n".join([
        f'set templateFolder to POSIX file "{escaped_template_dir}" as alias',
        'set pickedFile to choose file with prompt "Select email template" default location templateFolder',
        "return POSIX path of pickedFile",
    ])

    try:
        result = subprocess.run(
            ["osascript", "-e", script],
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
    except Exception as exc:
        return jsonify({"error": f"Unable to open email template chooser: {exc}"}), 500

    if result.returncode != 0:
        stderr = (result.stderr or "").strip()
        if "User canceled" in stderr:
            return jsonify({"cancelled": True})
        return jsonify({"error": stderr or "Email template chooser failed."}), 500

    selected_path = (result.stdout or "").strip()
    return jsonify({
        "path": selected_path,
        "name": os.path.basename(selected_path),
    })


@app.route('/choose-proposal-files', methods=['POST'])
def choose_proposal_files():
    if sys.platform != "darwin":
        return jsonify({"error": "The native proposal file chooser is only available on macOS."}), 400

    script = "\n".join([
        'set pickedFiles to choose file with prompt "Attach Files" with multiple selections allowed',
        'set selectedPaths to ""',
        'repeat with pickedFile in pickedFiles',
        'set selectedPaths to selectedPaths & POSIX path of pickedFile & linefeed',
        'end repeat',
        'return selectedPaths',
    ])

    try:
        result = subprocess.run(
            ["osascript", "-e", script],
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
    except Exception as exc:
        return jsonify({"error": f"Unable to open proposal file chooser: {exc}"}), 500

    if result.returncode != 0:
        stderr = (result.stderr or "").strip()
        if "User canceled" in stderr:
            return jsonify({"cancelled": True})
        return jsonify({"error": stderr or "Proposal file chooser failed."}), 500

    selected_paths = [
        path.strip()
        for path in (result.stdout or "").splitlines()
        if path.strip()
    ]
    return jsonify({
        "files": [
            {"path": path, "name": os.path.basename(path)}
            for path in selected_paths
        ]
    })


@app.route('/choose-distribution-list', methods=['POST'])
def choose_distribution_list():
    if sys.platform != "darwin":
        return jsonify({"error": "The native distribution list chooser is only available on macOS."}), 400

    if not os.path.isdir(EMAIL_LIST_DIR):
        return jsonify({"error": f"Email list folder not found: {EMAIL_LIST_DIR}"}), 404

    escaped_list_dir = _escape_applescript_string(EMAIL_LIST_DIR)
    script = "\n".join([
        f'set listFolder to POSIX file "{escaped_list_dir}" as alias',
        'set pickedFile to choose file with prompt "Select distribution list" default location listFolder of type {"public.comma-separated-values-text", "csv"}',
        "return POSIX path of pickedFile",
    ])

    try:
        result = subprocess.run(
            ["osascript", "-e", script],
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
    except Exception as exc:
        return jsonify({"error": f"Unable to open distribution list chooser: {exc}"}), 500

    if result.returncode != 0:
        stderr = (result.stderr or "").strip()
        if "User canceled" in stderr:
            return jsonify({"cancelled": True})
        return jsonify({"error": stderr or "Distribution list chooser failed."}), 500

    selected_path = (result.stdout or "").strip()
    if not selected_path.lower().endswith(".csv"):
        return jsonify({"error": "Please choose a .csv distribution list."}), 400

    return jsonify({
        "path": selected_path,
        "name": os.path.basename(selected_path),
    })


def _is_path_inside(path, parent_dir):
    try:
        return os.path.commonpath([
            os.path.realpath(path),
            os.path.realpath(parent_dir),
        ]) == os.path.realpath(parent_dir)
    except Exception:
        return False


EMAIL_ADDRESS_PATTERN = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.IGNORECASE)


def _extract_distribution_emails(csv_path):
    emails = []
    seen = set()

    def _add_addresses(cell):
        for match in EMAIL_ADDRESS_PATTERN.findall(str(cell or "")):
            address = match.strip()
            key = address.lower()
            if key not in seen:
                emails.append(address)
                seen.add(key)

    try:
        with open(csv_path, newline="", encoding="utf-8-sig") as handle:
            reader = csv.reader(handle)
            for row in reader:
                for cell in row:
                    _add_addresses(cell)
    except UnicodeDecodeError:
        with open(csv_path, newline="", encoding="latin-1") as handle:
            reader = csv.reader(handle)
            for row in reader:
                for cell in row:
                    _add_addresses(cell)
    return emails


def _chunked(items, chunk_size):
    for start in range(0, len(items), chunk_size):
        yield items[start:start + chunk_size]


def _open_blast_email_template_draft(template_path, bcc_recipients, batch_number, total_batches):
    try:
        with open(template_path, "rb") as handle:
            message = BytesParser(policy=policy.default).parse(handle)
    except Exception as exc:
        raise RuntimeError(f"Unable to read email template: {exc}") from exc

    _stamp_and_verify_new_outlook_sender(message, OUTLOOK_SENDER_EMAIL)
    _replace_message_header(message, "Bcc", ", ".join(bcc_recipients))
    for stale_header in ("To", "Cc", "Date", "Message-ID", "Thread-Index", "X-MS-TNEF-Correlator"):
        if stale_header in message:
            del message[stale_header]

    draft_path = os.path.join(
        tempfile.gettempdir(),
        f"pcs-blast-email-{batch_number}-of-{total_batches}-{uuid.uuid4().hex}.emltpl",
    )
    with open(draft_path, "wb") as handle:
        handle.write(message.as_bytes(policy=policy.SMTP))

    try:
        subprocess.run(
            ["open", "-a", "Microsoft Outlook", draft_path],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
    except subprocess.CalledProcessError as exc:
        stderr_text = (exc.stderr or "").strip()
        details = stderr_text or str(exc)
        raise RuntimeError(f"Outlook draft creation failed: {details}") from exc


@app.route('/generate-blast-emails', methods=['POST'])
def generate_blast_emails():
    data = request.get_json(silent=True) or {}
    template_path = str(data.get("email_template_path") or "").strip()
    distribution_list_path = str(data.get("distribution_list_path") or "").strip()

    if not template_path:
        return jsonify({"error": "Choose an email template first."}), 400
    if not distribution_list_path:
        return jsonify({"error": "Choose a distribution list first."}), 400
    if not os.path.isfile(template_path) or not _is_path_inside(template_path, EMAIL_TEMPLATE_DIR):
        return jsonify({"error": "Selected email template is not available."}), 400
    if (
        not os.path.isfile(distribution_list_path)
        or not _is_path_inside(distribution_list_path, EMAIL_LIST_DIR)
        or not distribution_list_path.lower().endswith(".csv")
    ):
        return jsonify({"error": "Selected distribution list must be a .csv file from the Email Lists folder."}), 400

    emails = _extract_distribution_emails(distribution_list_path)
    if not emails:
        return jsonify({"error": "No valid email addresses were found in the distribution list."}), 400

    batches = list(_chunked(emails, 40))
    try:
        for index, batch in enumerate(batches, start=1):
            _open_blast_email_template_draft(template_path, batch, index, len(batches))
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

    return jsonify({
        "message": f"Created {len(batches)} Outlook draft email(s) for {len(emails)} recipient(s).",
        "draft_count": len(batches),
        "recipient_count": len(emails),
    })


def get_folder_last_modified(folder_path):
    latest_mtime = None

    try:
        latest_mtime = os.path.getmtime(folder_path)
    except OSError:
        return None

    for root, dirs, files in os.walk(folder_path):
        dirs[:] = [name for name in dirs if name.lower() != "archive"]
        for name in dirs + files:
            path = os.path.join(root, name)
            try:
                path_mtime = os.path.getmtime(path)
            except OSError:
                continue
            if latest_mtime is None or path_mtime > latest_mtime:
                latest_mtime = path_mtime

    if latest_mtime is None:
        return None

    return datetime.datetime.fromtimestamp(latest_mtime)


def _is_proposal_list_folder(folder_name):
    normalized = str(folder_name or "").strip().lower()
    return bool(normalized) and normalized != "archive"


def _as_dir_list(root_dirs):
    if isinstance(root_dirs, (str, bytes, os.PathLike)):
        return [root_dirs]
    return list(root_dirs or [])


def resolve_open_proposal_folder(folder_name):
    safe_folder = os.path.basename(str(folder_name or ""))
    if not safe_folder:
        return None

    for root_dir in OPEN_PROPOSAL_DIRS:
        folder_path = os.path.join(root_dir, safe_folder)
        if os.path.isdir(folder_path):
            return folder_path

    temp_folder_path = os.path.join(PROPOSAL_TEMP_DIR, safe_folder)
    if os.path.isdir(temp_folder_path):
        return temp_folder_path

    return None


def build_proposal_entries(root_dirs, recent_cutoff):
    entries = []

    for root_dir in _as_dir_list(root_dirs):
        try:
            folder_names = [
                f for f in os.listdir(root_dir)
                if os.path.isdir(os.path.join(root_dir, f)) and _is_proposal_list_folder(f)
            ]
        except Exception:
            continue

        for folder_name in folder_names:
            folder_path = os.path.join(root_dir, folder_name)
            last_modified = get_folder_last_modified(folder_path)
            entries.append(
                {
                    "name": folder_name,
                    "last_modified": last_modified,
                    "last_modified_display": (
                        last_modified.strftime("%m/%d/%Y") if last_modified else "Unavailable"
                    ),
                    "is_recent": bool(last_modified and last_modified >= recent_cutoff),
                }
            )

    return sorted(entries, key=lambda proposal: proposal["name"].lower())


def recent_proposals(proposals):
    return sorted(
        [proposal for proposal in proposals if proposal["is_recent"]],
        key=lambda proposal: (
            proposal["last_modified"] or datetime.datetime.min,
            proposal["name"].lower(),
        ),
        reverse=True,
    )

def compute_total_squares(flat_roof_squares, wall_squares, fallback_total=None, raw_flat=None, raw_wall=None):
    raw_flat_blank = raw_flat is None or str(raw_flat).strip() == ""
    raw_wall_blank = raw_wall is None or str(raw_wall).strip() == ""
    if raw_flat_blank and raw_wall_blank and fallback_total is not None:
        return fallback_total
    return (flat_roof_squares or 0.0) + (wall_squares or 0.0)

def get_copy_destination_for_submitter(submitted_by):
    submitter = str(submitted_by or "").strip()
    if submitter in ("Mark", "Richard", "Vern"):
        return PCS_PROPOSALS_DIR
    if submitter == "David":
        return DAVIDS_PROPOSALS_DIR
    if submitter == "Lydia":
        return LYDIAS_PROPOSALS_DIR
    if submitter == "Randy":
        return RANDYS_PROPOSALS_DIR
    return None


def get_submitter_proposal_folder(folder_name, submitted_by):
    destination_root = get_copy_destination_for_submitter(submitted_by)
    safe_folder = os.path.basename(str(folder_name or ""))
    if not destination_root or not safe_folder:
        return None
    return os.path.join(destination_root, safe_folder)


def get_email_recipients_for_submitter(submitted_by):
    recipients = ["mark@procoatingsystems.com"]
    submitter = str(submitted_by or "").strip()
    if submitter == "David":
        recipients.append("david@procoatingsystems.com")
    elif submitter == "Lydia":
        recipients.append("lydia@procoatingsystems.com")
    return recipients

def get_sender_email_for_submitter(submitted_by):
    return OUTLOOK_SENDER_EMAIL

def get_copy_destination_web_url_for_submitter(submitted_by):
    submitter = str(submitted_by or "").strip()
    if submitter in ("Mark", "Richard", "Vern"):
        return PCS_PROPOSALS_WEB_URL
    if submitter == "David":
        return DAVIDS_PROPOSALS_WEB_URL
    if submitter == "Lydia":
        return LYDIAS_PROPOSALS_WEB_URL
    if submitter == "Randy":
        return RANDYS_PROPOSALS_WEB_URL
    return ""

def _join_url_path(base_url, *segments):
    base = str(base_url or "").strip()
    if not base:
        return ""

    parsed = urlsplit(base)
    cleaned_path = parsed.path.rstrip("/")
    encoded_segments = [
        quote(str(segment or "").strip("/"), safe="")
        for segment in segments
        if str(segment or "").strip("/")
    ]
    joined_path = "/".join(part for part in [cleaned_path, *encoded_segments] if part)
    return urlunsplit((parsed.scheme, parsed.netloc, joined_path, parsed.query, parsed.fragment))

def build_proposal_folder_link(folder_path, submitted_by=None, folder_name=None):
    destination_web_url = get_copy_destination_web_url_for_submitter(submitted_by)
    destination_root = get_copy_destination_for_submitter(submitted_by)
    normalized_folder_path = os.path.normpath(str(folder_path or "").strip()) if folder_path else ""

    if destination_root and folder_name:
        should_use_submitter_path = not normalized_folder_path
        if normalized_folder_path:
            try:
                relative_path = os.path.relpath(normalized_folder_path, destination_root)
                should_use_submitter_path = relative_path.startswith("..")
            except Exception:
                should_use_submitter_path = True
        if should_use_submitter_path:
            normalized_folder_path = os.path.normpath(os.path.join(destination_root, folder_name))

    if destination_web_url:
        if normalized_folder_path and destination_root:
            try:
                relative_path = os.path.relpath(normalized_folder_path, destination_root)
                if relative_path != "." and not relative_path.startswith(".."):
                    return _join_url_path(destination_web_url, *relative_path.split(os.sep))
                return destination_web_url.rstrip("/")
            except Exception:
                pass
        if folder_name:
            return _join_url_path(destination_web_url, folder_name)

    if normalized_folder_path:
        try:
            return pathlib.Path(normalized_folder_path).resolve().as_uri()
        except Exception:
            return ""
    return ""

def build_proposal_email_subject(customer_name, street_address):
    return f"{str(customer_name or '').strip()} {str(street_address or '').strip()}".strip()

@lru_cache(maxsize=1)
def _load_proposal_summary_email_template():
    try:
        with open(PROPOSAL_SUMMARY_TEMPLATE_PATH, "rb") as handle:
            message = BytesParser(policy=policy.default).parse(handle)
    except Exception as exc:
        raise RuntimeError(f"Unable to read proposal summary email template: {exc}") from exc

    plain_template = ""
    html_template = ""
    for part in message.walk():
        if part.get_content_maintype() == "multipart":
            continue
        if part.get_content_type() == "text/plain" and not plain_template:
            plain_template = str(part.get_content() or "")
        elif part.get_content_type() == "text/html" and not html_template:
            html_template = str(part.get_content() or "")

    if not plain_template or not html_template:
        raise RuntimeError("Proposal summary email template must contain both text/plain and text/html bodies.")

    return plain_template, html_template

def _insert_proposal_summary_extras_plain(body_text, proposal_note, proposal_language):
    note_text = str(proposal_note or "").strip()
    language_text = str(proposal_language or "").strip()
    if not note_text and not language_text:
        return body_text

    newline = "\r\n" if "\r\n" in body_text else "\n"
    profit_label = "Daily 10-year profit - "
    start = body_text.find(profit_label)
    if start == -1:
        return body_text

    line_end = body_text.find(newline, start)
    if line_end == -1:
        line_end = len(body_text)

    extra_lines = []
    if note_text:
        extra_lines.extend(["", note_text])
    if language_text:
        extra_lines.append(language_text)
    inserted_line = f"{newline}{newline.join(extra_lines)}"
    return f"{body_text[:line_end]}{inserted_line}{body_text[line_end:]}"

def _insert_proposal_summary_extras_html(body_html, proposal_note, proposal_language):
    note_text = str(proposal_note or "").strip()
    language_text = str(proposal_language or "").strip()
    if not note_text and not language_text:
        return body_html

    insert_at = body_html.rfind("</ul>")
    if insert_at == -1:
        return body_html

    extra_blocks = []
    if note_text:
        note_html = html.escape(note_text).replace("\n", "<br>\n")
        extra_blocks.extend([
            '<div style="direction: ltr; text-align: left; text-indent: 0px; font-family: Aptos, Arial, Helvetica, sans-serif; font-size: 12pt; color: rgb(33, 33, 33);">',
            "<br>",
            "</div>",
            '<div style="direction: ltr; text-align: left; text-indent: 0px; font-family: Aptos, Arial, Helvetica, sans-serif; font-size: 12pt; color: rgb(192, 0, 0);"><strong>',
            note_html,
            "</strong></div>",
        ])
    if language_text:
        language_html = html.escape(language_text).replace("\n", "<br>\n")
        extra_blocks.extend([
            '<div style="direction: ltr; text-align: left; text-indent: 0px; font-family: Aptos, Arial, Helvetica, sans-serif; font-size: 12pt; color: rgb(0, 0, 0);">',
            language_html,
            "</div>",
        ])
    return f"{body_html[:insert_at]}</ul>{''.join(extra_blocks)}{body_html[insert_at + len('</ul>'):]}"

def _format_folder_link_html(folder_name, folder_link):
    folder_text = html.escape(str(folder_name or "Proposal Folder"))
    link = str(folder_link or "").strip()
    if not link:
        return folder_text
    return f'<a href="{html.escape(link, quote=True)}">{folder_text}</a>'

def _build_proposal_summary_email_bodies(customer_name,
                                         street_address,
                                         total_squares,
                                         flat_roof_squares,
                                         wall_squares,
                                         roof_type,
                                         daily_profit,
                                         proposal_note,
                                         proposal_language,
                                         folder_link=""):
    plain_template, html_template = _load_proposal_summary_email_template()
    folder_name = build_proposal_email_subject(customer_name, street_address) or "Proposal Folder"
    replacements = {
        "FolderName": folder_name,
        "TotalSquares": _format_square_count(total_squares),
        "RoofType": str(roof_type or "").strip(),
        "10YrProfit": _format_currency(daily_profit),
    }

    plain_body = plain_template
    html_body = html_template
    for placeholder, value in replacements.items():
        plain_body = plain_body.replace(placeholder, value)
        if placeholder == "FolderName":
            html_body = html_body.replace(placeholder, _format_folder_link_html(value, folder_link))
        else:
            html_body = html_body.replace(placeholder, html.escape(value))

    flat_roof_text = _format_square_count(flat_roof_squares)
    wall_text = _format_square_count(wall_squares)
    plain_body = plain_body.replace("XXX", flat_roof_text, 1).replace("XXX", wall_text, 1)
    html_body = html_body.replace("XXX", html.escape(flat_roof_text), 1).replace("XXX", html.escape(wall_text), 1)

    plain_body = _insert_proposal_summary_extras_plain(plain_body, proposal_note, proposal_language)
    html_body = _insert_proposal_summary_extras_html(html_body, proposal_note, proposal_language)
    return plain_body, html_body

def build_proposal_summary_email_html(customer_name,
                                      street_address,
                                      folder_link,
                                      total_squares,
                                      flat_roof_squares,
                                      wall_squares,
                                      roof_type,
                                      daily_profit,
                                      proposal_note,
                                      proposal_language):
    _, html_body = _build_proposal_summary_email_bodies(
        customer_name=customer_name,
        street_address=street_address,
        total_squares=total_squares,
        flat_roof_squares=flat_roof_squares,
        wall_squares=wall_squares,
        roof_type=roof_type,
        daily_profit=daily_profit,
        proposal_note=proposal_note,
        proposal_language=proposal_language,
        folder_link=folder_link,
    )
    return html_body

def build_proposal_summary_email_text(customer_name,
                                      street_address,
                                      folder_link,
                                      total_squares,
                                      flat_roof_squares,
                                      wall_squares,
                                      roof_type,
                                      daily_profit,
                                      proposal_note,
                                      proposal_language):
    plain_body, _ = _build_proposal_summary_email_bodies(
        customer_name=customer_name,
        street_address=street_address,
        total_squares=total_squares,
        flat_roof_squares=flat_roof_squares,
        wall_squares=wall_squares,
        roof_type=roof_type,
        daily_profit=daily_profit,
        proposal_note=proposal_note,
        proposal_language=proposal_language,
        folder_link=folder_link,
    )
    return plain_body

def _replace_message_header(message, header_name, value):
    if header_name in message:
        message.replace_header(header_name, value)
    else:
        message[header_name] = value

def _set_text_message_part(part, body_text, subtype):
    payload = base64.b64encode(str(body_text or "").encode("utf-8")).decode("ascii")
    part.set_type(f"text/{subtype}")
    part.set_param("charset", "utf-8", replace=True)
    if "Content-Transfer-Encoding" in part:
        part.replace_header("Content-Transfer-Encoding", "base64")
    else:
        part["Content-Transfer-Encoding"] = "base64"
    part.set_payload(payload)

def _stamp_and_verify_new_outlook_sender(message, sender_email):
    sender = str(sender_email or "").strip().lower()
    if not sender:
        return

    # New Outlook can inherit stale MAPI identity metadata from an .emltpl file
    # even after its visible From header is replaced. Remove that metadata and
    # stamp every standard identity header before the draft is opened.
    for header_name in list(message.keys()):
        lowered = header_name.lower()
        if lowered.startswith("x-ms-exchange-") or lowered == "x-ms-tnef-correlator":
            del message[header_name]
    for header_name in ("From", "Sender", "Reply-To"):
        _replace_message_header(message, header_name, sender)
    _replace_message_header(message, "X-Unsent", "1")

    for header_name in ("From", "Sender", "Reply-To"):
        actual = parseaddr(str(message.get(header_name) or ""))[1].strip().lower()
        if actual != sender:
            raise RuntimeError(
                f"Outlook draft sender verification failed for {header_name}: expected {sender}."
            )

def _open_new_outlook_template_draft(
    subject_text,
    plain_text_body,
    html_body,
    recipients,
    bcc_recipients=None,
    sender_email=None,
):
    recipient_text = ", ".join(
        str(address or "").strip()
        for address in (recipients or [])
        if str(address or "").strip()
    )
    bcc_recipient_text = ", ".join(
        str(address or "").strip()
        for address in (bcc_recipients or [])
        if str(address or "").strip()
    )
    try:
        with open(PROPOSAL_SUMMARY_TEMPLATE_PATH, "rb") as handle:
            message = BytesParser(policy=policy.default).parse(handle)
    except Exception as exc:
        raise RuntimeError(f"Unable to read proposal summary email template: {exc}") from exc

    _replace_message_header(message, "To", recipient_text)
    if sender_email:
        _stamp_and_verify_new_outlook_sender(message, sender_email)
    if bcc_recipient_text:
        _replace_message_header(message, "Bcc", bcc_recipient_text)
    _replace_message_header(message, "Subject", str(subject_text or "").strip())

    for stale_header in ("Date", "Message-ID", "Thread-Index", "X-MS-TNEF-Correlator"):
        if stale_header in message:
            del message[stale_header]

    found_plain = False
    found_html = False
    for part in message.walk():
        if part.get_content_maintype() == "multipart":
            continue
        if part.get_content_type() == "text/plain":
            _set_text_message_part(part, plain_text_body, "plain")
            found_plain = True
        elif part.get_content_type() == "text/html":
            _set_text_message_part(part, html_body, "html")
            found_html = True

    if not found_plain or not found_html:
        raise RuntimeError("Proposal summary email template must contain both text/plain and text/html bodies.")

    draft_dir = tempfile.gettempdir()
    draft_path = os.path.join(draft_dir, f"pcs-proposal-summary-{uuid.uuid4().hex}.emltpl")
    with open(draft_path, "wb") as handle:
        handle.write(message.as_bytes(policy=policy.SMTP))
    try:
        subprocess.run(
            ["open", "-a", "Microsoft Outlook", draft_path],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
    except subprocess.CalledProcessError as exc:
        stderr_text = (exc.stderr or "").strip()
        details = stderr_text or str(exc)
        raise RuntimeError(f"Outlook draft creation failed: {details}") from exc
    return "fallback:new-outlook-template"

def _build_outlook_draft_warning(status_text, sender_email):
    status = str(status_text or "").strip()
    if not status or status == "matched":
        return None

    verify_message = (
        f"Outlook draft was created, but verify the From address is {sender_email} before sending."
    )
    if status == "fallback:account-match-disabled":
        return (
            "Outlook draft was created, but automatic sender selection is unavailable in this Outlook mode. "
            f"Verify the From address is {sender_email} before sending."
        )
    if status == "fallback:new-outlook-mailto":
        return None
    if status == "fallback:new-outlook-template":
        return None
    if status == "fallback:no-scriptable-accounts":
        return (
            "Outlook draft was created, but Outlook did not expose any scriptable mail accounts. "
            f"Verify the From address is {sender_email} before sending."
        )
    if status == "fallback:account-set-failed":
        return (
            "Outlook draft was created, but Outlook would not apply the requested sender account automatically. "
            f"Verify the From address is {sender_email} before sending."
        )
    if status.startswith("fallback:account-not-found:"):
        available_accounts = status.split(":", 2)[2].strip()
        if available_accounts:
            return (
                f"Outlook draft was created, but no scriptable Outlook account matched {sender_email}. "
                f"Outlook exposed: {available_accounts}. Verify the From address before sending."
            )
        return verify_message
    return verify_message

def create_outlook_proposal_summary_draft(customer_name,
                                          street_address,
                                          submitted_by,
                                          total_squares,
                                          flat_roof_squares,
                                          wall_squares,
                                          roof_type,
                                          daily_profit,
                                          proposal_note,
                                          proposal_language,
                                          folder_name,
                                          folder_link=None):
    if sys.platform != "darwin":
        return

    destination_root = get_copy_destination_for_submitter(submitted_by)
    folder_path = os.path.join(destination_root, folder_name) if destination_root else ""
    if folder_link is None:
        folder_link = build_proposal_folder_link(
            folder_path=folder_path,
            submitted_by=submitted_by,
            folder_name=folder_name,
        )
    subject_text = build_proposal_email_subject(customer_name, street_address)
    html_body = build_proposal_summary_email_html(
        customer_name=customer_name,
        street_address=street_address,
        folder_link=folder_link,
        total_squares=total_squares,
        flat_roof_squares=flat_roof_squares,
        wall_squares=wall_squares,
        roof_type=roof_type,
        daily_profit=daily_profit,
        proposal_note=proposal_note,
        proposal_language=proposal_language,
    )
    recipients = get_email_recipients_for_submitter(submitted_by)
    recipients_blob = "||".join(recipients)

    sender_email = get_sender_email_for_submitter(submitted_by)
    sender_label = "Vern"
    if _is_running_new_outlook():
        plain_text_body = build_proposal_summary_email_text(
            customer_name=customer_name,
            street_address=street_address,
            folder_link=folder_link,
            total_squares=total_squares,
            flat_roof_squares=flat_roof_squares,
            wall_squares=wall_squares,
            roof_type=roof_type,
            daily_profit=daily_profit,
            proposal_note=proposal_note,
            proposal_language=proposal_language,
        )
        status = _open_new_outlook_template_draft(
            subject_text,
            plain_text_body,
            html_body,
            recipients,
            sender_email=sender_email,
        )
        return _build_outlook_draft_warning(status, sender_email)

    account_match_enabled = "1"

    script_lines = [
        "on run argv",
        "set subjectText to item 1 of argv",
        "set htmlBody to item 2 of argv",
        "set senderEmail to item 3 of argv",
        "set recipientBlob to item 4 of argv",
        "set senderLabel to item 5 of argv",
        "set accountMatchFlag to item 6 of argv",
        'set AppleScript\'s text item delimiters to "||"',
        "set recipientList to text items of recipientBlob",
        "set availableAccounts to {}",
        "set matchStatus to \"fallback:account-match-disabled\"",
        'tell application "Microsoft Outlook"',
        "activate",
        "set targetAccount to missing value",
        "set accountList to {}",
        'if accountMatchFlag is equal to "1" then',
        "try",
        "set accountList to accountList & (exchange accounts)",
        "end try",
        "try",
        "set accountList to accountList & (imap accounts)",
        "end try",
        "try",
        "set accountList to accountList & (pop accounts)",
        "end try",
        "end if",
        "repeat with acct in accountList",
        "try",
        "set acctEmail to email address of acct as string",
        "on error",
        "set acctEmail to \"\"",
        "end try",
        "try",
        "set acctName to name of acct as string",
        "on error",
        "set acctName to \"\"",
        "end try",
        "if acctName is not \"\" then",
        "set end of availableAccounts to acctName",
        "else if acctEmail is not \"\" then",
        "set end of availableAccounts to acctEmail",
        "end if",
        "if targetAccount is missing value then",
        "if acctEmail is not \"\" then",
        "ignoring case",
        "if acctEmail is equal to senderEmail then set targetAccount to acct",
        "end ignoring",
        "end if",
        "end if",
        "if targetAccount is missing value then",
        "if senderLabel is not \"\" then",
        "if acctName is not \"\" then",
        "ignoring case",
        "if acctName is equal to senderLabel then set targetAccount to acct",
        "end ignoring",
        "end if",
        "end if",
        "end if",
        "if targetAccount is missing value then",
        "if senderLabel is not \"\" then",
        "if acctName is not \"\" then",
        "ignoring case",
        "if acctName contains senderLabel then set targetAccount to acct",
        "end ignoring",
        "end if",
        "end if",
        "end if",
        "end repeat",
        "set newMessage to make new outgoing message with properties {subject:subjectText, content:htmlBody}",
        "if targetAccount is not missing value then",
        "try",
        "set account of newMessage to targetAccount",
        "set matchStatus to \"matched\"",
        "try",
        "set sender of newMessage to {address:senderEmail}",
        "end try",
        "on error",
        "set matchStatus to \"fallback:account-set-failed\"",
        "end try",
        "else",
        'if accountMatchFlag is equal to "1" then',
        "if (count of accountList) is 0 then",
        "set matchStatus to \"fallback:no-scriptable-accounts\"",
        "else",
        'set AppleScript\'s text item delimiters to ", "',
        "set availableAccountText to availableAccounts as string",
        'set AppleScript\'s text item delimiters to "||"',
        "set matchStatus to \"fallback:account-not-found:\" & availableAccountText",
        "end if",
        "else",
        "set matchStatus to \"fallback:account-match-disabled\"",
        "end if",
        "end if",
        "repeat with recipientAddress in recipientList",
        "set cleanAddress to (recipientAddress as string)",
        'if cleanAddress is not "" then',
        "make new to recipient at end of to recipients of newMessage with properties {email address:{address:cleanAddress}}",
        "end if",
        "end repeat",
        "open newMessage",
        "return matchStatus",
        "end tell",
        "end run",
    ]

    cmd = [
        "osascript",
        *sum((["-e", line] for line in script_lines), []),
        subject_text,
        html_body,
        sender_email,
        recipients_blob,
        sender_label,
        account_match_enabled,
    ]
    try:
        result = subprocess.run(
            cmd,
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
    except subprocess.CalledProcessError as exc:
        stderr_text = (exc.stderr or "").strip()
        details = stderr_text or str(exc)
        raise RuntimeError(f"Outlook draft creation failed: {details}") from exc
    return _build_outlook_draft_warning(result.stdout, sender_email)

def _sync_directory_contents(source_folder, dest_folder):
    try:
        if os.path.realpath(source_folder) == os.path.realpath(dest_folder):
            return
    except Exception:
        pass

    os.makedirs(dest_folder, exist_ok=True)

    try:
        source_entries = {entry.name: entry for entry in os.scandir(source_folder)}
    except FileNotFoundError:
        return

    try:
        dest_entries = {entry.name: entry for entry in os.scandir(dest_folder)}
    except FileNotFoundError:
        dest_entries = {}

    for name, dest_entry in dest_entries.items():
        if name in source_entries:
            continue
        if dest_entry.is_file() and not _is_generated_proposal_artifact(name):
            continue
        dest_path = dest_entry.path
        try:
            if dest_entry.is_dir(follow_symlinks=False):
                shutil.rmtree(dest_path)
            else:
                os.remove(dest_path)
        except Exception as exc:
            _safe_debug(f"[WARN] Could not remove stale destination item {dest_path}: {exc}")

    for name, source_entry in source_entries.items():
        source_path = source_entry.path
        dest_path = os.path.join(dest_folder, name)
        if source_entry.is_dir(follow_symlinks=False):
            _sync_directory_contents(source_path, dest_path)
            continue
        if not _is_generated_proposal_artifact(name):
            continue

        copy_required = True
        if os.path.exists(dest_path) and os.path.isfile(dest_path):
            try:
                src_stat = os.stat(source_path)
                dest_stat = os.stat(dest_path)
                copy_required = (
                    src_stat.st_size != dest_stat.st_size
                    or src_stat.st_mtime_ns != dest_stat.st_mtime_ns
                )
            except Exception:
                copy_required = True

        if copy_required:
            shutil.copy2(source_path, dest_path)


def copy_proposal_to_temp_dir(source_folder, folder_name):
    if not source_folder or not folder_name:
        return

    temp_folder = os.path.join(PROPOSAL_TEMP_DIR, folder_name)
    _sync_directory_contents(source_folder, temp_folder)


def _is_pdf_ready_for_copy(
    pdf_path: str,
    previous_snapshot: tuple[int, int] | None = None,
) -> tuple[bool, tuple[int, int] | None]:
    try:
        stat_result = os.stat(pdf_path)
    except FileNotFoundError:
        return False, None
    except Exception as exc:
        _safe_debug(f"[WARN] Could not stat generated PDF {pdf_path}: {exc}")
        return False, None

    try:
        with open(pdf_path, "rb") as handle:
            header = handle.read(5)
    except Exception as exc:
        _safe_debug(f"[WARN] Could not read generated PDF {pdf_path}: {exc}")
        return False, None

    if header != b"%PDF-":
        return False, (int(stat_result.st_size), int(stat_result.st_mtime_ns))

    snapshot = (int(stat_result.st_size), int(stat_result.st_mtime_ns))
    if previous_snapshot is None:
        return False, snapshot

    return snapshot == previous_snapshot, snapshot


def _wait_for_generated_pdfs(
    source_folder: str,
    timeout_seconds: float = 45.0,
    stable_checks_required: int = 2,
) -> None:
    try:
        docx_names = [
            entry.name
            for entry in os.scandir(source_folder)
            if entry.is_file() and entry.name.lower().endswith(".docx")
        ]
    except FileNotFoundError:
        return
    except Exception as exc:
        _safe_debug(f"[WARN] Could not scan for generated PDFs in {source_folder}: {exc}")
        return

    expected_pdfs = {
        os.path.join(source_folder, f"{os.path.splitext(name)[0]}.pdf")
        for name in docx_names
    }
    if not expected_pdfs:
        return

    previous_snapshots = {pdf_path: None for pdf_path in expected_pdfs}
    stable_counts = {pdf_path: 0 for pdf_path in expected_pdfs}
    deadline = time.perf_counter() + max(timeout_seconds, 0.0)
    while time.perf_counter() < deadline:
        all_ready = True
        for pdf_path in expected_pdfs:
            is_ready, snapshot = _is_pdf_ready_for_copy(
                pdf_path,
                previous_snapshots[pdf_path],
            )
            previous_snapshots[pdf_path] = snapshot
            if is_ready:
                stable_counts[pdf_path] += 1
            else:
                stable_counts[pdf_path] = 0
                all_ready = False
        if all_ready and all(
            stable_counts[pdf_path] >= stable_checks_required
            for pdf_path in expected_pdfs
        ):
            return
        time.sleep(0.25)

    missing = [pdf_path for pdf_path in expected_pdfs if not os.path.exists(pdf_path)]
    if missing:
        _safe_debug(
            f"[WARN] Timed out waiting for generated PDFs before destination sync: {missing}"
        )
        return

    incomplete = [
        pdf_path
        for pdf_path in expected_pdfs
        if stable_counts[pdf_path] < stable_checks_required
    ]
    if incomplete:
        _safe_debug(
            f"[WARN] Timed out waiting for stable generated PDFs before destination sync: {incomplete}"
        )


def copy_proposal_to_submitter_destination(
    source_folder,
    folder_name,
    submitted_by,
    wait_for_pdfs: bool = False,
):
    destination_root = get_copy_destination_for_submitter(submitted_by)
    if not destination_root:
        return

    if wait_for_pdfs:
        _wait_for_generated_pdfs(source_folder)

    dest_folder = os.path.join(destination_root, folder_name)
    _sync_directory_contents(source_folder, dest_folder)

def update_existing_tracker_row(
    folder_name,
    lead_value,
    submitted_by,
    estimate_completed_date=None,
    tracker_path=PROPOSAL_TRACKER,
):
    with TRACKER_IO_LOCK:
        source_path = _proposal_tracker_source_path(tracker_path)
        if not tracker_path or not source_path:
            return False

        wb = None
        temp_path = None
        try:
            wb = load_workbook(source_path)
            ws = wb.active
            _ensure_proposal_tracker_status_column(ws)
            columns = _proposal_tracker_column_map(ws)
            row_key = str(folder_name or "").strip().lower()
            if not row_key:
                return False

            for row in range(2, ws.max_row + 1):
                existing_key = str(
                    ws.cell(row=row, column=columns["customer"]).value or ""
                ).strip().lower()
                if existing_key != row_key:
                    continue
                ws.cell(row=row, column=columns["customer"]).value = folder_name
                ws.cell(row=row, column=columns["lead_source"]).value = lead_value or ""
                ws.cell(row=row, column=columns["submitted_by"]).value = submitted_by or ""
                ws.cell(row=row, column=columns["estimated_by"]).value = "Vern"
                status_cell = ws.cell(row=row, column=columns["status"])
                if not str(status_cell.value or "").strip():
                    status_cell.value = "Draft"
                estimate_cell = ws.cell(row=row, column=columns["estimate_date"])
                if not str(estimate_cell.value or "").strip():
                    estimate_cell.value = (
                        estimate_completed_date or datetime.date.today()
                    )
                    estimate_cell.number_format = "m/d/yyyy"
                temp_path = _proposal_tracker_temp_path(tracker_path)
                wb.save(temp_path)
                wb.close()
                wb = None
                _replace_proposal_tracker_file(temp_path, tracker_path)
                temp_path = None
                return True
            return False
        finally:
            try:
                if wb is not None:
                    wb.close()
            except Exception:
                pass
            if temp_path and os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except Exception:
                    pass

def _update_tracking_after_save_spreadsheet(folder_name,
                               customer_name,
                               street_address,
                               city,
                               state,
                               zip_code,
                               product,
                               roof_type,
                               total_squares,
                               warranty_incl,
                               submitted_by,
                               proposal_folder,
                               tp10,
                               tp15,
                               tp20,
                               lead_value):
    with TRACKER_IO_LOCK:
        if update_existing_tracker_row(
            folder_name,
            lead_value,
            submitted_by,
            datetime.date.today(),
        ):
            return

        _append_to_proposal_tracking_unlocked(
            created_date=datetime.date.today().strftime("%m/%d/%Y"),
            customer_name=customer_name,
            street_address=street_address,
            city=city,
            state=state,
            zip_code=zip_code,
            product=product,
            roof_type=roof_type,
            total_squares=total_squares,
            warranty_incl=warranty_incl,
            submitted_by=submitted_by,
            folder_name=folder_name,
            proposal_folder=proposal_folder,
            tp10=tp10,
            tp15=tp15,
            tp20=tp20,
            lead_value=lead_value,
        )


def update_tracking_after_save(folder_name,
                               customer_name,
                               street_address,
                               city,
                               state,
                               zip_code,
                               product,
                               roof_type,
                               total_squares,
                               warranty_incl,
                               submitted_by,
                               proposal_folder,
                               tp10,
                               tp15,
                               tp20,
                               lead_value):
    flags = load_proposal_tracking_cutover_flags()
    supabase_error = None
    if flags.writes_enabled:
        try:
            get_proposal_tracking_store().upsert_from_proposal_save(
                created_date=datetime.date.today(),
                customer_name=customer_name,
                street_address=street_address,
                city=city,
                state=state,
                zip_code=zip_code,
                submitted_by=submitted_by,
                folder_name=folder_name,
                lead_value=lead_value,
            )
        except Exception as exc:
            supabase_error = exc
            _safe_debug(f"[ERROR] Supabase tracker refresh failed for {folder_name}: {exc}")
    if flags.spreadsheet_writes_active:
        _update_tracking_after_save_spreadsheet(
            folder_name=folder_name,
            customer_name=customer_name,
            street_address=street_address,
            city=city,
            state=state,
            zip_code=zip_code,
            product=product,
            roof_type=roof_type,
            total_squares=total_squares,
            warranty_incl=warranty_incl,
            submitted_by=submitted_by,
            proposal_folder=proposal_folder,
            tp10=tp10,
            tp15=tp15,
            tp20=tp20,
            lead_value=lead_value,
        )
    elif supabase_error is not None:
        raise supabase_error

def find_profit_summary_file(folder_path):
    # Safely handle missing/non-existent folder
    if not folder_path or not os.path.isdir(folder_path):
        return None
    for f in os.listdir(folder_path):
        if f.startswith("Profit Summary") and f.endswith((".xlsx")):
            return os.path.join(folder_path, f)
    return None


_PROPOSAL_DRAFT_EXCLUDED_FIELDS = frozenset({
    "action",
    "database_proposal_id",
    "customer_name_existing",
    "read_only",
    "readonly",
    "selected_proposal_file_paths",
})


def _proposal_draft_detail_from_form() -> dict:
    """Capture safe proposal form values so a contact detour is lossless."""
    snapshot = {}
    for field_name in request.form:
        if field_name in _PROPOSAL_DRAFT_EXCLUDED_FIELDS:
            continue
        if not re.fullmatch(r"[A-Za-z][A-Za-z0-9_]{0,79}", field_name):
            continue
        snapshot[field_name] = str(request.form.get(field_name) or "")[:20000]
    return snapshot


def _proposal_draft_detail_for_display(draft_detail: dict) -> dict:
    """Restore form values with the types required by the detail template."""
    restored = dict(draft_detail) if isinstance(draft_detail, dict) else {}
    raw_office_fee = str(restored.get("office_fee_pct") or "").strip()
    if not raw_office_fee:
        restored["office_fee_pct"] = None
        return restored
    try:
        office_fee = float(
            raw_office_fee.replace("%", "").replace("$", "").replace(",", "")
        )
    except (TypeError, ValueError):
        restored["office_fee_pct"] = None
    else:
        restored["office_fee_pct"] = (
            office_fee / 100.0 if office_fee > 1 else office_fee
        )
    return restored



@app.route('/update-proposal/<folder_name>', methods=['POST'])
def update_proposal(folder_name):
    
    allow_blank = (folder_name in ("NEW", "__blank__"))
    folder_path = None if allow_blank else resolve_open_proposal_folder(folder_name)

    action = (request.form.get('action') or '').strip().lower()
    database_proposal_id = (request.form.get('database_proposal_id') or '').strip()

    if action == "contact":
        customer_name = " ".join(
            (request.form.get("customer_name") or "").split()
        )
        contact_search = customer_name
        street_address = " ".join(
            (request.form.get("street_address") or "").split()
        )
        proposal_name = (
            f"{customer_name} - {street_address}"
            if street_address else customer_name
        )
        try:
            proposal_store = get_proposal_tracking_store()
            if database_proposal_id:
                database_proposal_id = str(uuid.UUID(database_proposal_id))
                stored_customer_name = (
                    customer_name
                    or f"New Proposal {database_proposal_id[:8].upper()}"
                )
                if not proposal_name:
                    proposal_name = stored_customer_name
                proposal_store.upsert_from_proposal_save(
                    proposal_id=database_proposal_id,
                    created_date=None,
                    customer_name=stored_customer_name,
                    street_address=street_address,
                    city=request.form.get("city", ""),
                    state=request.form.get("state", ""),
                    zip_code=request.form.get("zip_code", ""),
                    submitted_by=request.form.get("submitted_by", ""),
                    folder_name=(proposal_name if allow_blank else folder_name),
                    lead_value=request.form.get("lead", ""),
                    estimated_by="",
                )
            else:
                if not customer_name:
                    draft_reference = uuid.uuid4().hex[:8].upper()
                    customer_name = f"New Proposal {draft_reference}"
                    proposal_name = customer_name
                database_proposal_id = (
                    proposal_store.upsert_from_proposal_save(
                        created_date=None,
                        customer_name=customer_name,
                        street_address=street_address,
                        city=request.form.get("city", ""),
                        state=request.form.get("state", ""),
                        zip_code=request.form.get("zip_code", ""),
                        submitted_by=request.form.get("submitted_by", ""),
                        folder_name=(
                            proposal_name if allow_blank else folder_name
                        ),
                        lead_value=request.form.get("lead", ""),
                        estimated_by="",
                    )
                )
            proposal_store.save_proposal_draft_detail(
                database_proposal_id,
                _proposal_draft_detail_from_form(),
            )
        except (ValueError, ContactStoreError, TenantAuthenticationError) as exc:
            flash(str(exc), "danger")
            if allow_blank:
                return redirect(url_for("proposal_details_new"))
            return redirect(url_for("proposal_details", folder_name=folder_name))
        return redirect(url_for(
            "contact_management",
            attach_to_proposal=database_proposal_id,
            proposal_name=proposal_name,
            q=(contact_search or None),
            return_to_detail="1",
            proposal_folder_name=(
                proposal_name if allow_blank else folder_name
            ),
            customer_was_blank=("1" if not request.form.get("customer_name", "").strip() else "0"),
        ))

    excel_file = None
    if not allow_blank:
        if not folder_path:
            return f"Open proposal folder not found: {folder_name}", 404
        excel_file = find_profit_summary_file(folder_path)
        if not excel_file:
            return f"No 'Profit Summary' Excel file found in {folder_name}", 404

    # Collect updated data from the form and convert to appropriate types
    def parse_float(val, default=0.0):
        try:
            if val is None:
                return default
            if isinstance(val, str):
                cleaned = val.replace('$', '').replace(',', '').strip()
                if cleaned == '':
                    return default
                return float(cleaned)
            return float(val)
        except (TypeError, ValueError):
            return default

    def parse_int(val, default=None):
        try:
            if val is None:
                return default
            s = str(val).strip()
            if s == "":
                return default
            return int(float(s))
        except (TypeError, ValueError):
            return default

    # --- Lead field from Proposal Details ---
    lead_val = (request.form.get('lead') or '').strip()
    if lead_val == '':
        lead_val = None
    selected_proposal_file_paths = _selected_proposal_file_paths_from_form()

    # If the Blank Proposal flow hits the Create button, build artifacts and redirect
    if allow_blank and action == 'create':
        finalized_proposal_id = ""
        if database_proposal_id:
            try:
                finalized_proposal_id = str(uuid.UUID(database_proposal_id))
            except ValueError:
                flash("That proposal draft could not be selected.", "danger")
                return redirect(url_for("proposal_details_new"))

        # Pull the minimal required fields from the posted form
        customer_name = (request.form.get('customer_name') or '').strip()
        street_address = (request.form.get('street_address') or '').strip()
        city = (request.form.get('city') or '').strip()
        state = (request.form.get('state') or '').strip()
        zip_code = (request.form.get('zip_code') or '').strip()
        roof_type = (request.form.get('current_roof') or request.form.get('roof_type') or '').strip()
        raw_flat_roof_squares = request.form.get('flat_roof_squares')
        raw_wall_squares = request.form.get('wall_squares')
        flat_roof_squares = parse_float(raw_flat_roof_squares)
        wall_squares = parse_float(raw_wall_squares)
        try:
            total_squares = int(compute_total_squares(
                flat_roof_squares,
                wall_squares,
                fallback_total=parse_float(request.form.get('squares'), 0),
                raw_flat=raw_flat_roof_squares,
                raw_wall=raw_wall_squares,
            ))
        except Exception:
            total_squares = 0
        warranty_incl = (request.form.get('warranty_incl') or 'No').strip()
        include_travel = (request.form.get('include_travel') or 'No').strip()
        previous_include_travel = (request.form.get('previous_include_travel') or include_travel).strip()
        previous_calc_travel_total = parse_float(request.form.get('previous_calc_travel_total'))
        product = (request.form.get('product') or '').strip()
        submitted_by = (request.form.get('submitted_by') or '').strip()
        includes_text = (request.form.get('includes_text') or '').strip()
        proposal_language = (request.form.get('proposal_language') or includes_text or '').strip()
        repair_costs_total = parse_float(request.form.get('repair_costs_total'))
        if repair_costs_total > 0:
            proposal_language = REPAIR_COSTS_PROPOSAL_LANGUAGE

        # Collect any additional mapped fields present on the form for initial write
        def _pf(name, default=None):
            val = request.form.get(name)
            if val is None or str(val).strip() == '':
                return default
            try:
                return float(val.replace('$','').replace(',',''))
            except Exception:
                return val

        def _clean_price(name):
            if str(request.form.get(f"manual_{name}") or "").strip().lower() != "yes":
                return None
            if str(request.form.get(f"ov_{name}") or "").strip().lower() != "yes":
                return None
            return _pf(name)

        mapped_data_full = {
            "flat_roof_squares": flat_roof_squares,
            "wall_squares": wall_squares,
            "price_per_sq_10": _pf("price_per_sq_10"),
            "labor_days": _pf("labor_days"),
            "silicone_units_10": _pf("silicone_units_10"),
            "gaco_patch_units": _pf("gaco_patch_units"),
            "bleed_trap_units": _pf("bleed_trap_units"),
            "gaco_e5320_units": _pf("gaco_e5320_units", 0.0),
            "gaco_e5320_price": _clean_price("gaco_e5320_price"),
            "sw_1flash_units": _pf("sw_1flash_units"),
            "sw_bleed_block_units": _pf("sw_bleed_block_units"),
            "drainage_mat_units": _pf("drainage_mat_units"),
            "foam_units": _pf("foam_units"),
            "silicone_price": _clean_price("silicone_price"),
            "gaco_patch_price": _clean_price("gaco_patch_price"),
            "bleed_trap_price": _clean_price("bleed_trap_price"),
            "sw_1flash_price": _clean_price("sw_1flash_price"),
            "sw_bleed_block_price": _clean_price("sw_bleed_block_price"),
            "drainage_mat_price": _clean_price("drainage_mat_price"),
            "foam_price": _clean_price("foam_price"),
            "rfc_labor_price": _clean_price("rfc_labor_price"),
            "pcs_labor_price": _clean_price("pcs_labor_price"),
            "scarifying_total": _pf("scarifying_total"),
            "travel_total": _pf("travel_total"),
            "repair_costs_total": repair_costs_total,
            "include_travel": include_travel,
            "previous_include_travel": previous_include_travel,
            "previous_calc_travel_total": previous_calc_travel_total,
            "adjusted_coverage": _pf("adjusted_coverage", 0.0),
            "previous_adjusted_coverage": _pf("previous_adjusted_coverage", 0.0),
            "previous_silicone_units_10": _pf("previous_silicone_units_10", 0.0),
            "proposal_note": (request.form.get("proposal_note") or "").strip(),
            "proposal_language": proposal_language,
            "total_price_10": _pf("total_price_10"),
            "total_price_15": _pf("total_price_15"),
            "total_price_20": _pf("total_price_20"),
            "lead": lead_val,
            "pcs_or_roofer_ind": (request.form.get("pcs_or_roofer_ind") or "").strip(),
        }
        # Remove Nones to avoid overwriting with blanks
        mapped_data_full = {k: v for k, v in mapped_data_full.items() if v is not None}

        # Create artifacts using the helper (same behavior as /new)
        new_folder = create_proposal_from_fields(
            customer_name=customer_name,
            street_address=street_address,
            city=city,
            state=state,
            zip_code=zip_code,
            roof_type=roof_type,
            total_squares=total_squares,
            warranty_incl=warranty_incl,
            product=product,
            proposal_language=proposal_language,
            submitted_by=submitted_by,
            mapped_data=mapped_data_full,
            pdf_async=False,
            use_libreoffice=True,
            update_tracking=not bool(finalized_proposal_id),
            copy_destination=False,
        )
        if finalized_proposal_id:
            proposal_store = get_proposal_tracking_store()
            proposal_store.upsert_from_proposal_save(
                proposal_id=finalized_proposal_id,
                created_date=datetime.date.today(),
                customer_name=customer_name,
                street_address=street_address,
                city=city,
                state=state,
                zip_code=zip_code,
                submitted_by=submitted_by,
                folder_name=new_folder,
                lead_value=lead_val,
            )
            try:
                proposal_store.clear_proposal_draft_detail(finalized_proposal_id)
            except (ContactStoreError, TenantAuthenticationError) as exc:
                _safe_debug(
                    "Could not clear finalized proposal draft detail: "
                    f"{exc}"
                )
        new_proposal_folder = os.path.join(PROPOSAL_TEMP_DIR, new_folder)
        copy_proposal_to_submitter_destination(
            new_proposal_folder,
            new_folder,
            submitted_by,
            wait_for_pdfs=False,
        )
        move_selected_proposal_files_to_folder(
            selected_proposal_file_paths,
            get_submitter_proposal_folder(new_folder, submitted_by),
        )
        return redirect(url_for('proposal_list'))

    raw_flat_roof_squares = request.form.get('flat_roof_squares')
    raw_wall_squares = request.form.get('wall_squares')
    flat_roof_squares = parse_float(raw_flat_roof_squares)
    wall_squares = parse_float(raw_wall_squares)
    squares = compute_total_squares(
        flat_roof_squares,
        wall_squares,
        fallback_total=parse_float(request.form.get('squares'), 0),
        raw_flat=raw_flat_roof_squares,
        raw_wall=raw_wall_squares,
    )
    product = request.form.get('product')
    roof_type = request.form.get('current_roof')
    labor_days = parse_int(request.form.get('labor_days'), None)    
    warranty_incl = request.form.get('warranty_incl', 'No').strip()
    include_travel = request.form.get('include_travel', 'No').strip()
    previous_include_travel = request.form.get('previous_include_travel', include_travel).strip()
    previous_calc_travel_total = parse_float(request.form.get('previous_calc_travel_total'))
    previous_warranty_incl = request.form.get('previous_warranty_incl', warranty_incl)
    price_per_sq_10 = parse_float(request.form.get('price_per_sq_10'))
    commission_pct = parse_float(request.form.get('commission_pct'))

    submitted_by = request.form.get('submitted_by')
    previous_submitted_by = request.form.get('previous_submitted_by', '')

    raw_office_fee_pct = request.form.get('office_fee_pct')
    if raw_office_fee_pct is None or str(raw_office_fee_pct).strip() == '':
        office_fee_pct = None  # allow defaulting based on Submitted By in calc
    else:
        cleaned_office = str(raw_office_fee_pct).replace('%', '').strip()
        office_fee_value = parse_float(cleaned_office)
        if office_fee_value is None:
            office_fee_pct = None
        elif office_fee_value > 1:
            office_fee_pct = office_fee_value / 100.0  # e.g., "5" -> 0.05
        else:
            office_fee_pct = office_fee_value         # already decimal like 0.05

    raw_adjusted_coverage = request.form.get('adjusted_coverage') or request.form.get('adjust_coverage')
    adjusted_coverage = None if raw_adjusted_coverage is None or str(raw_adjusted_coverage).strip() == '' else parse_float(raw_adjusted_coverage)
    raw_silicone_units_10 = request.form.get('silicone_units_10')
    silicone_units_10 = None if raw_silicone_units_10 is None or str(raw_silicone_units_10).strip() == '' else parse_float(raw_silicone_units_10)
    raw_silicone_price = request.form.get('silicone_price')
    silicone_price = None if raw_silicone_price is None or str(raw_silicone_price).strip() == '' else parse_float(raw_silicone_price)
    raw_gaco_patch_units = request.form.get('gaco_patch_units')
    gaco_patch_units = None if raw_gaco_patch_units is None or str(raw_gaco_patch_units).strip() == '' else parse_float(raw_gaco_patch_units)
    raw_gaco_patch_price = request.form.get('gaco_patch_price')
    gaco_patch_price = None if raw_gaco_patch_price is None or str(raw_gaco_patch_price).strip() == '' else parse_float(raw_gaco_patch_price)
    raw_bleed_trap_units = request.form.get('bleed_trap_units') or request.form.get('sw_bleed_trap_units')
    bleed_trap_units = None if raw_bleed_trap_units is None or str(raw_bleed_trap_units).strip() == '' else parse_float(raw_bleed_trap_units)
    raw_bleed_trap_price = request.form.get('bleed_trap_price') or request.form.get('sw_bleed_trap_price')
    bleed_trap_price = None if raw_bleed_trap_price is None or str(raw_bleed_trap_price).strip() == '' else parse_float(raw_bleed_trap_price)
    raw_gaco_e5320_units = request.form.get('gaco_e5320_units')
    gaco_e5320_units = 0.0 if raw_gaco_e5320_units is None or str(raw_gaco_e5320_units).strip() == '' else parse_float(raw_gaco_e5320_units)
    raw_gaco_e5320_price = request.form.get('gaco_e5320_price')
    gaco_e5320_price = None if raw_gaco_e5320_price is None or str(raw_gaco_e5320_price).strip() == '' else parse_float(raw_gaco_e5320_price)
    raw_sw_1flash_units = request.form.get('sw_1flash_units')
    sw_1flash_units = None if raw_sw_1flash_units is None or str(raw_sw_1flash_units).strip() == '' else parse_float(raw_sw_1flash_units)
    raw_sw_1flash_price = request.form.get('sw_1flash_price')
    sw_1flash_price = None if raw_sw_1flash_price is None or str(raw_sw_1flash_price).strip() == '' else parse_float(raw_sw_1flash_price)
    raw_sw_bleed_block_units = request.form.get('sw_bleed_block_units')
    sw_bleed_block_units = None if raw_sw_bleed_block_units is None or str(raw_sw_bleed_block_units).strip() == '' else parse_float(raw_sw_bleed_block_units)
    raw_sw_bleed_block_price = request.form.get('sw_bleed_block_price')
    sw_bleed_block_price = None if raw_sw_bleed_block_price is None or str(raw_sw_bleed_block_price).strip() == '' else parse_float(raw_sw_bleed_block_price)
    raw_drainage_mat_units = request.form.get('drainage_mat_units')
    drainage_mat_units = None if raw_drainage_mat_units is None or str(raw_drainage_mat_units).strip() == '' else parse_float(raw_drainage_mat_units)
    raw_drainage_mat_price = request.form.get('drainage_mat_price')
    drainage_mat_price = None if raw_drainage_mat_price is None or str(raw_drainage_mat_price).strip() == '' else parse_float(raw_drainage_mat_price)
    raw_foam_units = request.form.get('foam_units')
    foam_units = None if raw_foam_units is None or str(raw_foam_units).strip() == '' else parse_float(raw_foam_units)
    raw_foam_price = request.form.get('foam_price')
    foam_price = None if raw_foam_price is None or str(raw_foam_price).strip() == '' else parse_float(raw_foam_price)
    raw_rfc_labor_price = request.form.get('rfc_labor_price')
    rfc_labor_price = None if raw_rfc_labor_price is None or str(raw_rfc_labor_price).strip() == '' else parse_float(raw_rfc_labor_price)
    pcs_labor_price = parse_float(request.form.get('pcs_labor_price'))

    def _clear_non_overridden_price(field_name, value):
        return value if str(request.form.get(f"ov_{field_name}") or "").strip().lower() == "yes" else None

    silicone_price = _clear_non_overridden_price("silicone_price", silicone_price)
    gaco_patch_price = _clear_non_overridden_price("gaco_patch_price", gaco_patch_price)
    bleed_trap_price = _clear_non_overridden_price("bleed_trap_price", bleed_trap_price)
    gaco_e5320_price = _clear_non_overridden_price("gaco_e5320_price", gaco_e5320_price)
    sw_1flash_price = _clear_non_overridden_price("sw_1flash_price", sw_1flash_price)
    sw_bleed_block_price = _clear_non_overridden_price("sw_bleed_block_price", sw_bleed_block_price)
    drainage_mat_price = _clear_non_overridden_price("drainage_mat_price", drainage_mat_price)
    foam_price = _clear_non_overridden_price("foam_price", foam_price)
    rfc_labor_price = _clear_non_overridden_price("rfc_labor_price", rfc_labor_price)
    pcs_labor_price = _clear_non_overridden_price("pcs_labor_price", pcs_labor_price)

    raw_scarifying_total = request.form.get('scarifying_total')
    scarifying_total = parse_float(raw_scarifying_total)
    travel_total = parse_float(request.form.get('travel_total'))
    repair_costs_total = parse_float(request.form.get('repair_costs_total'))
    # Use explicit fallbacks that reflect a prior/blank state so changes are detectable
    _prev_sq_raw = request.form.get('previous_squares')
    previous_squares = parse_float(_prev_sq_raw, 0.0)  # default to 0, not current squares

    previous_roof_type = request.form.get('previous_roof_type', '')  # default to '' so a change is caught

    previous_product = request.form.get('previous_product', '')  # safe default; not used for labor_days but consistent

    _prev_adj_raw = request.form.get('previous_adjusted_coverage')
    previous_adjusted_coverage = parse_float(_prev_adj_raw, 0.0)

    _prev_units_raw = request.form.get('previous_silicone_units_10')
    # Default to current silicone_units_10 if the hidden field is missing on first render
    previous_silicone_units_10 = parse_float(_prev_units_raw, (silicone_units_10 or 0.0))

    # Simple text field; persist across recalcs
    proposal_note = (request.form.get('proposal_note') or '').strip()
    proposal_language = (request.form.get('proposal_language') or '').strip()
    if repair_costs_total > 0:
        proposal_language = REPAIR_COSTS_PROPOSAL_LANGUAGE
    customer_name = (request.form.get('customer_name') or '').strip()
    street_address = (request.form.get('street_address') or '').strip()
    city = (request.form.get('city') or '').strip()
    state = (request.form.get('state') or '').strip()
    zip_code = (request.form.get('zip_code') or '').strip()
    pcs_or_roofer_ind = (request.form.get('pcs_or_roofer_ind') or '').strip()
    previous_pcs_or_roofer_ind = (request.form.get('previous_pcs_or_roofer_ind') or '').strip()

    # Use proposal_language as the single source of truth for downstream Word/Excel writes
    includes_text = proposal_language

    # Carry read-only flag through POST round-trips, supporting both new and legacy formats
    # New: read_only = "Yes"/"No"; Legacy: readonly = "1"/"0"
    read_only_param = request.form.get('read_only')
    if read_only_param is not None:
        readonly = (read_only_param.strip().lower() == 'yes')
    else:
        readonly = (request.form.get('readonly') == '1')

    # Prepare data dictionary for template (may include more fields as needed)
    data = {
        'flat_roof_squares': flat_roof_squares,
        'wall_squares': wall_squares,
        'squares': squares,
        'product': product,
        'current_roof': roof_type,
        'labor_days': labor_days,
        'warranty_incl': warranty_incl,
        'include_travel': include_travel,
        'previous_include_travel': previous_include_travel,
        'previous_calc_travel_total': previous_calc_travel_total,
        'price_per_sq_10': price_per_sq_10,
        'commission_pct': commission_pct,
        'adjusted_coverage': adjusted_coverage,
        'silicone_units_10': silicone_units_10,
        'silicone_price': silicone_price,
        'gaco_patch_units': gaco_patch_units,
        'gaco_patch_price': gaco_patch_price,
        'sw_1flash_units': sw_1flash_units,
        'sw_1flash_price': sw_1flash_price,
        'bleed_trap_units': bleed_trap_units,
        'bleed_trap_price': bleed_trap_price,
        'gaco_e5320_units': gaco_e5320_units,
        'gaco_e5320_price': gaco_e5320_price,
        'sw_bleed_block_units': sw_bleed_block_units,
        'sw_bleed_block_price': sw_bleed_block_price,
        'drainage_mat_units': drainage_mat_units,
        'drainage_mat_price': drainage_mat_price,
        'foam_units': foam_units,
        'foam_price': foam_price,
        'rfc_labor_price': rfc_labor_price,
        'pcs_labor_price': pcs_labor_price,
        'scarifying_total': scarifying_total,
        'travel_total': travel_total,
        'repair_costs_total': repair_costs_total,
        'previous_squares': previous_squares,
        'previous_roof_type': previous_roof_type,
        'previous_product': previous_product,
        'previous_warranty_incl': previous_warranty_incl,
        'previous_adjusted_coverage': previous_adjusted_coverage,
        'previous_silicone_units_10': previous_silicone_units_10,
        'coverage_10': 0,
        'coverage_15': 0,
        'coverage_20': 0,
        'submitted_by': submitted_by,
        'office_fee_pct': office_fee_pct,
        'previous_submitted_by': previous_submitted_by,
        'proposal_note': proposal_note,
        'proposal_language': proposal_language,
        'customer_name': customer_name,
        'street_address': street_address,
        'city': city,
        'state': state,
        'zip_code': zip_code,
        'includes_text': includes_text,
        'pcs_or_roofer_ind': pcs_or_roofer_ind,
        'previous_pcs_or_roofer_ind': previous_pcs_or_roofer_ind,
        'selected_proposal_file_paths': json.dumps(selected_proposal_file_paths),
    }
    for _price_field in UNIT_PRICE_FIELDS:
        data[f"manual_{_price_field}"] = request.form.get(f"manual_{_price_field}") or "No"

    # If saving an existing proposal, archive old artifacts and regenerate in the same folder
    if action == 'save' and not allow_blank and folder_name:
        save_started = time.perf_counter()
        proposal_folder = folder_path
        archive_started = time.perf_counter()
        _archive_existing_artifacts(proposal_folder)
        _log_timing(f"artifact archive for {folder_name}", archive_started)
        # Collect any additional mapped fields present on the form for initial write
        def _pf(name, default=None):
            val = request.form.get(name)
            if val is None or str(val).strip() == '':
                return default
            try:
                return float(val.replace('$','').replace(',',''))
            except Exception:
                return val

        def _clean_price(name):
            if str(request.form.get(f"ov_{name}") or "").strip().lower() != "yes":
                return None
            return _pf(name)

        mapped_data_full = {
            "flat_roof_squares": flat_roof_squares,
            "wall_squares": wall_squares,
            "price_per_sq_10": _pf("price_per_sq_10"),
            "labor_days": _pf("labor_days"),
            "silicone_units_10": _pf("silicone_units_10"),
            "gaco_patch_units": _pf("gaco_patch_units"),
            "bleed_trap_units": _pf("bleed_trap_units"),
            "gaco_e5320_units": _pf("gaco_e5320_units", 0.0),
            "gaco_e5320_price": _clean_price("gaco_e5320_price"),
            "sw_1flash_units": _pf("sw_1flash_units"),
            "sw_bleed_block_units": _pf("sw_bleed_block_units"),
            "drainage_mat_units": _pf("drainage_mat_units"),
            "foam_units": _pf("foam_units"),
            "silicone_price": _clean_price("silicone_price"),
            "gaco_patch_price": _clean_price("gaco_patch_price"),
            "bleed_trap_price": _clean_price("bleed_trap_price"),
            "sw_1flash_price": _clean_price("sw_1flash_price"),
            "sw_bleed_block_price": _clean_price("sw_bleed_block_price"),
            "drainage_mat_price": _clean_price("drainage_mat_price"),
            "foam_price": _clean_price("foam_price"),
            "rfc_labor_price": _clean_price("rfc_labor_price"),
            "pcs_labor_price": _clean_price("pcs_labor_price"),
            "scarifying_total": _pf("scarifying_total"),
            "travel_total": _pf("travel_total"),
            "repair_costs_total": _pf("repair_costs_total"),
            "include_travel": include_travel,
            "previous_include_travel": previous_include_travel,
            "previous_calc_travel_total": previous_calc_travel_total,
            "adjusted_coverage": adjusted_coverage,
            "previous_adjusted_coverage": previous_adjusted_coverage,
            "previous_silicone_units_10": previous_silicone_units_10,
            "previous_roof_type": previous_roof_type,
            "previous_product": previous_product,
            "previous_squares": previous_squares,
            "previous_pcs_or_roofer_ind": previous_pcs_or_roofer_ind,
            "proposal_note": (request.form.get("proposal_note") or "").strip(),
            "proposal_language": proposal_language,
            "total_price_10": _pf("total_price_10"),
            "total_price_15": _pf("total_price_15"),
            "total_price_20": _pf("total_price_20"),
            "lead": lead_val,
            "pcs_or_roofer_ind": (request.form.get("pcs_or_roofer_ind") or "").strip(),
        }
        # Remove Nones to avoid overwriting with blanks
        mapped_data_full = {k: v for k, v in mapped_data_full.items() if v is not None}

        create_proposal_from_fields(
            customer_name=customer_name,
            street_address=street_address,
            city=city,
            state=state,
            zip_code=zip_code,
            roof_type=roof_type,
            total_squares=int(squares) if squares else 0,
            warranty_incl=warranty_incl,
            product=product,
            proposal_language=proposal_language,
            submitted_by=submitted_by,
            target_folder=proposal_folder,
            mapped_data=mapped_data_full,
            pdf_async=False,
            use_libreoffice=True,
            update_tracking=False,
            copy_destination=False,
        )
        _log_timing(f"save core regeneration for {folder_name}", save_started)

        total_price_10 = parse_float(mapped_data_full.get("total_price_10"), 0)
        total_price_15 = parse_float(mapped_data_full.get("total_price_15"), 0)
        total_price_20 = parse_float(mapped_data_full.get("total_price_20"), 0)

        _run_background_task(
            f"tracker refresh for {folder_name}",
            lambda: update_tracking_after_save(
                folder_name=folder_name,
                customer_name=customer_name,
                street_address=street_address,
                city=city,
                state=state,
                zip_code=zip_code,
                product=product,
                roof_type=roof_type,
                total_squares=int(squares) if squares else 0,
                warranty_incl=warranty_incl,
                submitted_by=submitted_by,
                proposal_folder=proposal_folder,
                tp10=total_price_10,
                tp15=total_price_15,
                tp20=total_price_20,
                lead_value=lead_val,
            ),
        )
        copy_started = time.perf_counter()
        copy_proposal_to_submitter_destination(
            proposal_folder,
            folder_name,
            submitted_by,
            wait_for_pdfs=False,
        )
        move_selected_proposal_files_to_folder(
            selected_proposal_file_paths,
            get_submitter_proposal_folder(folder_name, submitted_by),
        )
        _log_timing(f"destination copy for {folder_name}", copy_started)
        return redirect(url_for('proposal_list'))

    # Call calculation_routine and merge results
    calc_result = calculation_routine(
        squares,
        product,
        roof_type,
        labor_days,
        warranty_incl,
        include_travel,
        price_per_sq_10,
        commission_pct,
        submitted_by=submitted_by,
        previous_submitted_by=previous_submitted_by,
        office_fee_pct=office_fee_pct,
        adjusted_coverage=adjusted_coverage,
        silicone_units_10=silicone_units_10,
        silicone_price=silicone_price,
        gaco_patch_units=gaco_patch_units,
        gaco_patch_price=gaco_patch_price,
        sw_1flash_units=sw_1flash_units,
        sw_1flash_price=sw_1flash_price,
        bleed_trap_units=bleed_trap_units,
        bleed_trap_price=bleed_trap_price,
        gaco_e5320_units=gaco_e5320_units,
        gaco_e5320_price=gaco_e5320_price,
        sw_bleed_block_units=sw_bleed_block_units,
        sw_bleed_block_price=sw_bleed_block_price,
        drainage_mat_units=drainage_mat_units,
        drainage_mat_price=drainage_mat_price,
        foam_units=foam_units,
        foam_price=foam_price,
        rfc_labor_price=rfc_labor_price,
        pcs_labor_price=pcs_labor_price,
        scarifying_total=scarifying_total,
        travel_total=travel_total,
        repair_costs_total=repair_costs_total,
        previous_squares=previous_squares,
        previous_roof_type=previous_roof_type,
        previous_product=previous_product,
        previous_adjusted_coverage=previous_adjusted_coverage,
        previous_silicone_units_10=previous_silicone_units_10,
        proposal_note=proposal_note,
        pcs_or_roofer_ind=pcs_or_roofer_ind,
        previous_pcs_or_roofer_ind=previous_pcs_or_roofer_ind,
        previous_include_travel=previous_include_travel,
        previous_calc_travel_total=previous_calc_travel_total,
    )
    # Persist key header fields and note across round trip so they are not lost
    calc_result.update({
        "flat_roof_squares": flat_roof_squares,
        "wall_squares": wall_squares,
        "customer_name": customer_name,
        "street_address": street_address,
        "city": city,
        "state": state,
        "zip_code": zip_code,
        "proposal_note": proposal_note,
        "proposal_language": includes_text,
        "includes_text": includes_text,
        "pcs_or_roofer_ind": pcs_or_roofer_ind,
        "previous_pcs_or_roofer_ind": previous_pcs_or_roofer_ind,
        "previous_include_travel": include_travel,
        "previous_calc_travel_total": calc_result.get("calc_travel_total", 0),
    })

    data.update(calc_result)

    # --- Display fallbacks for fields saved as formulas (no cached value yet) ---
    try:
        # Derive office fee pct and commission pct for fallback math
        _submitted_by_disp = str(data.get("submitted_by") or "")
        _office_fee_pct_disp = office_fee_pct_for_submitter(_submitted_by_disp)
        _commission_pct_disp = commission_pct_for_submitter(_submitted_by_disp)

        # Call calculation_routine with "previous_*" equal to current to avoid resetting overrides
        calc_disp = calculation_routine(
            squares=float(data.get("squares") or 0.0),
            product=str(data.get("product") or ""),
            roof_type=str(data.get("current_roof") or ""),
            labor_days=float(data.get("labor_days") or 0.0),
            warranty_incl=str(data.get("warranty_incl") or "No"),
            include_travel=str(data.get("include_travel") or "No"),
            price_per_sq_10=float(data.get("price_per_sq_10") or 0.0),
            commission_pct=float(_commission_pct_disp),
            submitted_by=_submitted_by_disp,
            previous_submitted_by=_submitted_by_disp,
            office_fee_pct=float(_office_fee_pct_disp),
            adjusted_coverage=float(data.get("adjusted_coverage") or 0.0),
            silicone_units_10=float(data.get("silicone_units_10") or 0.0),
            silicone_price=float(data.get("silicone_price") or 0.0),
            gaco_patch_units=float(data.get("gaco_patch_units") or 0.0),
            gaco_patch_price=float(data.get("gaco_patch_price") or 0.0),
            sw_1flash_units=float(data.get("sw_1flash_units") or 0.0),
            sw_1flash_price=float(data.get("sw_1flash_price") or 0.0),
            bleed_trap_units=float(data.get("bleed_trap_units") or 0.0),
            bleed_trap_price=float(data.get("bleed_trap_price") or 0.0),
            gaco_e5320_units=float(data.get("gaco_e5320_units") or 0.0),
            gaco_e5320_price=float(data.get("gaco_e5320_price") or GACO_E5320_PRICE),
            sw_bleed_block_units=float(data.get("sw_bleed_block_units") or 0.0),
            sw_bleed_block_price=float(data.get("sw_bleed_block_price") or 0.0),
            drainage_mat_units=float(data.get("drainage_mat_units") or 0.0),
            drainage_mat_price=float(data.get("drainage_mat_price") or 0.0),
            foam_units=float(data.get("foam_units") or 0.0),
            foam_price=float(data.get("foam_price") or 0.0),
            rfc_labor_price=float(data.get("rfc_labor_price") or 0.0),
            pcs_labor_price=float(data.get("pcs_labor_price") or 0.0),
            scarifying_total=float(data.get("scarifying_total") or 0.0),
            travel_total=float(data.get("travel_total") or 0.0),
            repair_costs_total=float(data.get("repair_costs_total") or 0.0),
            previous_squares=float(data.get("squares") or 0.0),
            previous_roof_type=str(data.get("current_roof") or ""),
            previous_product=str(data.get("product") or ""),
            previous_adjusted_coverage=float(data.get("adjusted_coverage") or 0.0),
            previous_silicone_units_10=float(data.get("silicone_units_10") or 0.0),
            proposal_note=str(data.get("proposal_note") or ""),
            pcs_or_roofer_ind=pcs_or_roofer_ind,
            previous_pcs_or_roofer_ind=pcs_or_roofer_ind,
            previous_include_travel=str(data.get("include_travel") or "No"),
            previous_calc_travel_total=data.get("calc_travel_total") or 0,
        )

        # Fields saved as formulas that may be blank/0 when the workbook hasn't cached results
        _fallback_fields = [
            # Inputs that may be stored as formulas (or blank when N/A) but must show a number on screen
            "labor_days",
            "silicone_units_10","silicone_units_15","silicone_units_20","gaco_patch_units","bleed_trap_units","gaco_e5320_units",
            "sw_1flash_units","sw_bleed_block_units","drainage_mat_units","foam_units",

            # Unit prices that may be formulas or overridden numbers
            "silicone_price","gaco_patch_price","bleed_trap_price","gaco_e5320_price",
            "sw_1flash_price","sw_bleed_block_price","drainage_mat_price",
            "foam_price","rfc_labor_price","pcs_labor_price",

            # Price per square and totals
            "price_per_sq_10","price_per_sq_15","price_per_sq_20",
            "total_price_10","total_price_15","total_price_20",

            # Cost/fee totals and downstream profit metrics
            "total_cost","total_cost_15","total_cost_20",
            "warranty_10_total","office_fee_total","office_fee_15_total","office_fee_20_total",
            "silicone_total","silicone_15_total","silicone_20_total","gaco_patch_total","bleed_trap_total","gaco_e5320_total",
            "sw_1flash_total","sw_bleed_block_total","drainage_mat_total",
            "foam_total","rfc_labor_total","pcs_labor_total",
            "commission_amt","commission_amt_15","commission_amt_20",
            "profit_share","profit_share_15","profit_share_20",
            "daily_profit","daily_profit_15","daily_profit_20",
            "profit_pct","profit_pct_15","profit_pct_20",
            "pcs_profit","pcs_profit_15","pcs_profit_20"
        ]

        def _is_missing(v):
            try:
                if v is None:
                    return True
                if isinstance(v, float) and math.isnan(v):
                    return True
                return float(v) == 0.0
            except Exception:
                # Strings (formulas) won't be missing here; only numbers can be missing
                return False

        for _k in _fallback_fields:
            if _is_missing(data.get(_k)):
                if _k in calc_disp:
                    data[_k] = calc_disp[_k]
        data["calc_travel_total"] = calc_disp.get("calc_travel_total", data.get("calc_travel_total", 0))
        data["previous_calc_travel_total"] = data["calc_travel_total"]
        data["previous_include_travel"] = data.get("include_travel") or "No"
    except Exception as _fallback_e:
        # Non-fatal: if anything goes wrong, we just skip the display fallback
        _safe_debug(f"[DEBUG] display fallback error: {_fallback_e}")

    # --- Merge evaluated display data (prefers ghost cells) before rendering (POST) ---
    try:
        _folder_path = resolve_open_proposal_folder(folder_name) or os.path.join(PROPOSAL_TEMP_DIR, folder_name)
        data = merge_display_fallbacks(
            data,
            _folder_path,
            folder_name,
            prefer_saved_derived=False,
        )
        try:
            with open(APP_ERROR_LOG, "a", encoding="utf-8") as _f:
                _f.write(
                    f"\n[PD POST] ghost merge {folder_name}: sw1(u/p/t)="
                    f"({data.get('sw_1flash_units')},{data.get('sw_1flash_price')},{data.get('sw_1flash_total')}) "
                    f"sil10={data.get('silicone_units_10')}\n"
                )
        except Exception:
            pass
    except Exception:
        pass

    # POST recalculations should display the freshly computed derived values. Saved
    # workbook/ghost values are only fallbacks and can be stale after submitter changes.
    for _k in (
        "total_price_10", "total_price_15", "total_price_20",
        "warranty_10_total", "office_fee_total", "office_fee_15_total", "office_fee_20_total",
        "total_cost", "total_cost_15", "total_cost_20",
        "silicone_total", "silicone_15_total", "silicone_20_total",
        "gaco_patch_total", "bleed_trap_total", "gaco_e5320_total", "sw_1flash_total",
        "sw_bleed_block_total", "drainage_mat_total",
        "foam_total", "rfc_labor_total", "pcs_labor_total",
        "commission_amt", "commission_amt_15", "commission_amt_20",
        "profit_share", "profit_share_15", "profit_share_20",
        "daily_profit", "daily_profit_15", "daily_profit_20",
        "profit_pct", "profit_pct_15", "profit_pct_20",
        "pcs_profit", "pcs_profit_15", "pcs_profit_20",
        "commission_pct",
    ):
        if _k in calc_result:
            data[_k] = calc_result[_k]

    # Recalculate and auto-submit POSTs should only refresh on-screen values.
    # Any file replacement/copy work belongs exclusively to the save/create flows above.

    # Ensure Lead value is reflected back to the UI
    try:
        if 'data' in locals() and isinstance(data, dict):
            data['lead'] = lead_val
    except Exception:
        pass

    if allow_blank and action in ("full_detail", "full_detail_preview"):
        return render_template(
            "proposal_full_detail.html",
            data=data,
            folder_name="NEW",
            is_preview=True,
        )

    return render_template(
        "proposal_details.html",
        data=data,
        **data,
        customer_organization_names=proposal_customer_organization_names(),
        folder_name=folder_name,
        readonly=readonly,
        is_blank=(folder_name in ("NEW", "__blank__")),
        database_proposal_id=database_proposal_id or None,
    )


    

@app.route('/proposal_details/new', methods=['GET'])
def proposal_details_new():
    data = make_blank_data()
    # Support both new and legacy readonly query parameters
    read_only_param = request.args.get('read_only')
    if read_only_param is not None:
        readonly = (read_only_param.strip().lower() == 'yes')
    else:
        readonly = (request.args.get('readonly') == '1')
    return render_template(
        "proposal_details.html",
        data=data,
        **data,
        customer_organization_names=proposal_customer_organization_names(),
        folder_name="NEW",
        readonly=readonly,
        is_blank=True,
    )

@app.route('/proposal_details')
def proposal_details_query():
    folder_name = (request.args.get('folder_name') or '').strip()
    if not folder_name:
        return redirect(url_for('proposal_list'))
    if PROPOSAL_DATABASE_SOURCE_ENABLED and not _resolve_existing_proposal_folder(folder_name):
        proposal_id = (request.args.get('proposal_id') or '').strip()
        try:
            proposal = get_proposal_tracking_store().get_management_proposal(
                proposal_id
            )
        except (ContactStoreError, TenantAuthenticationError) as exc:
            flash(str(exc), "danger")
            return redirect(url_for('proposal_list'))
        if proposal:
            data = make_blank_data()
            data.update(
                _proposal_draft_detail_for_display(
                    proposal.get("draft_detail") or {}
                )
            )
            data.update({
                "customer_name": proposal.get("customer_name", ""),
                "street_address": proposal.get("project_street_address", ""),
                "city": proposal.get("project_city", ""),
                "state": proposal.get("project_state", ""),
                "zip_code": proposal.get("project_zip_code", ""),
                "submitted_by": proposal.get("submitted_by", ""),
                "previous_submitted_by": proposal.get("submitted_by", ""),
                "lead": proposal.get("lead_source", ""),
                "proposal_note": proposal.get("response_notes", ""),
            })
            if request.args.get("customer_was_blank") == "1":
                data["customer_name"] = ""
            read_only_param = request.args.get('read_only')
            readonly = (
                read_only_param is not None
                and read_only_param.strip().lower() == 'yes'
            )
            return render_template(
                "proposal_details.html",
                data=data,
                **data,
                customer_organization_names=proposal_customer_organization_names(),
                folder_name="__blank__",
                readonly=readonly,
                is_blank=True,
                database_proposal_id=proposal["id"],
            )
    return proposal_details(folder_name)

def _resolve_existing_proposal_folder(folder_name: str) -> str | None:
    safe_folder = os.path.basename(str(folder_name or "").strip())
    if not safe_folder:
        return None
    proposals_path = resolve_open_proposal_folder(safe_folder)
    contracts_path = os.path.join(CONTRACTS_DIR, safe_folder)
    completed_path = os.path.join(COMPLETED_DIR, safe_folder)
    deadfile_path = os.path.join(DEADFILE_DIR, safe_folder)
    for candidate in (proposals_path, contracts_path, completed_path, deadfile_path):
        if candidate and os.path.isdir(candidate):
            return candidate
    return None

@app.route('/proposal_full_detail/<folder_name>')
def proposal_full_detail(folder_name):
    if folder_name in ("__blank__", "NEW"):
        return redirect(url_for("proposal_details_new"))

    safe_folder = os.path.basename(str(folder_name or "").strip())
    folder_path = _resolve_existing_proposal_folder(safe_folder)
    if not folder_path:
        return f"Folder not found: {safe_folder}", 404

    data = read_profit_summary_for_display(folder_path)
    if not data:
        return f"No Profit Summary file found in folder: {safe_folder}", 404

    data = merge_display_fallbacks(data, folder_path, safe_folder)
    return render_template(
        "proposal_full_detail.html",
        data=data,
        folder_name=safe_folder,
    )

@app.route('/proposal_details/<folder_name>')
def proposal_details(folder_name):
    # Serve blank form if requested
    if folder_name == "__blank__":
        data = make_blank_data()
        # Reuse the blank POST flow by setting folder_name to NEW for the template's form action
        # Support both new and legacy readonly query parameters
        read_only_param = request.args.get('read_only')
        if read_only_param is not None:
            readonly = (read_only_param.strip().lower() == 'yes')
        else:
            readonly = (request.args.get('readonly') == '1')
        return render_template(
            "proposal_details.html",
            data=data,
            **data,
            customer_organization_names=proposal_customer_organization_names(),
            folder_name="NEW",
            readonly=readonly,
            is_blank=True,
        )

    # --- Deadfile indicator logic ---
    dead_ind = request.args.get('dead_ind')
    if dead_ind is not None and str(dead_ind).strip().lower() in ('yes', 'true', '1'):
        # Attempt to move the folder from proposals to deadfile
        src_path = resolve_open_proposal_folder(folder_name)
        dest_path = os.path.join(DEADFILE_DIR, folder_name)
        try:
            # Check if source exists
            if not src_path or not os.path.exists(src_path):
                flash(f"Source folder for '{folder_name}' does not exist.", "error")
            elif os.path.exists(dest_path):
                flash(f"Target folder '{dest_path}' already exists.", "error")
            else:
                shutil.move(src_path, dest_path)
                flash(f"Proposal '{folder_name}' moved to dead file.", "success")
        except Exception as e:
            flash(f"Error moving proposal: {e}", "error")
        return redirect(url_for('proposal_list'))

    # --- Contract indicator logic ---
    contract_ind = request.args.get('contract_ind')
    if contract_ind is not None and str(contract_ind).strip().lower() in ('yes', 'true', '1'):
        # Attempt to move the folder from proposals to contracts
        src_path = resolve_open_proposal_folder(folder_name)
        dest_path = os.path.join(CONTRACTS_DIR, folder_name)
        try:
            # Check if source exists
            if not src_path or not os.path.exists(src_path):
                flash(f"Source folder for '{folder_name}' does not exist.", "error")
            elif os.path.exists(dest_path):
                flash(f"Target folder '{dest_path}' already exists.", "error")
            else:
                shutil.move(src_path, dest_path)
                flash(f"Proposal '{folder_name}' moved to contracts.", "success")
        except Exception as e:
            flash(f"Error moving proposal: {e}", "error")
        return redirect(url_for('proposal_list'))

    # --- Close contract indicator logic ---
    close_ind = request.args.get('close_ind')
    if close_ind is not None and str(close_ind).strip().lower() in ('yes', 'true', '1'):
        # Attempt to move the folder from contracts to completed
        src_path = os.path.join(CONTRACTS_DIR, folder_name)
        dest_path = os.path.join(COMPLETED_DIR, folder_name)
        try:
            if not os.path.exists(src_path):
                flash(f"Source folder '{src_path}' does not exist.", "error")
            elif os.path.exists(dest_path):
                flash(f"Target folder '{dest_path}' already exists.", "error")
            else:
                shutil.move(src_path, dest_path)
                flash(f"Contract '{folder_name}' closed and moved to Completed.", "success")
        except Exception as e:
            flash(f"Error closing contract: {e}", "error")
        return redirect(url_for('proposal_list', status='under'))

    # Determine source root (Open Proposals vs Contracts) by checking where the folder exists
    safe_folder = os.path.basename(folder_name)
    database_proposal_id = (request.args.get('proposal_id') or '').strip()
    if PROPOSAL_DATABASE_SOURCE_ENABLED and not database_proposal_id:
        try:
            database_proposal = (
                get_proposal_tracking_store().get_management_proposal_by_folder(
                    safe_folder
                )
            )
            if database_proposal:
                database_proposal_id = str(database_proposal.get("id") or "")
        except (ContactStoreError, TenantAuthenticationError):
            database_proposal_id = ""
    proposals_path = resolve_open_proposal_folder(safe_folder)
    contracts_path = os.path.join(CONTRACTS_DIR, safe_folder)

    if proposals_path and os.path.isdir(proposals_path):
        folder_path = proposals_path
    elif os.path.isdir(contracts_path):
        folder_path = contracts_path
    else:
        return f"Folder not found in open proposal directories or CONTRACTS_DIR: {safe_folder}", 404

    # Find the first file in the folder that starts with 'Profit Summary'
    profit_files = [f for f in os.listdir(folder_path) if f.startswith("Profit Summary") and f.endswith((".xlsm", ".xlsx"))]
    if not profit_files:
        return f"No Profit Summary file found in folder: {folder_name}"

    file_path = os.path.join(folder_path, profit_files[0])

    # Read the Excel file using openpyxl with cached values (data_only=True)
    try:
        wb_import = load_workbook(file_path, data_only=True)
        ws = wb_import.worksheets[0]
    except Exception as _e:
        return f"Unable to read Profit Summary with openpyxl: {_e}", 500
    cell_map = detect_profit_summary_cell_map(ws)

    def _cell(addr, default=None):
        try:
            v = ws[addr].value
            return v if v is not None else default
        except Exception:
            return default

    def _field(field, default=None):
        cell = cell_map.get(field)
        if not cell:
            return default
        return _cell(cell, default)

    hidden_values = read_hidden_sheet_values(wb_import, default=0)

    # Read Proposal Note and Language
    _proposal_note_import = _field("proposal_note", "")
    _proposal_language_import = _field("proposal_language", "")
    _repair_costs_import = _field("repair_costs_total", 0)
    _travel_total_import = _field("travel_total", 0)
    _include_travel_import = include_travel_from_travel_total(_travel_total_import)
    _adjusted_coverage_import = infer_adjusted_spread_rate(
        wb_import,
        str(_field("product", "") or ""),
        str(_field("current_roof", "") or ""),
    )
    try:
        _repair_costs_import_num = float(str(_repair_costs_import or 0).replace("$", "").replace(",", "").strip() or 0)
    except Exception:
        _repair_costs_import_num = 0.0
    if _repair_costs_import_num > 0:
        _proposal_language_import = REPAIR_COSTS_PROPOSAL_LANGUAGE

    data = {
        "flat_roof_squares": hidden_values.get("flat_roof_squares", 0),
        "wall_squares": hidden_values.get("wall_squares", 0),
        "adjusted_coverage": _adjusted_coverage_import,
        "squares": _field("squares", 0),
        "product": _field("product", ""),
        "price_per_sq_10": _field("price_per_sq_10", 0),
        "total_price_10": _field("total_price_10", 0),
        "current_roof": _field("current_roof", ""),
        "warranty_incl": _field("warranty_incl", "No"),
        "include_travel": _include_travel_import,
        "price_per_sq_15": _field("price_per_sq_15", 0),
        "total_price_15": _field("total_price_15", 0),
        "labor_days": _field("labor_days", None),
        "price_per_sq_20": _field("price_per_sq_20", 0),
        "total_price_20": _field("total_price_20", 0),
        "includes_text": _proposal_language_import,
        "proposal_language": _proposal_language_import,
        "submitted_by": _field("submitted_by", ""),
        "previous_submitted_by": _field("submitted_by", ""),
        "silicone_units_10": _field("silicone_units_10", 0),
        "silicone_units_15": _field("silicone_units_15", 0),
        "silicone_units_20": _field("silicone_units_20", 0),
        "silicone_price": _field("silicone_price", 0),
        "silicone_total": _field("silicone_total", 0),
        "silicone_15_total": _field("silicone_15_total", 0),
        "silicone_20_total": _field("silicone_20_total", 0),
        "gaco_patch_units": _field("gaco_patch_units", None),
        "gaco_patch_price": _field("gaco_patch_price", 0),
        "gaco_patch_total": _field("gaco_patch_total", 0),
        "bleed_trap_units": _field("bleed_trap_units", 0),
        "bleed_trap_price": _field("bleed_trap_price", 0),
        "bleed_trap_total": _field("bleed_trap_total", 0),
        "gaco_e5320_units": _field("gaco_e5320_units", 0),
        "gaco_e5320_price": _field("gaco_e5320_price", 0),
        "gaco_e5320_total": _field("gaco_e5320_total", 0),
        "sw_1flash_units": _field("sw_1flash_units", 0),
        "sw_1flash_price": _field("sw_1flash_price", 0),
        "sw_1flash_total": _field("sw_1flash_total", 0),
        "sw_bleed_block_units": _field("sw_bleed_block_units", 0),
        "sw_bleed_block_price": _field("sw_bleed_block_price", 0),
        "sw_bleed_block_total": _field("sw_bleed_block_total", 0),
        "drainage_mat_units": _field("drainage_mat_units", 0),
        "drainage_mat_price": _field("drainage_mat_price", 0),
        "drainage_mat_total": _field("drainage_mat_total", 0),
        "foam_units": _field("foam_units", None),
        "foam_price": _field("foam_price", 0),
        "foam_total": _field("foam_total", 0),
        "rfc_labor_price": _field("rfc_labor_price", 0),
        "rfc_labor_total": _field("rfc_labor_total", 0),
        "scarifying_total": _field("scarifying_total", 0),
        "pcs_labor_price": _field("pcs_labor_price", 0),
        "pcs_labor_total": _field("pcs_labor_total", 0),
        "travel_total": _travel_total_import,
        "repair_costs_total": _repair_costs_import,
        "warranty_10_total": _field("warranty_10_total", 0),
        "warranty_15_total": _field("warranty_15_total", 0),
        "warranty_20_total": _field("warranty_20_total", 0),
        "office_fee_total": _field("office_fee_total", 0),
        "office_fee_15_total": _field("office_fee_15_total", 0),
        "office_fee_20_total": _field("office_fee_20_total", 0),
        "total_cost": _field("total_cost", 0),
        "total_cost_15": _field("total_cost_15", 0),
        "total_cost_20": _field("total_cost_20", 0),
        "pcs_profit": _field("pcs_profit", 0),
        "pcs_profit_15": _field("pcs_profit_15", 0),
        "pcs_profit_20": _field("pcs_profit_20", 0),
        "profit_pct": _field("profit_pct", 0),
        "profit_pct_15": _field("profit_pct_15", 0),
        "profit_pct_20": _field("profit_pct_20", 0),
        "daily_profit": _field("daily_profit", 0),
        "daily_profit_15": _field("daily_profit_15", 0),
        "daily_profit_20": _field("daily_profit_20", 0),
        "profit_share": _field("profit_share", 0),
        "profit_share_15": _field("profit_share_15", 0),
        "profit_share_20": _field("profit_share_20", 0),
        "commission_amt": _field("commission_amt", 0),
        "commission_amt_15": _field("commission_amt_15", 0),
        "commission_amt_20": _field("commission_amt_20", 0),
        "customer_name": _field("customer_name", ""),
        "pcs_or_roofer_ind": _field("pcs_or_roofer_ind", ""),    
        "street_address": _field("street_address", ""),
        "city": _field("city", ""),
        "state": _field("state", ""),
        "zip_code": _field("zip_code", ""),
        "lead": _field("lead", ""),
        "proposal_note": _proposal_note_import,
    }

    pcs_or_roofer_ind = str(data.get("pcs_or_roofer_ind") or "").strip()

    # --- Fallback: if key computed fields are missing (no cached values), compute them for display ---
    def _to_float(v, d=0.0):
        try:
            return float(v)
        except Exception:
            return d

    # Consider pcs_profit/profit_pct/daily_profit/profit_share as the key set
    missing_core = any(
        (data.get(k) is None) for k in ("pcs_profit", "profit_pct", "daily_profit", "profit_share")
    )

    if missing_core:
        # Derive office fee % and commission % from Submitted By
        submitted_by_import = str(data.get("submitted_by") or "").strip()
        office_fee_pct_import = office_fee_pct_for_submitter(submitted_by_import)
        commission_pct_import = commission_pct_for_submitter(submitted_by_import)

        # Prepare inputs for calculation_routine using the imported sheet values
        calc_result = calculation_routine(
            squares=_to_float(data.get("squares"), 0.0),
            product=str(data.get("product") or ""),
            roof_type=str(data.get("current_roof") or ""),
            labor_days=int(_to_float(data.get("labor_days"), 0)),
            warranty_incl=str(data.get("warranty_incl") or "No"),
            include_travel=str(data.get("include_travel") or "No"),
            price_per_sq_10=_to_float(data.get("price_per_sq_10"), 0.0),
            commission_pct=_to_float(commission_pct_import, 0.0),
            submitted_by=submitted_by_import,
            previous_submitted_by=submitted_by_import,
            office_fee_pct=_to_float(office_fee_pct_import, 0.0),
            adjusted_coverage=_to_float(data.get("adjusted_coverage"), 0.0),
            silicone_units_10=_to_float(data.get("silicone_units_10"), 0.0),
            silicone_price=_to_float(data.get("silicone_price"), 0.0),
            gaco_patch_units=_to_float(data.get("gaco_patch_units"), 0.0),
            gaco_patch_price=_to_float(data.get("gaco_patch_price"), 0.0),
            gaco_e5320_units=_to_float(data.get("gaco_e5320_units"), 0.0),
            gaco_e5320_price=_to_float(data.get("gaco_e5320_price"), GACO_E5320_PRICE),
            sw_1flash_units=_to_float(data.get("sw_1flash_units"), 0.0),
            sw_1flash_price=_to_float(data.get("sw_1flash_price"), 0.0),
            bleed_trap_units=_to_float(data.get("bleed_trap_units"), 0.0),
            bleed_trap_price=_to_float(data.get("bleed_trap_price"), 0.0),
            sw_bleed_block_units=_to_float(data.get("sw_bleed_block_units"), 0.0),
            sw_bleed_block_price=_to_float(data.get("sw_bleed_block_price"), 0.0),
            drainage_mat_units=_to_float(data.get("drainage_mat_units"), 0.0),
            drainage_mat_price=_to_float(data.get("drainage_mat_price"), 0.0),
            foam_units=_to_float(data.get("foam_units"), 0.0),
            foam_price=_to_float(data.get("foam_price"), 0.0),
            rfc_labor_price=_to_float(data.get("rfc_labor_price"), 0.0),
            pcs_labor_price=_to_float(data.get("pcs_labor_price"), 0.0),
            scarifying_total=_to_float(data.get("scarifying_total"), 0.0),
            travel_total=_to_float(data.get("travel_total"), 0.0),
            repair_costs_total=_to_float(data.get("repair_costs_total"), 0.0),
            previous_squares=_to_float(data.get("squares"), 0.0),
            previous_roof_type=str(data.get("current_roof") or ""),
            previous_product=str(data.get("product") or ""),
            previous_adjusted_coverage=_to_float(data.get("adjusted_coverage"), 0.0),
            previous_silicone_units_10=_to_float(data.get("silicone_units_10"), 0.0),
            proposal_note=str(data.get("proposal_note") or ""),
            pcs_or_roofer_ind=pcs_or_roofer_ind,
            previous_pcs_or_roofer_ind=pcs_or_roofer_ind,
            previous_include_travel=str(data.get("include_travel") or "No"),
            previous_calc_travel_total=data.get("calc_travel_total") or 0,
        )

        # Overlay computed outputs onto data for display without writing back to the file
        overlay_keys = [
            "labor_days", "price_per_sq_10", "price_per_sq_15", "price_per_sq_20",
            "total_price_10", "total_price_15", "total_price_20",
            "silicone_units_10", "silicone_units_15", "silicone_units_20", "silicone_price", "silicone_total", "silicone_15_total", "silicone_20_total",
            "gaco_patch_units", "gaco_patch_price", "gaco_patch_total",
            "bleed_trap_units", "bleed_trap_price", "bleed_trap_total",
            "gaco_e5320_units", "gaco_e5320_price", "gaco_e5320_total",
            "sw_1flash_units", "sw_1flash_price", "sw_1flash_total",
            "sw_bleed_block_units", "sw_bleed_block_price", "sw_bleed_block_total",
            "drainage_mat_units", "drainage_mat_price", "drainage_mat_total",
            "foam_units", "foam_price", "foam_total",
            "rfc_labor_price", "rfc_labor_total", "pcs_labor_price", "pcs_labor_total",
            "scarifying_total", "travel_total", "repair_costs_total",
            "office_fee_total", "office_fee_15_total", "office_fee_20_total",
            "pcs_profit", "pcs_profit_15", "pcs_profit_20",
            "profit_pct", "profit_pct_15", "profit_pct_20",
            "daily_profit", "daily_profit_15", "daily_profit_20",
            "profit_share", "profit_share_15", "profit_share_20",
            "warranty_10_total", "warranty_15_total", "warranty_20_total",
            "coverage_10", "coverage_15", "coverage_20",
            "commission_amt", "commission_amt_15", "commission_amt_20",
            "commission_pct", "total_cost", "total_cost_15", "total_cost_20", "warranty_incl",
            "office_fee_pct", "adjusted_coverage",
        ]
        for k in overlay_keys:
            if k in calc_result:
                data[k] = calc_result[k]

    # --- Ensure Labor Days is visible for existing files with a formula but no cached value ---
    try:
        ld_val = data.get("labor_days")
        # Treat None/blank/zero as missing for display
        _ld_missing = (ld_val is None) or (str(ld_val).strip() == "") or (float(ld_val) == 0.0)
    except Exception:
        _ld_missing = True
    if _ld_missing:
        try:
            _squares_disp = float(data.get("squares") or 0.0)
        except Exception:
            _squares_disp = 0.0
        _roof_disp = str(data.get("current_roof") or "")
        if _roof_disp in ("Ballasted 60 mil", "Ballasted 45 mil"):
            _calc_ld_disp = int(math.ceil(_squares_disp / 30.0))
        elif _roof_disp == "Rock/Foam/Coat":
            _calc_ld_disp = int(math.ceil(_squares_disp / 75.0))
        else:
            _calc_ld_disp = int(math.ceil(_squares_disp / 45.0))
        data["labor_days"] = _calc_ld_disp
        data["calc_labor_days"] = _calc_ld_disp
        data["ov_labor_days"] = False

    # --- Ensure Silicone Units 10 is visible when the cell contains a formula without cached value ---
    try:
        su10_val = data.get("silicone_units_10")
        _su_missing = (su10_val is None) or (str(su10_val).strip() == "") or (float(su10_val) == 0.0)
    except Exception:
        _su_missing = True
    if _su_missing:
        try:
            _squares_disp = float(data.get("squares") or 0.0)
        except Exception:
            _squares_disp = 0.0
        _prod_disp = str(data.get("product") or "")
        _roof_disp = str(data.get("current_roof") or "")
        try:
            _adj_cov_disp = float(data.get("adjusted_coverage") or 0.0)
        except Exception:
            _adj_cov_disp = 0.0
        eff_cov_disp = adjusted_coverage_rates(_prod_disp, _roof_disp, _adj_cov_disp).get(10, 0.0)
        _calc_su10_disp = int(math.ceil((_squares_disp / 5.0) * eff_cov_disp))
        data["silicone_units_10"] = _calc_su10_disp
        data["calc_silicone_units_10"] = _calc_su10_disp
        data["ov_silicone_units_10"] = False

    # Ensure required keys exist for the template & triggers (Excel import init only)
    data.setdefault("coverage_10", 0)
    data.setdefault("coverage_15", 0)
    data.setdefault("coverage_20", 0)
    data.setdefault("adjusted_coverage", 0)

    # Initialize previous_* to current for first round-trip after Excel import
    data["previous_roof_type"] = str(data.get("current_roof") or "")
    data["previous_product"] = str(data.get("product") or "")
    try:
        data["previous_squares"] = float(data.get("squares") or 0)
    except Exception:
        data["previous_squares"] = 0.0
    try:
        data["previous_adjusted_coverage"] = float(data.get("adjusted_coverage") or 0)
    except Exception:
        data["previous_adjusted_coverage"] = 0.0

    # Derive Office Fee % and Commission % from Submitted By (Excel import should not overwrite with wrong cell)
    submitted_by_import = str(data.get("submitted_by") or "").strip()
    office_fee_pct_import = office_fee_pct_for_submitter(submitted_by_import)
    data["office_fee_pct"] = office_fee_pct_import

    data["commission_pct"] = commission_pct_for_submitter(submitted_by_import)

    # Recompute Office Fee total from the pre-office-fee subtotal so it aligns with P3 = subtotal + E25.
    office_fee_subtotal_10 = (
        (_to_float(data.get("squares"), 0.0) * _to_float(data.get("price_per_sq_10"), 0.0))
        + _to_float(data.get("warranty_10_total"), 0.0)
        + _to_float(data.get("travel_total"), 0.0)
        + _to_float(data.get("repair_costs_total"), 0.0)
    )
    data["office_fee_total"] = excel_round(office_fee_subtotal_10 * data["office_fee_pct"], 0)

    # Carry read-only flag through GET round-trips, supporting both new and legacy formats
    read_only_param = request.args.get('read_only')
    if read_only_param is not None:
        readonly = (read_only_param.strip().lower() == 'yes')
    else:
        readonly = (request.args.get('readonly') == '1')

    # --- Display fallbacks for fields saved as formulas (no cached value yet) ---
    try:
        # If the GET route has already prepared `data`, we reuse it here.
        _submitted_by_disp = str(data.get("submitted_by") or "")
        _office_fee_pct_disp = office_fee_pct_for_submitter(_submitted_by_disp)
        _commission_pct_disp = commission_pct_for_submitter(_submitted_by_disp)

        # Compute display values without resetting overrides (previous_* = current)
        calc_disp = calculation_routine(
            squares=float(data.get("squares") or 0.0),
            product=str(data.get("product") or ""),
            roof_type=str(data.get("current_roof") or ""),
            labor_days=float(data.get("labor_days") or 0.0),
            warranty_incl=str(data.get("warranty_incl") or "No"),
            include_travel=str(data.get("include_travel") or "No"),
            price_per_sq_10=float(data.get("price_per_sq_10") or 0.0),
            commission_pct=float(_commission_pct_disp),
            submitted_by=_submitted_by_disp,
            previous_submitted_by=_submitted_by_disp,
            office_fee_pct=float(_office_fee_pct_disp),
            adjusted_coverage=float(data.get("adjusted_coverage") or 0.0),
            silicone_units_10=float(data.get("silicone_units_10") or 0.0),
            silicone_price=float(data.get("silicone_price") or 0.0),
            gaco_patch_units=float(data.get("gaco_patch_units") or 0.0),
            gaco_patch_price=float(data.get("gaco_patch_price") or 0.0),
            gaco_e5320_units=float(data.get("gaco_e5320_units") or 0.0),
            gaco_e5320_price=float(data.get("gaco_e5320_price") or GACO_E5320_PRICE),
            sw_1flash_units=float(data.get("sw_1flash_units") or 0.0),
            sw_1flash_price=float(data.get("sw_1flash_price") or 0.0),
            bleed_trap_units=float(data.get("bleed_trap_units") or 0.0),
            bleed_trap_price=float(data.get("bleed_trap_price") or 0.0),
            sw_bleed_block_units=float(data.get("sw_bleed_block_units") or 0.0),
            sw_bleed_block_price=float(data.get("sw_bleed_block_price") or 0.0),
            drainage_mat_units=float(data.get("drainage_mat_units") or 0.0),
            drainage_mat_price=float(data.get("drainage_mat_price") or 0.0),
            foam_units=float(data.get("foam_units") or 0.0),
            foam_price=float(data.get("foam_price") or 0.0),
            rfc_labor_price=float(data.get("rfc_labor_price") or 0.0),
            pcs_labor_price=float(data.get("pcs_labor_price") or 0.0),
            scarifying_total=float(data.get("scarifying_total") or 0.0),
            travel_total=float(data.get("travel_total") or 0.0),
            repair_costs_total=float(data.get("repair_costs_total") or 0.0),
            previous_squares=float(data.get("squares") or 0.0),
            previous_roof_type=str(data.get("current_roof") or ""),
            previous_product=str(data.get("product") or ""),
            previous_adjusted_coverage=float(data.get("adjusted_coverage") or 0.0),
            previous_silicone_units_10=float(data.get("silicone_units_10") or 0.0),
            proposal_note=str(data.get("proposal_note") or ""),
            pcs_or_roofer_ind=pcs_or_roofer_ind,
            previous_pcs_or_roofer_ind=pcs_or_roofer_ind,
            previous_include_travel=str(data.get("include_travel") or "No"),
            previous_calc_travel_total=data.get("calc_travel_total") or 0,
        )

        # These are saved as formulas; show calculated values if Excel hasn't cached them
        _fallback_fields = [
            # Inputs that may be stored as formulas (or blank when N/A) but must show a number on screen
            "labor_days",
            "silicone_units_10","silicone_units_15","silicone_units_20","gaco_patch_units","bleed_trap_units","gaco_e5320_units",
            "sw_1flash_units","sw_bleed_block_units","drainage_mat_units","foam_units",

            # Unit prices that may be stored as formulas (must display a number)
            "silicone_price","gaco_patch_price","bleed_trap_price","gaco_e5320_price",
            "sw_1flash_price","sw_bleed_block_price","drainage_mat_price",
            "foam_price","rfc_labor_price","pcs_labor_price",

            # Existing entries (keep totals/prices/etc.)
            "price_per_sq_10","price_per_sq_15","price_per_sq_20",
            "total_price_10","total_price_15","total_price_20",
            "total_cost","total_cost_15","total_cost_20",
            "warranty_10_total","office_fee_total","office_fee_15_total","office_fee_20_total",
            "silicone_total","silicone_15_total","silicone_20_total","gaco_patch_total","bleed_trap_total","gaco_e5320_total",
            "sw_1flash_total","sw_bleed_block_total","drainage_mat_total",
            "foam_total","rfc_labor_total","pcs_labor_total",
            "commission_amt","commission_amt_15","commission_amt_20",
            "profit_share","profit_share_15","profit_share_20",
            "daily_profit","daily_profit_15","daily_profit_20",
            "profit_pct","profit_pct_15","profit_pct_20",
            "pcs_profit","pcs_profit_15","pcs_profit_20"
        ]

        def _is_missing(v):
            try:
                if v is None:
                    return True
                if isinstance(v, float) and math.isnan(v):
                    return True
                return float(v) == 0.0
            except Exception:
                # Strings (formulas) won't be missing here; only numbers can be missing
                return False

        for _k in _fallback_fields:
            if _is_missing(data.get(_k)):
                if _k in calc_disp:
                    data[_k] = calc_disp[_k]
        data["calc_travel_total"] = calc_disp.get("calc_travel_total", data.get("calc_travel_total", 0))
        data["previous_calc_travel_total"] = data["calc_travel_total"]
        data["previous_include_travel"] = data.get("include_travel") or "No"
    except Exception as _fallback_e:
        # Non-fatal: if anything goes wrong, skip the fallback
        _safe_debug(f"[DEBUG] display fallback (GET) error: {_fallback_e}")

    # --- Merge evaluated display data (prefers ghost cells) before rendering (GET) ---
    try:
        _folder_path = resolve_open_proposal_folder(folder_name) or os.path.join(PROPOSAL_TEMP_DIR, folder_name)
        data = merge_display_fallbacks(data, _folder_path, folder_name)
        try:
            with open(APP_ERROR_LOG, "a", encoding="utf-8") as _f:
                _f.write(
                    f"\n[PD GET] ghost merge {folder_name}: sw1(u/p/t)="
                    f"({data.get('sw_1flash_units')},{data.get('sw_1flash_price')},{data.get('sw_1flash_total')}) "
                    f"sil10={data.get('silicone_units_10')}\n"
                )
        except Exception:
            pass
    except Exception:
        pass

    return render_template(
        "proposal_details.html",
        data=data,
        **data,
        customer_organization_names=proposal_customer_organization_names(),
        folder_name=folder_name,
        readonly=readonly,
        is_blank=False,
        database_proposal_id=database_proposal_id or None,
    )

        
def replace_placeholder_blocks(doc, replacements):
    def replace_text_in_block(paragraph_or_cell):
        full_text = ''.join(run.text for run in paragraph_or_cell.runs)
        for key, val in replacements.items():
            full_text = full_text.replace(key, str(val))
        for run in paragraph_or_cell.runs:
            run.text = ''
        if paragraph_or_cell.runs:
            paragraph_or_cell.runs[0].text = full_text

    for para in doc.paragraphs:
        if any(key in para.text for key in replacements):
            replace_text_in_block(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if any(key in para.text for key in replacements):
                        replace_text_in_block(para)


if __name__ == "__main__":
    import os

    # Default to port 5050 so we never collide with things that grab 5000.
    # You can still override with: PORT=5088 python pcs_proposal_web.py
    port = int(os.environ.get("PORT", 5050))
    app.run(host="127.0.0.1", port=port, debug=False, use_reloader=False)
